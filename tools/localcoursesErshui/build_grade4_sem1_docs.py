from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from grade4_sem1_lesson_data import SOURCES, UNITS


ROOT = Path.cwd()
OUT_DIR = ROOT / "在地課程4年級上學期教案"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# compact_reference_guide preset + named lesson-title override.
FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft JhengHei"
BODY_SIZE = 11
NAVY = "1F4E79"
TEAL = "0F6B6D"
DARK = "24313A"
MUTED = "5B6B77"
WHITE = "FFFFFF"
HEADER_FILL = "E8EEF5"
PALE_TEAL = "E9F5F4"
PALE_AMBER = "FFF2D8"
PALE_BLUE = "EFF5FA"
PALE_GREY = "F5F7F8"
BORDER = "B8C5CE"
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=BODY_SIZE, bold=False, color=DARK, italic=False):
    run.font.name = FONT_LATIN
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)
    return run


def set_paragraph_format(paragraph, before=0, after=4, line=1.25, keep=False):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep
    pf.widow_control = True
    return paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != PAGE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {PAGE_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def set_repeat_table_header(row):
    repeat_header(row)
    prevent_row_split(row)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(paragraph, after=0, line=1.0)
    set_run_font(paragraph.add_run("第 "), 8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    set_run_font(paragraph.add_run(" 頁"), 8.5, color=MUTED)


def add_hyperlink(paragraph, text, url, color=TEAL, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)
    rpr.append(rfonts)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color)
    rpr.append(col)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    rpr.append(size)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_bottom_border(paragraph, color=TEAL, size="12"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "5")
        node.set(qn("w:color"), BORDER)
        borders.append(node)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_format(p, before=13, after=6, line=1.0, keep=True)
        set_run_font(p.add_run(text), 16, True, NAVY)
        add_bottom_border(p, TEAL, "14")
    elif level == 2:
        set_paragraph_format(p, before=10, after=5, line=1.05, keep=True)
        set_run_font(p.add_run(text), 13, True, TEAL)
    else:
        set_paragraph_format(p, before=7, after=4, line=1.1, keep=True)
        set_run_font(p.add_run(text), 12, True, NAVY)
    return p


def add_lesson_title(doc, text):
    # Named override requested by the user: visibly larger and separated.
    p = doc.add_paragraph(style="Lesson Title")
    set_paragraph_format(p, before=0, after=12, line=1.0, keep=True)
    run = p.add_run(text)
    set_run_font(run, 18, True, NAVY)
    add_bottom_border(p, TEAL, "22")
    return p


def add_body(doc, text, bold_label=None, fill=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=2 if fill else 0, after=7 if fill else 5, line=1.25)
    if fill:
        shade_paragraph(p, fill)
    if bold_label:
        set_run_font(p.add_run(bold_label), BODY_SIZE, True, NAVY)
    set_run_font(p.add_run(text), BODY_SIZE)
    return p


def add_real_bullet(container, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = container.add_paragraph(style=style)
    set_paragraph_format(p, after=3, line=1.22)
    set_run_font(p.add_run(text), BODY_SIZE)
    return p


def add_text_cell(cell, text, size=9.2, bold=False, color=DARK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    set_paragraph_format(p, after=0, line=1.16)
    set_run_font(p.add_run(str(text)), size, bold, color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_margins(cell)
    set_cell_border(cell)
    return p


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [1900, 7460]
    for idx, value in enumerate(("課前項目", "教學配置內容")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        add_text_cell(
            cell,
            value,
            size=9.4,
            bold=True,
            color=WHITE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        set_cell_shading(cells[0], HEADER_FILL)
        add_text_cell(cells[0], label, size=9.4, bold=True, color=NAVY)
        add_text_cell(cells[1], value, size=9.3)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_grid_table(doc, headers, rows, widths, font_size=8.8, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        add_text_cell(cell, header, size=font_size, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(row):
            if r_idx % 2 == 1:
                set_cell_shading(cells[idx], PALE_GREY)
            add_text_cell(cells[idx], value, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_resource_table(doc, source_keys):
    rows = []
    for key in source_keys:
        src = SOURCES[key]
        rows.append((src["label"], src["note"], src["url"]))
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["資源名稱", "教學用途與使用提醒", "網址"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        add_text_cell(cell, header, size=8.9, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for row_idx, (label, note, url) in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        if row_idx % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, PALE_GREY)
        add_text_cell(cells[0], label, size=8.7)
        add_text_cell(cells[1], note, size=8.7)
        p = add_text_cell(cells[2], "", size=8.2)
        add_hyperlink(p, "開啟來源", url)
        p.add_run().add_break()
        set_run_font(p.add_run(url), 7.5, color=MUTED)
    set_table_geometry(table, [2350, 3810, 3200])
    doc.add_paragraph()
    return table


def format_curriculum_entries(entries):
    return "\n".join(f"{code}　{text}" for code, text in entries)


def add_lesson(doc, lesson):
    doc.add_page_break()
    add_lesson_title(doc, lesson["title"])
    add_body(doc, lesson["focus"], bold_label="本節焦點｜", fill=PALE_TEAL)
    add_heading(doc, "課前教學配置", 2)
    prep_rows = [
        ("主要英語口說／句型", lesson["patterns"]),
        ("句型來源與使用界線", lesson["pattern_source"]),
        ("生活用語", lesson["daily"]),
        ("核心英文字詞", lesson["core"]),
        ("相關英文字詞＋中文", lesson["related"]),
        ("數位教材", lesson["digital"]),
        ("手作教材", lesson["handmade"]),
        ("本節學習單", lesson["worksheet"]),
    ]
    add_kv_table(doc, prep_rows)
    add_heading(doc, "40分鐘詳細教學流程", 2)
    flow_rows = []
    for item in lesson["stages"]:
        flow_rows.append([
            f"{item['name']}\n{item['minutes']}分鐘",
            item["title"],
            item["teacher"],
            item["students"],
            item["assessment"],
        ])
    add_grid_table(
        doc,
        ["階段／時間", "主要活動標題", "教師如何教與如何提問", "學生如何學與如何互動", "形成性評量／成果"],
        flow_rows,
        [900, 1450, 3000, 2570, 1440],
        font_size=8.4,
    )
    total = sum(item["minutes"] for item in lesson["stages"])
    if total != 40:
        raise ValueError(f"{lesson['title']} totals {total} minutes")
    add_body(
        doc,
        "本節英語口說以「能在地方學習任務中聽懂、指出、輪流說」為達成標準；不把專有名詞拼寫或尚未正式教過的文法列為必要通過條件。",
        bold_label="四年級適切性檢核｜",
        fill=PALE_AMBER,
    )


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(BODY_SIZE)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(BODY_SIZE)
        srpr = style.element.get_or_add_rPr()
        srfonts = srpr.rFonts
        if srfonts is None:
            srfonts = OxmlElement("w:rFonts")
            srpr.insert(0, srfonts)
        srfonts.set(qn("w:ascii"), FONT_LATIN)
        srfonts.set(qn("w:hAnsi"), FONT_LATIN)
        srfonts.set(qn("w:eastAsia"), FONT_CJK)
    if "Lesson Title" not in [s.name for s in styles]:
        lesson_style = styles.add_style("Lesson Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lesson_style = styles["Lesson Title"]
    lesson_style.font.name = FONT_LATIN
    lesson_style.font.size = Pt(18)
    lesson_style.font.bold = True
    lesson_style.font.color.rgb = RGBColor.from_string(NAVY)


def configure_section(doc, unit):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(header_p, after=0, line=1.0)
    set_run_font(
        header_p.add_run(f"二水國小四年級在地課程｜上學期單元{unit['unit_no']}"),
        8.5,
        color=MUTED,
    )
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc, unit):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=10, after=8, line=1.0)
    set_run_font(p.add_run("LOCAL RAILWAY × ENGLISH"), 10, True, TEAL)
    p = doc.add_paragraph()
    set_paragraph_format(p, before=30, after=10, line=1.0)
    set_run_font(p.add_run(f"四年級在地課程教案\n〈{unit['title']}〉"), 28, True, NAVY)
    p = doc.add_paragraph()
    set_paragraph_format(p, after=20, line=1.2)
    set_run_font(p.add_run("英語領域 × 社會領域 × 綜合活動領域"), 14, True, TEAL)
    add_bottom_border(p, TEAL, "24")
    metrics = [
        ["教學期程", unit["weeks"], "節數", f"{unit['lessons_count']}節"],
        ["每節時間", "40分鐘", "對象", "二水國小四年級"],
        ["課程主軸", "原在地課程活動不做重大變更", "版本", "2026-07-26"],
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for idx, value in enumerate(("項目", "內容", "項目", "內容")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        add_text_cell(
            cell,
            value,
            9.2,
            True,
            WHITE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(metrics):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(row):
            if idx in (0, 2):
                set_cell_shading(cells[idx], HEADER_FILL)
                add_text_cell(cells[idx], value, 9.5, True, NAVY)
            else:
                add_text_cell(cells[idx], value, 9.5)
    set_table_geometry(table, [1350, 3330, 1350, 3330])
    doc.add_paragraph()
    add_body(
        doc,
        unit["position"],
        bold_label="單元定位｜",
        fill=PALE_TEAL,
    )
    add_body(
        doc,
        "本教案先以地方知識與原活動為主，再選用學生已學或當期正在學的英語。專有名詞可看圖說；完整地方史、資料查證與反思主要以中文表達。",
        bold_label="設計底線｜",
        fill=PALE_AMBER,
    )
    doc.add_page_break()


def build_unit(unit):
    doc = Document()
    configure_styles(doc)
    configure_section(doc, unit)
    add_cover(doc, unit)

    add_heading(doc, "一、單元定位與核心成果", 1)
    add_body(doc, unit["position"])
    add_heading(doc, "學習目標", 2)
    for goal in unit["goals"]:
        add_real_bullet(doc, goal)
    add_body(doc, unit["core_output"], bold_label="核心成果｜", fill=PALE_BLUE)
    add_body(doc, unit["alignment_note"], bold_label="英語配對判準｜", fill=PALE_AMBER)

    add_heading(doc, "二、英語部定／彈性課程融入比對", 1)
    add_grid_table(
        doc,
        ["在地活動", "英語來源", "本單元採用語言", "採用理由與界線"],
        unit["english_match"],
        [1500, 1850, 3100, 2910],
        font_size=8.8,
    )
    add_body(
        doc,
        "彈性英語課程中的分組海報、資料比較、角色互動等學習方法可移入；主題字彙仍須以鐵道地方內容為準。若課本主題不相干，便不以『配進度』為由硬套。",
        bold_label="彈性課程運用｜",
        fill=PALE_TEAL,
    )

    add_heading(doc, "三、英語聽說評量手冊融入", 1)
    add_grid_table(
        doc,
        ["項目", "本單元選用內容", "教學與評量方式"],
        unit["manual_integration"],
        [1600, 3300, 4460],
        font_size=9.0,
    )

    add_heading(doc, "四、建議影片、繪本、照片與資料", 1)
    add_resource_table(doc, unit["source_keys"])
    add_body(
        doc,
        "所有網站、時刻、票價、營業與開放資訊都可能變動，教師須於授課前重新開啟官方頁確認。圖片只截取教學所需範圍並保留來源，不移除浮水印、不整頁複製受著作權保護的繪本。",
        bold_label="使用提醒｜",
        fill=PALE_AMBER,
    )

    add_heading(doc, "五、領域課綱對應", 1)
    add_grid_table(
        doc,
        ["領域", "學習表現（代碼＋完整敘述）", "學習內容（代碼＋完整敘述）", "在本單元中的具體證據"],
        [
            [
                "英語文",
                format_curriculum_entries(unit["curriculum"]["english"]["performance"]),
                format_curriculum_entries(unit["curriculum"]["english"]["content"]),
                unit["curriculum"]["english"]["evidence"],
            ],
            [
                "社會",
                format_curriculum_entries(unit["curriculum"]["social"]["performance"]),
                format_curriculum_entries(unit["curriculum"]["social"]["content"]),
                unit["curriculum"]["social"]["evidence"],
            ],
            [
                "綜合活動",
                format_curriculum_entries(unit["curriculum"]["comprehensive"]["performance"]),
                format_curriculum_entries(unit["curriculum"]["comprehensive"]["content"]),
                unit["curriculum"]["comprehensive"]["evidence"],
            ],
        ],
        [1000, 3100, 2800, 2460],
        font_size=8.5,
    )

    add_heading(doc, "六、逐節教學流程", 1)
    add_body(
        doc,
        f"本單元共{unit['lessons_count']}節。每節先列出英語口說、生活用語、核心與相關詞彙，以及數位／手作／學習單準備；其下再列完整40分鐘流程。",
        fill=PALE_BLUE,
    )
    for lesson in unit["lessons"]:
        add_lesson(doc, lesson)

    doc.add_page_break()
    add_heading(doc, "七、教材、圖卡、範例答案與學習單附錄", 1)
    for appendix in unit["appendices"]:
        add_heading(doc, appendix["title"], 2)
        for item in appendix["body"]:
            add_real_bullet(doc, item)

    add_heading(doc, "八、教師授課前最後檢核", 1)
    checklist = [
        "重新開啟官方來源，確認時刻、開放資訊與網頁是否仍有效；截圖標示擷取日期。",
        "確認英語主句型來自三年級已學內容或四年級上學期同步內容；支架語塊不列為未教文法測驗。",
        "確認每個英文詞都服務地方觀察、資料整理或口語互動；刪除與在地活動無關的主題句型。",
        "備妥紙本替代方案；數位設備失效時仍可用照片、圖卡與資料卡完成同一學習目標。",
        "學習單先由教師試作一次，確認四年級能在預定時間內完成，字體、圖例與書寫空間足夠。",
        "若不進行實地踏查，清楚稱為「數位觀察／資料觀察」，不讓學生誤以為影像等同親身現場經驗。",
    ]
    for item in checklist:
        add_real_bullet(doc, item)

    out_path = OUT_DIR / f"{unit['file_stub']}.docx"
    doc.save(out_path)
    return out_path


def main():
    outputs = []
    for unit in UNITS:
        outputs.append(build_unit(unit))
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
