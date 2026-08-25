from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "在地課程4年級上學期教案" / "02_第3-10週_坐火車趣集集.docx"
PROMPT_PATH = (
    ROOT
    / "在地課程4年級上學期教案"
    / "02_第3-10週_坐火車趣集集 教材"
    / "00_查證與清冊"
    / "prompts"
    / "L1_12_storyboard_flow_prompts_bilingual_v2.md"
)

BLUE = RGBColor(31, 78, 121)
BODY_FONT = "Microsoft JhengHei"
ENGLISH_FONT = "Comic Sans MS"


def parse_prompts(markdown: str) -> dict[str, dict[str, str]]:
    section_pattern = re.compile(
        r"(?ms)^## 分鏡 (?P<num>\d{2})｜(?P<title>[^\n]+)\n(?P<body>.*?)(?=^---\n\n## 分鏡 |\Z)"
    )
    result: dict[str, dict[str, str]] = {}
    for match in section_pattern.finditer(markdown):
        body = match.group("body")
        zh_match = re.search(
            r"(?ms)^### 中文版 Flow 動態提示詞\s*\n+(.*?)\n+### English Flow motion prompt",
            body,
        )
        en_match = re.search(
            r"(?ms)^### English Flow motion prompt\s*\n+(.*?)\n+後製字幕：",
            body,
        )
        subtitle_match = re.search(r"(?m)^後製字幕：(.*)$", body)
        narration_match = re.search(r"(?m)^後製旁白：(.*)$", body)
        if not all((zh_match, en_match, subtitle_match, narration_match)):
            raise ValueError(f"Prompt section {match.group('num')} is incomplete")
        result[match.group("num")] = {
            "title": match.group("title").strip(),
            "zh": " ".join(zh_match.group(1).split()),
            "en": " ".join(en_match.group(1).split()),
            "subtitle": subtitle_match.group(1).strip(),
            "narration": narration_match.group(1).strip(),
        }
    if set(result) != {f"{i:02d}" for i in range(1, 13)}:
        raise ValueError(f"Expected 12 storyboard sections, found {sorted(result)}")
    return result


def set_run_font(run, font_name: str, size_pt: float, *, bold: bool = False, color=None):
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), BODY_FONT if font_name == ENGLISH_FONT else font_name)


def clear_paragraph(paragraph: Paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def write_labeled_paragraph(
    paragraph: Paragraph,
    label: str,
    content: str,
    *,
    content_font: str = BODY_FONT,
    size_pt: float = 10.5,
):
    clear_paragraph(paragraph)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.keep_together = False

    label_run = paragraph.add_run(label)
    set_run_font(label_run, BODY_FONT, size_pt, bold=True, color=BLUE)
    content_run = paragraph.add_run(content)
    set_run_font(content_run, content_font, size_pt)


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if paragraph.style is not None:
        new_paragraph.style = paragraph.style
    return new_paragraph


def find_paragraph(document: Document, predicate, start: int = 0) -> Paragraph:
    for paragraph in document.paragraphs[start:]:
        if predicate(paragraph.text):
            return paragraph
    raise ValueError("Required paragraph was not found")


def replace_text_paragraph(paragraph: Paragraph, new_text: str):
    clear_paragraph(paragraph)
    run = paragraph.add_run(new_text)
    set_run_font(run, BODY_FONT, 10.5)


def update_document():
    prompts = parse_prompts(PROMPT_PATH.read_text(encoding="utf-8"))
    document = Document(DOCX_PATH)

    workflow = find_paragraph(
        document,
        lambda text: text.startswith("每張Flow提示詞皆完整寫入站名"),
    )
    replace_text_paragraph(
        workflow,
        "每張分鏡皆提供中文版與英文版 Flow 動態提示詞。Flow 只負責動態，不生成站名、字幕、旁白、對話框或鏡間轉場；"
        "需要的文字先正確烘焙於靜態首幀，精確字幕與旁白依各分鏡所列內容後製加入。",
    )

    recognition = find_paragraph(
        document,
        lambda text: text.startswith("試播時以暫停點檢核："),
    )
    replace_text_paragraph(
        recognition,
        "試播時以暫停點檢核：二水、集集、水里、車埕四個教學重點站均能在2秒內辨認；"
        "順序固定為二水→集集→水里→車埕；不得出現第五個節點、重複站名、亂碼或片尾自行換景。",
    )

    unified_prompt = find_paragraph(
        document,
        lambda text: text.startswith("12張分鏡統一靜態提示詞："),
    )
    replace_text_paragraph(
        unified_prompt,
        "12張分鏡統一靜態提示詞：16:9、1920×1080、四年級教學影片分鏡首幀。"
        "半寫實3D兒童動畫與溫暖繪本質感，忠實保留臺灣中部集集線的站體比例、月臺與地方景觀；"
        "同一列黃色列車、同一晴朗早晨、同一柔和色盤。路線圖只顯示四個教學重點站："
        "二水 Ershui、集集 Jiji、水里 Shuili、車埕 Checheng，順序固定且不得增加節點。"
        "所有站名與標題必須在送入 Flow 前烘焙於首幀，列為完全靜止區。"
        "分鏡01、03、12必須使用修正後的四站版首幀。禁止錯誤站序、重複站名、亂碼、簡體字、"
        "模型自行生成文字、片尾換景、人物進入軌道、浮水印與不相關地標。",
    )

    for num in [f"{i:02d}" for i in range(1, 13)]:
        item = prompts[num]
        heading = find_paragraph(
            document,
            lambda text, n=num: text.startswith(f"分鏡{n}｜"),
        )
        replace_text_paragraph(heading, f"分鏡{num}｜{item['title']}")
        heading.style = "Heading 3"

        paragraphs = document.paragraphs
        heading_index = next(
            i for i, paragraph in enumerate(paragraphs) if paragraph._element is heading._element
        )
        next_heading_index = next(
            (
                i
                for i in range(heading_index + 1, len(paragraphs))
                if paragraphs[i].text.startswith("分鏡") and "｜" in paragraphs[i].text
            ),
            len(paragraphs),
        )
        block = paragraphs[heading_index + 1 : next_heading_index]
        old_flow = next(
            p for p in block if p.text.startswith("個別Google Flow動態提示詞：")
        )
        old_post = next(p for p in block if p.text.startswith("後製備援（逐字照用）："))

        write_labeled_paragraph(
            old_flow,
            "中文版 Google Flow 動態提示詞：",
            item["zh"],
        )
        english_paragraph = insert_paragraph_after(old_flow)
        write_labeled_paragraph(
            english_paragraph,
            "English Google Flow motion prompt: ",
            item["en"],
            content_font=ENGLISH_FONT,
        )
        write_labeled_paragraph(old_post, "後製字幕：", item["subtitle"])
        narration_paragraph = insert_paragraph_after(old_post)
        write_labeled_paragraph(
            narration_paragraph,
            "後製旁白：",
            item["narration"],
        )

    temp_path = DOCX_PATH.with_name(DOCX_PATH.stem + "_flow_update.tmp.docx")
    backup_path = DOCX_PATH.with_name(
        DOCX_PATH.stem
        + "_更新Flow提示詞前備份_"
        + datetime.now().strftime("%Y%m%d-%H%M%S")
        + ".docx"
    )
    shutil.copy2(DOCX_PATH, backup_path)
    document.save(temp_path)
    temp_path.replace(DOCX_PATH)
    return backup_path


if __name__ == "__main__":
    backup = update_document()
    print(f"Updated: {DOCX_PATH}")
    print(f"Backup: {backup}")
