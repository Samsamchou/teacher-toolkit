# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DOCX_PATH = (
    PROJECT_ROOT
    / "在地課程4年級上學期教案"
    / "02_第3-10週_坐火車趣集集.docx"
)
START_MARKER = "第1節自製教材｜二水站暖身圖、90秒路線影片與三站閃卡"
END_MARKER = "八張學習單設計初稿與繪製提示詞"
QUERY_DATE = "2026-07-28"


SOURCES = [
    (
        "臺鐵官網首頁／路線與站點",
        "確認集集線順序為二水→源泉→濁水→龍泉→集集→水里→車埕；教材僅突出二水、集集、車埕，仍不可跳接或顛倒。",
        "https://www.railway.gov.tw/tra-tip-web/tip",
    ),
    (
        "臺鐵車站代碼資料",
        "確認二水3430、源泉3431、濁水3432、龍泉3433、集集3434、水里3435、車埕3436。",
        "https://www.railway.gov.tw/tra-tip-web/tip/tip00C/tipC21/view?subCode=8ae4cac3756b7b41017573dc65f31899",
    ),
    (
        "交通部：集集線全線復駛",
        "確認二水至車埕為完整區間；2026年資料顯示集集線已全線復駛。實際乘車仍須課前查最新公告與時刻。",
        "https://www.motc.gov.tw/ch/app/news_list/view?id=14&module=news&serno=c752d930-f8a2-4740-a67e-3115a45deb92",
    ),
    (
        "參山國家風景區：二水車站",
        "確認二水是集集支線起點、路線沿濁水溪谷北岸前往車埕，並可用山川、田野、綠色隧道說明沿線景觀。",
        "https://www.trimt-nsa.gov.tw/zh-tw/attraction/118/",
    ),
    (
        "南投旅遊網：鐵道主題旅遊",
        "交叉確認沿線七站與地方景觀；水里為重要轉乘點，車埕為終點。",
        "https://travel.nantou.gov.tw/theme-train-trip/",
    ),
    (
        "南投旅遊網：集集車站",
        "供第1、6節核對集集車站外觀與地方特色。",
        "https://travel.nantou.gov.tw/attractions/jiji-station/",
    ),
    (
        "南投旅遊網：車埕車站",
        "供第1、6節核對車埕終點、木業／木造景觀等特徵。",
        "https://travel.nantou.gov.tw/attractions/checheng-station/",
    ),
    (
        "臺鐵列車時刻查詢",
        "第2節模擬表不得冒充現行時刻；若教師延伸真實查詢，必須以此頁重新確認並標示查詢日期。",
        "https://www.railway.gov.tw/tra-tip-web/tip/tip001/tip112/gobytime?lang=zh_TW",
    ),
    (
        "臺鐵網路訂票說明",
        "第4節只模擬「選起點、選終點、選日期、查詢、選車次、確認摘要」，不蒐集身分資料、不付款、不成立真實訂票。",
        "https://www.railway.gov.tw/tra-tip-web/tip/tip00C/tipC21/view?subCode=8ae4cac3756b7b41017573dba4e4186f",
    ),
    (
        "臺鐵雙語服務用語",
        "可核對候車時不要超越黃線等月臺安全用語；學生版仍以短句與圖像呈現。",
        "https://www.railway.gov.tw/tra-tip-web/tip/tip00C/tipC20/view70",
    ),
    (
        "Google Flow：模型與支援功能",
        "確認Frames to Video可使用8秒片段，實際可用長度與功能仍依帳號、模型及當日介面為準。",
        "https://support.google.com/flow/answer/16352836?hl=en",
    ),
    (
        "Google Flow：Scenebuilder",
        "確認可排序、調整片段前後端、預覽與下載；本案以12段約8秒素材剪成90秒。",
        "https://support.google.com/labs/answer/16935718?hl=en",
    ),
    (
        "Google Flow官方提示技巧",
        "以固定角色／列車／色彩／構圖參考及Frames to Video維持跨鏡頭一致性。",
        "https://blog.google/innovation-and-ai/products/flow-video-tips/",
    ),
]


FLOW_COMMON_STILL = (
    "16:9、1920×1080、四年級教學影片分鏡首幀。半寫實3D兒童動畫與溫暖繪本質感，"
    "忠實保留臺灣中部集集線的窄軌鐵道、站體比例、月臺與地方景觀；同一列紅黃主色柴油列車、"
    "同一晴朗早晨、同一柔和色盤。畫面資訊分成前景人物／列車、中景站體或田野、後景地形。"
    "站名、路線線條、字幕與箭頭只保留乾淨安全區，文字後製，不讓生成模型自行亂寫。"
    "禁止：錯誤站序、把高鐵或捷運畫成集集線、左駕公路列車、虛構月臺、簡體字、浮水印、"
    "過度觀光化、人物站上軌道、人物跨越安全線。"
)

FLOW_CONTINUITY = (
    "固定Ingredients／參考素材：A. 同一列紅黃主色集集線柴油列車三視圖；"
    "B. 二水、集集、水里、車埕站體官方照片參考板；C. 白底紅線集集線路線圖與同一暖色3D風格板。"
    "所有鏡頭維持列車車頭、車窗數量、色彩、天空方向、路線圖符號與字幕安全區一致。"
)


FLOW_SHOTS = [
    {
        "id": "01",
        "time": "原始0–8秒；成片保留0–5秒",
        "title": "旅程開場：從二水出發",
        "station": "二水／集集／車埕（標題路線卡）",
        "still": "俯視簡化臺灣中部教學地圖，二水在左下、集集在中段、車埕在右上；三站以圓點標示，紅黃小火車停在二水，大片留白放標題。",
        "route": "紅色路線由二水圓點亮起，沿正確方向依序延伸到集集與車埕；中間站以小圓點帶過，不跳站。",
        "subtitle": "從二水出發，認識集集線",
        "narration": "今天，我們從二水出發，搭上集集線。",
        "audio": "輕快木琴開場，低音量車站環境聲；第1秒一聲柔和列車鈴。",
        "camera": "由地圖全景緩慢推近二水圓點，再小幅向右上平移。",
        "action": "小火車圖示在二水輕輕晃動，路線依序點亮。",
        "transition": "紅色路線變成下一鏡頭月臺旁的實體鋼軌，柔和溶接。",
    },
    {
        "id": "02",
        "time": "5–13秒",
        "title": "二水車站照片尋線索",
        "station": "二水站 Ershui Station",
        "still": "依官方現況照片重建二水車站正面與局部月臺，站名牌、鐵軌、月臺三個線索清楚但不加提示圈；一組四年級學生站在安全區觀察。",
        "route": "右下角小路線條只亮二水站，箭頭指向集集線方向。",
        "subtitle": "二水站：找找看站名牌、鐵軌、月臺",
        "narration": "這是二水站。你找到三個火車站線索了嗎？",
        "audio": "自然車站廣播遠景、鳥鳴與輕微風聲；不使用可辨識真實廣播內容。",
        "camera": "先看站體全景，再緩慢推近站名牌，停在學生可指認的構圖。",
        "action": "學生依序指向站名牌、鐵軌、月臺，人物始終在安全線後。",
        "transition": "鏡頭推近站名牌後，站名牌化成地圖上的二水標記。",
    },
    {
        "id": "03",
        "time": "13–21秒",
        "title": "認識支線起點與方向",
        "station": "二水站 Ershui Station",
        "still": "簡化教學地圖與二水站局部同框，縱貫線為灰色、集集線為紅色，二水交會點清楚；北向箭頭固定在右上角。",
        "route": "灰色南北線保持靜止，紅色集集線從二水向東側山谷延伸，依序通過源泉、濁水、龍泉、集集、水里到車埕。",
        "subtitle": "二水是集集線的起點",
        "narration": "二水是集集線的起點，火車從這裡轉進山谷。",
        "audio": "輕微紙張展開聲、短促指示音；背景音樂延續。",
        "camera": "俯視固定鏡頭，先聚焦交會點，再沿紅線平滑追蹤。",
        "action": "紅黃列車磁鐵圖示從灰線轉入紅線，無瞬間跳點。",
        "transition": "地圖紅線逐漸變成田野間真實軌道。",
    },
    {
        "id": "04",
        "time": "21–29秒",
        "title": "列車離開二水進入田野",
        "station": "二水站至源泉方向",
        "still": "紅黃列車離開二水，行駛在彰化南端與濁水溪谷附近的田野，稻田、灌溉水圳與遠山層次清楚，沒有海岸景觀。",
        "route": "左下角小路線動畫由二水亮到源泉，列車點同步移動。",
        "subtitle": "支線連接車站、田野與聚落",
        "narration": "列車離開二水，把車站、田野和聚落連起來。",
        "audio": "規律但柔和的車輪聲、田野蟲鳴與風聲；不鳴長笛。",
        "camera": "低角度側向跟拍列車，速度平穩，保留田野全景。",
        "action": "列車向右行駛，田野與水圳自然向左滑過。",
        "transition": "以一叢稻葉遮鏡，接到河谷遠景。",
    },
    {
        "id": "05",
        "time": "29–37秒",
        "title": "沿濁水溪谷前進",
        "station": "濁水／龍泉沿線",
        "still": "高處斜俯視列車沿濁水溪谷北岸前進，河道、河岸平原、低山與軌道相對位置清楚；景觀為教學示意但不誇張。",
        "route": "小路線由源泉依序亮到濁水、龍泉；每站亮起時只出現一個正確圓點。",
        "subtitle": "沿著濁水溪谷前進",
        "narration": "集集線沿著濁水溪谷前進，窗外有河流和山。",
        "audio": "河水遠聲、列車行進聲與輕柔音樂；無暴雨、無急流警報。",
        "camera": "空拍式緩慢平移，但高度穩定，不高速俯衝。",
        "action": "列車沿河谷彎線行駛，小路線點同步移動。",
        "transition": "鏡頭順著軌道彎道帶入綠色隧道。",
    },
    {
        "id": "06",
        "time": "37–45秒",
        "title": "綠色隧道與地方運輸",
        "station": "龍泉至集集沿線",
        "still": "列車通過由老樹形成的綠色隧道，光影斑駁；畫面角落以淡色歷史小圖示呈現早期工程材料與農產品運輸，不出現年代數字。",
        "route": "小路線由龍泉穩定亮向集集，列車圖示與實景方向一致。",
        "subtitle": "以前運材料和農產品，今天也服務旅遊",
        "narration": "這條鐵路以前運送材料和農產品，今天也陪旅客看風景。",
        "audio": "樹葉沙沙、輪軌聲；歷史小圖示出現時有兩聲輕柔木塊音。",
        "camera": "列車前方視角緩慢前進，樹影流動柔和，避免閃爍。",
        "action": "材料、農產品圖示先淡入再淡出，列車持續前進。",
        "transition": "綠色樹冠向上掀開，露出集集車站。",
    },
    {
        "id": "07",
        "time": "45–53秒",
        "title": "抵達集集站",
        "station": "集集站 Jiji Station",
        "still": "依官方照片保留集集車站外觀與月臺特徵，列車緩慢進站，學生在安全區；站名後製區位於上方，不讓模型亂寫。",
        "route": "小路線的集集站圓點放大並亮起，前段二水至集集保持已完成色，後段未亮。",
        "subtitle": "集集站：車站連接小鎮生活",
        "narration": "這是集集站。車站連接小鎮生活，也迎接旅客。",
        "audio": "列車減速聲、短促進站鈴、低音量人群環境聲。",
        "camera": "月臺安全區側向固定鏡頭，列車由左入畫並平穩停靠。",
        "action": "列車減速停穩；人物不奔跑、不靠近月臺邊緣。",
        "transition": "站名圓點擴大成下一鏡的小鎮觀察框。",
    },
    {
        "id": "08",
        "time": "53–61秒",
        "title": "集集小鎮景觀",
        "station": "集集站 Jiji Station",
        "still": "集集站周邊小鎮的車站、街道、單車與綠意組成四格觀察畫面；自然景觀與人文景觀各有清楚例子，沒有不相關知名地標。",
        "route": "上方細路線保持集集站發光，右向箭頭提示旅程尚未結束。",
        "subtitle": "找一找：自然景觀和人文景觀",
        "narration": "在集集，你看見哪些自然景觀？哪些是人們建造的？",
        "audio": "自行車鈴一聲、街道低音量環境聲、鳥鳴。",
        "camera": "由四格中心緩慢平移，依序停留田野、樹木、車站、街道。",
        "action": "自然景觀格出現綠色框，人文景觀格出現橘色框，各停留一秒。",
        "transition": "右向箭頭帶動路線繼續往水里。",
    },
    {
        "id": "09",
        "time": "61–69秒",
        "title": "經過水里進入山城",
        "station": "水里站 Shuili Station",
        "still": "列車靠近水里站與山城聚落，山谷、河流、聚落與月臺清楚；水里作為途中重要車站，但不搶走三站出口說主焦點。",
        "route": "小路線由集集依序亮到水里，水里圓點短暫放大，再沿線指向車埕。",
        "subtitle": "水里是沿線的重要車站",
        "narration": "列車經過水里，山谷中的聚落和交通在這裡相遇。",
        "audio": "山谷風聲、列車輪軌聲、遠處車站環境聲。",
        "camera": "從山谷全景緩慢下降到車站中景，保持地形關係。",
        "action": "小路線上的列車點通過水里後繼續前進。",
        "transition": "鏡頭沿軌道向山谷深處推進到終點。",
    },
    {
        "id": "10",
        "time": "69–77秒",
        "title": "抵達車埕終點",
        "station": "車埕站 Checheng Station",
        "still": "依官方照片保留車埕車站木造與山城終點氣氛，列車慢速進站，終點止衝擋位置合理；無高樓都市天際線。",
        "route": "小路線最後一段由水里亮到車埕，車埕圓點以雙圈顯示終點。",
        "subtitle": "車埕站：集集線的終點",
        "narration": "這是車埕站，也是集集線的終點。",
        "audio": "列車減速、短鈴、森林鳥鳴；抵達時一聲柔和完成音。",
        "camera": "從軌道方向緩慢推近月臺，列車停止後鏡頭保持穩定。",
        "action": "列車停穩，終點雙圈亮起，學生在安全區揮手。",
        "transition": "木造屋簷向上擦拭，接到車埕景觀近景。",
    },
    {
        "id": "11",
        "time": "77–85秒",
        "title": "車埕木業與山林景觀",
        "station": "車埕站 Checheng Station",
        "still": "車埕木造車站、山林與木業文化線索的三層景觀，使用木材堆疊圖示而非危險伐木場面；學生以放大鏡觀察。",
        "route": "上方細路線完整顯示二水到車埕，車埕保持發光，其餘站點依序可見。",
        "subtitle": "車站會保存地方故事",
        "narration": "車埕的木造景觀，讓我們看見地方過去的故事。",
        "audio": "林間風聲、鳥鳴與柔和木琴；無電鋸聲。",
        "camera": "從木造站體緩慢橫移到山林，再回到學生觀察表情。",
        "action": "木材圖示與歷史照片框淡入，停一秒後淡出。",
        "transition": "放大鏡鏡片變成最後的完整路線圖。",
    },
    {
        "id": "12",
        "time": "原始85–93秒；成片保留前5秒，總長90秒",
        "title": "三站回顧與出口說準備",
        "station": "二水站／集集站／車埕站",
        "still": "完整簡化路線圖，二水、集集、車埕三張圓形照片依路線排列；三站下方各留一個大字區，右下留What’s this?口說泡泡。",
        "route": "紅色路線從二水快速依序亮到集集再到車埕，三站照片依序放大一次。",
        "subtitle": "二水—集集—車埕　What’s this?",
        "narration": "二水、集集、車埕。你能說出三個站名嗎？",
        "audio": "三次短促閃卡音、最後一聲完成鈴；背景音樂在第5秒淡出。",
        "camera": "固定正面教學卡構圖，依序小幅推近三站照片。",
        "action": "三站依序亮起，口說泡泡最後出現。",
        "transition": "第5秒停格成教師提問畫面，成片在此結束。",
    },
]


CLOCK_CARDS = [
    ("C01", "基礎整點組", "09:00", "nine o’clock", "短針正指9，長針正指12"),
    ("C02", "基礎整點組", "10:00", "ten o’clock", "短針正指10，長針正指12"),
    ("C03", "基礎整點組", "11:00", "eleven o’clock", "短針正指11，長針正指12"),
    ("C04", "半點挑戰組", "09:00", "nine o’clock", "短針正指9，長針正指12"),
    ("C05", "半點挑戰組", "09:30", "nine thirty", "短針位於9與10正中間，長針正指6"),
    ("C06", "半點挑戰組", "10:00", "ten o’clock", "短針正指10，長針正指12"),
]


ROLE_CARDS = [
    ("R01", "STATION WORKER", "站務人員", "站在服務台旁，手指電子看板，協助旅客找月臺；制服中性、不放真實徽章。"),
    ("R02", "CONDUCTOR", "列車長", "站在已停穩列車門旁，確認上下車秩序；不站在軌道上。"),
    ("R03", "PASSENGER", "旅客", "手持模擬車票與小背包，在安全線後排隊。"),
    ("R04", "SAFETY MONITOR", "安全員", "四年級學生戴可辨識的黃色安全員臂章，舉起『WAIT』圖示卡提醒同學。"),
]


BOARDING_CARDS = [
    ("S01", "CHECK THE TICKET", "看票", "學生站在安全區，雙手拿模擬票確認from與to，票面不得有個資。"),
    ("S02", "CHECK THE BOARD", "看板", "學生抬頭查看簡化電子看板，手指正確月臺欄；看板只放三行模擬資訊。"),
    ("S03", "LINE UP", "排隊", "三位學生在安全線後沿腳印標記排成一列，保留間距。"),
    ("S04", "LET PEOPLE GET OFF FIRST", "先下後上", "列車已停穩，車門開啟；下車箭頭先向外，上車學生在兩側等待。"),
    ("S05", "TAKE YOUR SEAT", "入座", "學生進入車廂後坐在座位上，把走道留空。"),
    ("S06", "KEEP YOUR THINGS SAFE", "保管物品", "學生把背包抱在腿上或放妥，不擋走道，離開前回頭檢查。"),
]


DANGER_CARDS = [
    ("D01", "站在安全線上", "學生腳踩或越過黃色安全線，列車尚未進站；紅色驚嘆號，不畫碰撞。", "退回安全線後等待"),
    ("D02", "在月臺奔跑", "學生在月臺跑動，鞋旁有速度線；與軌道保持距離，不畫跌倒或受傷。", "慢慢走，注意四周"),
    ("D03", "先上後下", "車門開啟，上車學生擋住下車旅客；以交叉箭頭呈現衝突，不推擠。", "先讓人下車，再依序上車"),
    ("D04", "靠門嬉戲", "兩名學生在列車門邊玩鬧，門仍開啟；不夾傷、不驚悚。", "離開門邊，站穩或入座"),
]


STATION_CARDS = [
    (
        "二水",
        "Ershui",
        "集集線起點與鐵路交會線索；田野、濁水溪谷入口。",
        "以前：鐵路協助運送工程材料與沿線農產品。現在：也是居民交通與鐵道旅遊的起點。",
        "二水站官方外觀／站名牌、月臺與軌道、周邊田野或路線起點圖。",
    ),
    (
        "集集",
        "Jiji",
        "小鎮車站、街道、綠意與地方生活。",
        "以前：車站服務沿線居民與物資運輸。現在：仍服務交通，也讓旅客認識集集小鎮。",
        "集集站官方外觀、月臺或站名牌、小鎮街道與自然／人文景觀各一張。",
    ),
    (
        "水里",
        "Shuili",
        "山谷聚落與沿線重要轉乘、服務節點。",
        "以前：鐵路連接山區聚落與物資運輸。現在：水里仍是沿線重要車站與轉乘地點。",
        "水里站官方外觀、月臺與交通轉乘線索、山谷或聚落景觀。",
    ),
    (
        "車埕",
        "Checheng",
        "集集線終點、木造車站、山林與木業文化線索。",
        "以前：鐵路與工程、木業等地方發展有關。現在：終點車站與木造景觀成為認識地方故事的入口。",
        "車埕站官方外觀、終點或軌道線索、木造／山林景觀。",
    ),
]


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_font(run, size=10.5, bold=False, color=None, east_asia="Microsoft JhengHei") -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def format_paragraph(paragraph, size=10.5, bold=False, color=None, after=4, before=0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in paragraph.runs:
        apply_font(run, size=size, bold=bold, color=color)


def relocate_before(element, marker_paragraph) -> None:
    marker_paragraph._p.addprevious(element)


def add_paragraph(doc, marker, text="", style=None, *, size=10.5, bold=False, color=None, after=4, before=0):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    format_paragraph(p, size=size, bold=bold, color=color, after=after, before=before)
    relocate_before(p._p, marker)
    return p


def add_heading(doc, marker, text, level=2, page_break=False):
    style = f"Heading {min(max(level, 1), 3)}"
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    size = {1: 18, 2: 16, 3: 13}.get(level, 12)
    apply_font(run, size=size, bold=True, color=(31, 78, 121))
    p.paragraph_format.space_before = Pt(10 if level == 2 else 6)
    p.paragraph_format.space_after = Pt(5)
    if page_break:
        p.paragraph_format.page_break_before = True
    relocate_before(p._p, marker)
    return p


def add_label_text(doc, marker, label, text, *, size=10.5):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    apply_font(r1, size=size, bold=True, color=(31, 78, 121))
    r2 = p.add_run(text)
    apply_font(r2, size=size)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    relocate_before(p._p, marker)
    return p


def add_bullets(doc, marker, items: Iterable[str], level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        add_paragraph(doc, marker, item, style=style, size=10.5, after=2)


def add_numbered(doc, marker, items: Iterable[str]):
    for item in items:
        add_paragraph(doc, marker, item, style="List Number", size=10.5, after=2)


def add_table(doc, marker, rows, *, widths=(3.2, 12.8), header=None, font_size=9.5):
    row_count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=row_count, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cursor = 0
    if header:
        for j, value in enumerate(header):
            cell = table.rows[0].cells[j]
            cell.text = value
            set_cell_shading(cell, "D9EAF7")
            for p in cell.paragraphs:
                format_paragraph(p, size=font_size, bold=True, color=(31, 78, 121), after=0)
        set_repeat_table_header(table.rows[0])
        cursor = 1
    for i, (key, value) in enumerate(rows, start=cursor):
        left, right = table.rows[i].cells
        left.text = str(key)
        right.text = str(value)
        set_cell_shading(left, "FFF2CC")
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for cell in (left, right):
            set_cell_margins(cell)
        for p in left.paragraphs:
            format_paragraph(p, size=font_size, bold=True, after=0)
        for p in right.paragraphs:
            format_paragraph(p, size=font_size, after=0)
    for row in table.rows:
        row.cells[0].width = Cm(widths[0])
        row.cells[1].width = Cm(widths[1])
    relocate_before(table._tbl, marker)
    add_paragraph(doc, marker, "", size=4, after=0)
    return table


def add_prompt(doc, marker, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.right_indent = Cm(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    r1 = p.add_run(f"{title}：")
    apply_font(r1, size=10, bold=True, color=(192, 80, 77))
    r2 = p.add_run(text)
    apply_font(r2, size=10)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F7F7F7")
    p_pr.append(shd)
    relocate_before(p._p, marker)
    return p


def remove_old_appendix_section(doc):
    start_p = next((p for p in doc.paragraphs if p.text.strip() == START_MARKER), None)
    end_p = next((p for p in doc.paragraphs if p.text.strip() == END_MARKER), None)
    if start_p is None or end_p is None:
        raise RuntimeError("Cannot locate appendix replacement markers.")
    body = doc._body._element
    children = list(body)
    start_idx = children.index(start_p._p)
    end_idx = children.index(end_p._p)
    if start_idx >= end_idx:
        raise RuntimeError("Appendix markers are out of order.")
    for child in children[start_idx:end_idx]:
        body.remove(child)
    return end_p


def replace_text_everywhere(doc, replacements):
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                for run in p.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                if old in p.text:
                    original_style = p.style
                    p.text = p.text.replace(old, new)
                    p.style = original_style
                    format_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacements.items():
                    if old in cell.text:
                        text = cell.text.replace(old, new)
                        cell.text = text
                        for p in cell.paragraphs:
                            format_paragraph(p, size=9.5, after=0)


def add_source_register(doc, marker):
    add_heading(doc, marker, "共通資料查證與使用原則", level=2)
    add_bullets(
        doc,
        marker,
        [
            f"查詢日期：{QUERY_DATE}。重要路線、站名、票務與安全資訊採「一個官方主要來源＋一個可信交叉來源」。",
            "圖片生成不是事實來源。站體、站序、路線方向、票務流程與安全行為須先以官方資料定稿，再進行插畫化。",
            "所有地圖均標示「教學示意圖，非依實際比例繪製」；所有時刻表與車票均標示「模擬教材，不可作為真實乘車資訊」。",
            "AI生成圖中的中文、英文、時間、站名、箭頭、條碼與圖例一律以Word／Canva／Figma後製；不得直接採用生成模型的錯字。",
            "官方照片只做教師端查證或依授權條款使用。若無法確認再利用權利，教材改用連結、課堂投影或自行拍攝照片，不把網路照片重新發布到公開網站。",
            "教師授課前再次開啟臺鐵公告與時刻查詢頁；若路線營運或介面已變更，以授課當日資料為準。",
        ],
    )
    rows = [(name, f"{purpose}\n網址：{url}") for name, purpose, url in SOURCES]
    add_table(doc, marker, rows, widths=(4.0, 12.0), header=("來源", "本案用途與網址"), font_size=8.7)


def add_lesson1(doc, marker):
    add_heading(doc, marker, "第1節自製教材｜二水站暖身圖、90秒路線影片與三站閃卡", level=2, page_break=True)
    add_heading(doc, marker, "一、教學流程重新分析", level=3)
    add_table(
        doc,
        marker,
        [
            ("Warm-up 5分鐘", "學生不是猜景點名稱，而是從「站名牌、鐵軌、月臺」找可見證據。先不放提示圈，第二輪才逐步揭露。教師依序問What’s this?／What’s that?，學生可先指、再用It’s a station／train作答。"),
            ("Presentation 10分鐘", "90秒影片的功能是建立路線方向與地方功能：二水起點→田野與濁水溪谷→集集小鎮→水里山城→車埕終點。每到二水、集集、車埕停格，教師用火車磁鐵沿簡圖移動，避免學生把三張景點照誤認為彼此相鄰。"),
            ("Production 20分鐘", "學生以人體路線、站名卡與景觀卡重建順序；每放一張卡要說明「我把它放這裡，因為……」，同伴用官方路線小卡核對。"),
            ("Wrap-up 5分鐘", "三站閃卡採「看1秒→同伴說→全班說→路線定位」四步；教師混用What’s this?／What’s that?，最後請學生用中文說一項支線對地方的功能。"),
        ],
    )

    add_heading(doc, marker, "二、二水站暖身圖", level=3)
    add_table(
        doc,
        marker,
        [
            ("教學用途", "第一輪投影完整圖30秒靜默觀察；第二輪遮住站名，只讓學生找站名牌、鐵軌、月臺；第三輪依學生回答顯示三個半透明提示圈。"),
            ("教師操作", "先問What do you see?允許中文，接著以Point to the sign／track／platform引導指認，再問What’s this?。不先說答案，不要求學生辨認建築年代。"),
            ("學生任務", "個人找3個線索→兩人互指→用「我看見＿＿，所以我認為這是火車站」說證據。"),
            ("成品規格", "16:9，1920×1080；製作A無標記版、B三提示圈版、C答案標籤版；投影時文字不小於36pt。"),
            ("畫面固定文字", "後製：二水站 Ershui Station；線索1 站名牌 sign；線索2 鐵軌 track；線索3 月臺 platform。"),
            ("建議檔名", "L1_Warmup_ErshuiStation_A_clean.png／B_clues.png／C_answer.png"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "共通製作提示詞",
        "使用授課前重新查證的二水車站官方現況照片作為構圖參考，不改變屋頂、入口、站名牌、月臺與軌道的相對位置。"
        "製作四年級課堂投影用16:9照片式教學圖：左側清楚看到站體與站名牌，右側延伸到月臺與兩條平行鐵軌，"
        "前景保留可放問題的淺色區；日間自然光、真實比例、細節清楚、無遊客臉部特寫。"
        "先輸出完全無文字無標記版本，再由後製加上精確站名與提示圈。禁止虛構第二棟站房、把月臺畫成捷運月臺、"
        "增加高鐵列車、改動軌道方向、簡體字、浮水印、人物站在軌道或跨越安全線。",
    )
    add_prompt(
        doc,
        marker,
        "圖片後製提示詞",
        "在同一張底圖建立三個可獨立顯示的圖層：黃色圓角框圈住站名牌、藍色圓角框圈住鐵軌、綠色圓角框圈住月臺；"
        "答案版在框旁加「站名牌 sign」「鐵軌 track」「月臺 platform」。標籤不遮住證據，箭頭尖端精確落在物件上。"
        "底部加一條深藍問題帶：What’s this?　It’s a station.；右下加小字「照片／圖像僅供課堂觀察，請依來源授權使用」。",
    )
    add_label_text(doc, marker, "驗收清單｜", "①三個線索不被裁切；②不看標籤也能辨認車站；③站體與軌道方向通過官方照片比對；④投影2公尺外可讀；⑤無錯字、簡體字與危險站位。")

    add_heading(doc, marker, "三、90秒路線影片製作方式", level=3)
    add_bullets(
        doc,
        marker,
        [
            "先建立三張一致性Ingredient／參考板：列車三視圖、四站官方照片參考板、路線圖與色彩風格板。生成後鎖定同一專案使用。",
            "依下列12張16:9分鏡首幀逐張生成。每張先做靜態事實檢核，再用Frames to Video轉為約8秒動態片段。",
            "每張Flow提示詞皆完整寫入站名、路線動畫、精確字幕、精確旁白、環境音／音效、鏡頭、動作、轉場與禁止項目；若Flow未正確產生文字或語音，以同欄「後製備援」覆蓋。",
            "剪輯順序01→12；第01段保留5秒，第02–11段各保留8秒，第12段保留5秒，總長90秒。轉場不得額外增加時長。",
            "輸出1920×1080、H.264 MP4、25或30fps；旁白峰值清楚高於環境聲，字幕至少42px、白字深色描邊，置於下方安全區。",
            "試播時以暫停點檢核：二水、集集、車埕三站能在2秒內辨認；路線點不能跳站；所有人物保持月臺安全。",
        ],
    )
    add_prompt(doc, marker, "12張分鏡統一靜態提示詞", FLOW_COMMON_STILL)
    add_prompt(doc, marker, "Flow一致性Ingredient規格", FLOW_CONTINUITY)

    for shot in FLOW_SHOTS:
        add_heading(doc, marker, f"分鏡{shot['id']}｜{shot['title']}", level=3)
        add_table(
            doc,
            marker,
            [
                ("時間", shot["time"]),
                ("畫面站名", shot["station"]),
                ("精確字幕", shot["subtitle"]),
                ("精確旁白", shot["narration"]),
                ("環境音／音效", shot["audio"]),
                ("建議檔名", f"L1_SB{shot['id']}_still.png；L1_SB{shot['id']}_flow.mp4"),
            ],
            font_size=9.3,
        )
        add_prompt(doc, marker, "個別分鏡靜態提示詞", f"{FLOW_COMMON_STILL} 個別畫面：{shot['still']} {FLOW_CONTINUITY}")
        independent_flow = (
            f"8秒、16:9、以分鏡{shot['id']}靜態圖作為首幀，使用固定Ingredients保持列車、站體、路線圖與色盤一致。"
            f"【畫面站名】{shot['station']}。【路線動畫】{shot['route']}。【字幕】只顯示繁體中文與英文："
            f"「{shot['subtitle']}」，置於下方安全區，白字深藍描邊，不得改字。【旁白】臺灣華語、溫暖清楚、四年級語速："
            f"「{shot['narration']}」。【環境音與音效】{shot['audio']}。【攝影機】{shot['camera']}。【畫面動作】{shot['action']}。"
            f"【轉場】{shot['transition']}。不新增其他站名、不顛倒路線、不讓列車瞬移、不讓人物進入軌道或跨越安全線；"
            "不要浮水印、不要簡體字、不要自動改寫字幕、不要可辨識的真實車站廣播。若模型不能準確產生字幕或旁白，保留乾淨安全區與自然環境聲，供後製加入精確文字與聲音。"
        )
        add_prompt(doc, marker, "個別Google Flow動態提示詞", independent_flow)
        add_prompt(
            doc,
            marker,
            "後製備援（逐字照用）",
            f"字幕：{shot['subtitle']}｜旁白：{shot['narration']}｜聲音：{shot['audio']}｜轉場：{shot['transition']}｜"
            "字幕不得由剪輯者改寫；旁白錄音與字幕逐字一致。",
        )
        add_label_text(
            doc,
            marker,
            "單鏡驗收｜",
            "首幀構圖一致；站名與路線位置正確；8秒內旁白說完且不急促；字幕無錯字；聲音不蓋過旁白；轉場不造成跳站；無危險行為。",
        )

    add_heading(doc, marker, "四、三站閃卡｜出口說", level=3)
    add_table(
        doc,
        marker,
        [
            ("教學操作", "每張先閃1秒不讀字→學生兩人回答→顯示站名→全班說It’s Ershui／Jiji／Checheng Station.→把卡貼回路線正確位置。"),
            ("尺寸", "A5直式或1080×1350；站景占70%，站名後製區占20%，底部口說條占10%；印刷版四角圓角、3mm出血。"),
            ("固定文字", "正面：中文站名＋官方英文拼法；背面：What’s this?／What’s that?　It’s _____ Station.；教師版另有路線位置小圖。"),
            ("共通檔名", "L1_Flash_Ershui／Jiji／Checheng_front.png與_back_teacher.png"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "三站閃卡共通提示詞",
        "製作三張同系列A5直式四年級英語口說閃卡，溫暖3D兒童繪本風但忠實參考各站官方照片。"
        "正面只放一個最具辨識度的站體或月臺景觀，構圖、光線、邊框、列車比例一致；上方留中文站名與官方英文站名後製區，"
        "底部留口說句型區。不要把三站地標混在同一張卡，不要自行生成站名文字。人物僅作比例參考並站在安全區。",
    )
    flash_individual = [
        ("二水站 Ershui Station", "呈現二水站正面或月臺與分歧起點線索，背景可有田野，但不可加車埕木造景觀。", "紅色起點圓點"),
        ("集集站 Jiji Station", "呈現集集車站具辨識度的站體與小鎮綠意，不加入水里或車埕終點符號。", "橘色中段圓點"),
        ("車埕站 Checheng Station", "呈現車埕木造站體、山林與終點氣氛，不畫都市高樓或海岸。", "紫色終點雙圈"),
    ]
    for i, (name, visual, route_symbol) in enumerate(flash_individual, start=1):
        add_prompt(
            doc,
            marker,
            f"閃卡{i}個別提示詞｜{name}",
            f"{visual} 教師版背面右上放正確路線位置小圖：{route_symbol}；正面後製精確文字「{name}」。"
            "底部句型後製為「What’s this?」「It’s _____ Station.」。圖片不得包含其他站名。",
        )
    add_label_text(doc, marker, "驗收清單｜", "三卡遠看即可區分；站名與站景一致；英文拼法正確；貼回路線時位置唯一；學生能遮字口說；無錯誤地標、站序與生成文字。")


def add_lesson2(doc, marker):
    add_heading(doc, marker, "第2節自製教材｜時鐘圖卡、簡化時刻表與時刻偵探學習單", level=2, page_break=True)
    add_heading(doc, marker, "一、教學流程重新分析", level=3)
    add_table(
        doc,
        marker,
        [
            ("Warm-up 5分鐘", "先用整點組09:00／10:00／11:00確認短針、長針與at nine o’clock；再用挑戰組09:00／09:30／10:00讓學生比較半點。教師只要求辨時與口說，不先教時刻表。"),
            ("Presentation 10分鐘", "教師用三色透明片示範「先找站名欄→沿車次列向右找時間→比條件→圈車次」，並大聲思考：我要九點後第一班從二水出發，先看哪一欄？"),
            ("Production 20分鐘", "兩人一組扮時刻偵探：讀三個唯一答案任務，必須圈出證據時間並說The train leaves at…；同伴以Answer, please.邀請作答。"),
            ("Wrap-up 5分鐘", "換一題新條件做出口票，學生在30秒內圈站名、框時間、寫車次；教師收回檢查是否沿正確列／欄讀表。"),
        ],
    )
    add_heading(doc, marker, "二、六張時鐘圖卡", level=3)
    add_table(
        doc,
        marker,
        [
            ("成品規格", "兩組各3張，共6張；1:1、1200×1200，印刷10×10公分；鐘面直徑占80%，下方保留數位時間與英文後製區。"),
            ("操作方式", "基礎組正面只看指針、背面看數位時間與英文；挑戰組先比較09:00→09:30→10:00，再請學生用手臂模仿指針。"),
            ("固定色碼", "基礎整點組深藍框；半點挑戰組橘框。相同09:00、10:00仍保留兩張，讓兩組活動能各自完整使用。"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "六卡共通提示詞",
        "製作方形可剪裁教學時鐘卡，白底、粗深藍刻度、12個數字完整且方向正確，一個大鐘面置中，短針明顯較短、長針明顯較長，"
        "無秒針。外框以可愛但不干擾讀時的火車站鐘造型；右下只有小型列車圖示。數字時間與英文由後製加入，"
        "禁止羅馬數字、額外指針、歪斜鐘面、卡通臉遮住刻度、簡體字與錯誤時間。",
    )
    for code, group, time, english, hands in CLOCK_CARDS:
        add_prompt(
            doc,
            marker,
            f"{code}個別提示詞｜{group} {time}",
            f"沿用六卡共通風格；{hands}，指針尖端精確指向刻度。外框為{'深藍色' if '基礎' in group else '橘色'}。"
            f"後製底部精確加入「{time}」「{english}」；建議檔名L2_{code}_{time.replace(':','')}.png。",
        )
    add_label_text(doc, marker, "驗收清單｜", "逐卡以紙製指針覆核角度；09:30短針不可仍正指9；12個數字無缺漏重複；英語與數位時間一致；縮印10公分仍可辨讀。")

    add_heading(doc, marker, "三、三站三車次簡化時刻表", level=3)
    timetable_text = (
        "車次A：二水08:40、集集09:18、車埕09:48；"
        "車次B：二水09:20、集集09:58、車埕10:28；"
        "車次C：二水10:10、集集10:48、車埕11:18。"
    )
    add_table(
        doc,
        marker,
        [
            ("教學定位", "只練讀表策略，不代表真實班次。每次投影與印刷都要顯示「模擬教材，不可作為真實乘車資訊」。"),
            ("表格方向", "列＝車次A/B/C；欄＝二水Ershui、集集Jiji、車埕Checheng。左上角以小圖示提醒「先找站名欄，再沿車次列讀時間」。"),
            ("指定資料", timetable_text),
            ("透明片色碼", "二水欄黃色、集集欄綠色、車埕欄藍色；教師示範時只覆蓋目標欄，避免全表同時上色。"),
            ("成品規格", "投影版16:9；學習單版A4橫向表格寬16公分；時間至少18pt、站名至少16pt。"),
            ("建議檔名", "L2_SimulatedTimetable_teacher_16x9.png；L2_SimulatedTimetable_student_A4.png"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "時刻表製作提示詞",
        "設計四年級使用的三站三車次簡化時刻表，白底、深藍粗格線、列高充足。上方標題「集集線模擬時刻表」，"
        "右上紅框警語「模擬教材，不可作為真實乘車資訊」。欄順序固定為二水Ershui／集集Jiji／車埕Checheng，"
        "列順序固定車次A／B／C；所有時間依指定資料後製輸入，不讓AI自行生成。"
        "在表格左側放四步小圖示：①找站名欄②沿列移動③比較時間④圈車次。保留教師使用三色透明片的清楚空間。",
    )
    add_label_text(doc, marker, "驗收清單｜", "站名與時間完全照指定資料；表格無跨欄錯位；三題任務都有唯一答案；警語醒目；投影與黑白影印都能分辨。")

    add_heading(doc, marker, "四、《旅行包2｜時刻偵探》學習單", level=3)
    add_table(
        doc,
        marker,
        [
            ("學生版上區", "姓名、班級、座號；六張小時鐘配對區（整點3題＋半點3題），數位時間選項可剪貼或連線。"),
            ("學生版中區", "完整模擬時刻表；三題任務旁各有「圈站名→框證據時間→寫車次」三格。"),
            ("唯一答案任務", "①九點後第一班從二水出發是哪一班？②十點前抵達集集的班次中，最晚從二水出發的是哪一班？③哪一班最早抵達車埕？"),
            ("學生版下區", "口說紀錄：The train leaves at _____.／It arrives at _____. 同伴簽名；自評「我有沿同一列找時間」。"),
            ("教師答案版", "同版面，以半透明黃框標任務1車次B與09:20；綠框標任務2車次B、09:20與09:58；藍框標任務3車次A與09:48；頁尾列推理句。"),
            ("規格與檔名", "A4直式雙面或A3單面；學生版L2_WS_TimetableDetective_student.docx／pdf；教師版L2_WS_TimetableDetective_answer.pdf。"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "學習單共通提示詞",
        "A4直式、四年級偵探任務風，藍黃綠配色但不滿版著色，標題大字「旅行包2｜時刻偵探」。"
        "版面由上到下：姓名欄與學習目標、六時鐘配對、三站三車次模擬表、三題唯一答案任務、兩句口說紀錄。"
        "每題都有可圈、可框、可寫的實體空間；繁體中文清楚，英文使用Comic Relief或相近圓體。"
        "時刻、站名與答案全部由後製輸入；頁尾固定紅框警語「模擬教材，不可作為真實乘車資訊」。",
    )
    add_prompt(
        doc,
        marker,
        "學生版提示詞",
        "刪除所有答案、色框與勾選；三題後各保留「車次：____」「證據時間：____」「我的理由：因為____」；"
        "口說區保留同伴簽名線。不要提供會直接暴露答案的箭頭。",
    )
    add_prompt(
        doc,
        marker,
        "教師答案版提示詞",
        "複製學生版，不移動任何題目；只新增半透明答案框、任務答案與一句推理："
        "1 B，因為09:20是九點後最早；2 B，因為09:58前抵達且09:20最晚出發；3 A，因為09:48最早抵達。"
        "答案層可單獨隱藏，方便同一底稿輸出學生版。",
    )
    add_label_text(doc, marker, "驗收清單｜", "學生版零答案；答案版題目位置完全相同；三題唯一答案；至少留8毫米手寫高度；40分鐘流程內可完成；口說句與表格證據一致。")


def add_lesson3(doc, marker):
    add_heading(doc, marker, "第3節自製教材｜票券比較、放大票面與車票解碼學習單", level=2, page_break=True)
    add_heading(doc, marker, "一、教學流程重新分析", level=3)
    add_table(
        doc,
        marker,
        [
            ("Warm-up 5分鐘", "並列公車票、展覽票、模擬火車票，學生先找共同資訊，再說哪一張可能用來搭車。重點是「用可見欄位提出證據」，不是背票種。"),
            ("Presentation 10分鐘", "放大票面依from→to→date→time→car→seat逐欄揭露。每揭一欄，教師說Point to…，學生在個人小票上指同欄，再用What’s this?回答。"),
            ("Production 20分鐘", "兩人扮售票員／旅客解碼正確票與錯票；學生必須指出哪一欄不符旅程條件，並用中文說修正理由。"),
            ("Wrap-up 5分鐘", "完成三項隱私判斷：條碼、訂票代碼、個人資料不拍照、不分享；模擬票一律標SAMPLE。"),
        ],
    )
    add_heading(doc, marker, "二、三張票券比較卡", level=3)
    add_table(
        doc,
        marker,
        [
            ("成品", "公車票、展覽票、模擬火車票各1張；A5橫式，三卡同尺寸、同視角、同背景，方便並列比較。"),
            ("共同欄位", "日期、時間、起點／地點、票種名稱；只有模擬火車票完整呈現from、to、date、time、car、seat。"),
            ("安全設計", "全部為虛構票號；條碼為不可掃描裝飾並加SAMPLE；不放姓名、身分證、電話、真實訂票代碼。"),
            ("操作", "第一輪遮住票種名稱，只看欄位猜用途；第二輪揭露名稱；第三輪圈共同資訊。"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "共通提示詞",
        "製作三張同系列A5橫式兒童票券比較卡，平視桌面、白色票面、粗體欄位、可愛但不幼稚的交通圖示；"
        "每張左上留票種名稱、中央留主要欄位、右側留不可掃描的SAMPLE假條碼。所有文字與數字由後製精確加入，"
        "禁止真實公司商標、真實QR code、姓名、電話、身分證、訂票代碼與可辨識個資。",
    )
    ticket_prompts = [
        ("公車票 Bus Ticket", "FROM Ershui School／TO Ershui Park／DATE 2026-10-08／TIME 08:10；加入小公車圖示，不放car與seat。"),
        ("展覽票 Exhibition Ticket", "PLACE Railway Classroom／DATE 2026-10-08／TIME 13:30／ZONE A；加入小型展覽框圖示，不放from與to。"),
        ("模擬火車票 Sample Train Ticket", "FROM Ershui／TO Checheng／DATE 2026-10-08／TIME 09:20／CAR 2／SEAT 08A；大字SAMPLE，不使用真實車次與票價。"),
    ]
    for i, (name, details) in enumerate(ticket_prompts, start=1):
        add_prompt(doc, marker, f"票券{i}個別提示詞｜{name}", f"{details} 版面與其他兩卡完全一致，僅欄位種類和交通圖示不同。")
    add_label_text(doc, marker, "驗收清單｜", "遮住名稱仍可從欄位推理；火車票六欄齊全；三卡無個資與可掃描碼；所有日期時間符合學習單；縮印仍可讀。")

    add_heading(doc, marker, "三、放大模擬火車票", level=3)
    add_table(
        doc,
        marker,
        [
            ("固定資料", "FROM Ershui；TO Checheng；DATE 2026-10-08；TIME 09:20；CAR 2；SEAT 08A；票面醒目SAMPLE。"),
            ("六色解碼", "from黃、to綠、date藍、time橘、car紫、seat粉；教師用六張覆蓋片逐欄揭露，顏色同學習單。"),
            ("規格", "投影版16:9，票面占畫面85%；列印版A3橫式；每欄至少28pt。"),
            ("隱私提醒", "右下固定：No photos. No sharing.／不拍照、不分享；假條碼上覆SAMPLE斜紋。"),
            ("檔名", "L3_EnlargedSampleTicket_teacher.png；L3_EnlargedSampleTicket_blank.png"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "放大票面提示詞",
        "設計A3橫式放大模擬火車票，六欄按閱讀順序由左到右、由上到下排列：from→to→date→time→car→seat；"
        "每欄有一致圖示、英文小標與大字資料，六種淡色底框不得互換。右側放不可掃描的假條碼與巨大半透明SAMPLE，"
        "下方放Point to from／to／date／time／car／seat六個教師指令按鈕。不要仿製真實票面防偽、票價、車次或訂票代碼。",
    )
    add_label_text(doc, marker, "驗收清單｜", "六欄順序與色碼固定；投影最後一排可讀；假條碼不可掃描；SAMPLE醒目；沒有真實個資；Point to指令可逐欄使用。")

    add_heading(doc, marker, "四、《旅行包3｜車票解碼》學習單", level=3)
    add_table(
        doc,
        marker,
        [
            ("學生版上區", "姓名欄、旅程條件卡：二水→車埕、2026-10-08、09:20、CAR 2、SEAT 08A。"),
            ("學生版中區", "一張正確模擬票；六個解碼格依序抄寫from、to、date、time、car、seat並連到票面。"),
            ("錯票診斷", "錯票A：FROM Checheng／TO Ershui（起訖顛倒）；錯票B：DATE 2026-10-09（日期不符）。學生圈錯欄並寫正確資料。"),
            ("隱私判斷", "勾選可分享／不可分享：我的學習單答案、真實條碼、訂票代碼、姓名與電話；教師說明課堂模擬票不等於真票。"),
            ("教師答案版", "在同版面標示六欄答案、錯票A from/to、錯票B date，並於隱私題標出真實條碼／訂票代碼／姓名與電話不可分享。"),
            ("規格與檔名", "A4直式；L3_WS_TicketDecoder_student與L3_WS_TicketDecoder_answer；學生版與答案版使用同一底稿。"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "學習單共通提示詞",
        "A4直式、可愛驗票員偵探風，標題大字「旅行包3｜車票解碼」。上方旅程條件卡，中段SAMPLE模擬票與六色解碼格，"
        "下段兩張錯票診斷，底部隱私判斷。每個答案格至少8毫米高，箭頭不交叉，票面不使用真實商標或可掃描條碼；"
        "固定警語「模擬教材，不可作為真實乘車資訊」。",
    )
    add_prompt(doc, marker, "學生版提示詞", "刪除所有答案與示範圈；保留六色空格、錯欄圈選空間、修正線與隱私勾選框。條件卡保留，因它是作答依據。")
    add_prompt(doc, marker, "教師答案版提示詞", "複製學生版並鎖定版面；以半透明色框連出六欄；錯票A圈from/to並改為Ershui／Checheng，錯票B圈date並改為2026-10-08；隱私題加理由短句。")
    add_label_text(doc, marker, "驗收清單｜", "學生版無答案；錯票各只有一類錯誤；六欄與放大票面相同；隱私題不讓學生誤以為真票可公開；可在15分鐘內完成。")


def add_lesson4(doc, marker):
    add_heading(doc, marker, "第4節數位教材｜iPad模擬線上購票網站設計規劃書（完善版）", level=2, page_break=True)
    add_heading(doc, marker, "一、教學與產品定位", level=3)
    add_table(
        doc,
        marker,
        [
            ("課堂目標", "學生依序完成起點二水→目的地→日期→查詢→選車次→確認摘要，能發現起訖顛倒並回上一步修正；到確認頁即停止。"),
            ("非目標", "不連接臺鐵真實訂票、不付款、不建立會員、不輸入姓名、身分證、電話、電子郵件、信用卡或真實訂票代碼。"),
            ("使用情境", "教師投影示範1次；兩人共用1台iPad，一人操作、一人使用六步檢核卡；完成後交換角色。"),
            ("完成回饋", "確認頁顯示SAMPLE模擬票、彩帶與2秒完成鈴；可靜音；明確文字「練習完成，沒有真的訂票」。"),
            ("部署邊界", "純前端HTML／CSS／JavaScript，可部署Firebase Hosting；不使用Firestore、Authentication、Cloud Functions與分析追蹤。"),
        ],
    )
    add_heading(doc, marker, "二、六步學生流程與畫面", level=3)
    steps = [
        ("1 選起點", "預設空白；學生點選二水Ershui。若點其他站，仍可繼續，但教師任務卡指定二水。", "大按鈕、站名中英並列、路線圖上二水亮起。"),
        ("2 選目的地", "可選集集Jiji／水里Shuili／車埕Checheng；練習主線選車埕。", "選取後路線由二水亮到目的地。"),
        ("3 選日期", "只顯示三個虛構教學日期，不呼叫真實日曆；預設2026-10-08。", "醒目顯示「模擬日期」。"),
        ("4 查詢", "先驗證起訖不可相同；若起訖顛倒，顯示友善提示並提供「交換」或「回上一步」。", "不使用紅色羞辱訊息；以橘色提示和箭頭。"),
        ("5 選車次", "顯示第2節同一組A/B/C模擬車次；每張卡列出出發、抵達、所需時間，資料旁標SAMPLE。", "選取區至少48×48px，整張卡可點。"),
        ("6 確認摘要", "逐列顯示from、to、date、train、depart、arrive；學生勾選六格後按「完成練習」。", "沒有購買、付款或送出訂票按鈕。"),
    ]
    for title, logic, ui in steps:
        add_table(doc, marker, [("操作邏輯", logic), ("畫面要求", ui)], header=(title, "規格"), font_size=9.2)

    add_heading(doc, marker, "三、教師示範與互動設計", level=3)
    add_bullets(
        doc,
        marker,
        [
            "教師模式網址參數?teacher=1：顯示步驟編號、教學旁白提示、錯誤示範按鈕、重設按鈕；學生模式隱藏答案與旁白。",
            "教師先故意選「起點車埕、目的地二水」，問What’s wrong?，再按「交換起訖」示範修正；系統保留橘色箭頭3秒後淡出。",
            "每一步完成才解鎖「下一步」；若缺欄位，焦點移到缺漏位置並朗讀簡短提示，不能只靠顏色。",
            "雙人分工卡：Operator操作iPad；Checker逐項說Check from／to／date／train／time；第2輪交換。",
            "完成畫面停留，學生抄寫摘要到《旅行包4》，再由教師統一重設；網站不保存學生姓名或操作紀錄。",
        ],
    )

    add_heading(doc, marker, "四、資訊架構與資料契約", level=3)
    add_table(
        doc,
        marker,
        [
            ("頁面／元件", "首頁任務卡、六步進度條、站點選擇器、日期卡、模擬班次卡、確認摘要、SAMPLE車票、完成回饋、教師提示抽屜。"),
            ("config.js", "集中管理站名中英、站序、三個模擬日期、A/B/C時刻、提示文字、音效開關、版本號與警語；畫面不可硬編資料。"),
            ("狀態", "sessionStorage只保存本機當次步驟、選項與靜音狀態；關閉分頁後可清除。不得放localStorage長期追蹤。"),
            ("驗證", "起點必填、目的地必填、起訖不同、日期必填、車次必選、確認六欄全勾；每一錯誤都有文字訊息。"),
            ("路線規則", "站序依二水→源泉→濁水→龍泉→集集→水里→車埕；學生選目的地時路線只亮到該站，不跨跳。"),
            ("模擬資料", "沿用第2節A/B/C時刻。首頁與結果頁固定顯示「模擬教材，不可作為真實乘車資訊」。"),
        ],
    )
    add_heading(doc, marker, "五、視覺、iPad與無障礙規格", level=3)
    add_bullets(
        doc,
        marker,
        [
            "iPad直式與橫式皆可用；主要內容寬度不超過960px；按鈕觸控區至少48×48px，間距至少8px，正文至少18px。",
            "顏色：二水黃、集集綠、水里藍、車埕紫；同時用站名、圖示和形狀區分，不以顏色作唯一線索。",
            "每一步有h1標題、可見鍵盤焦點、正確label、aria-live錯誤提示；可用Tab／Enter完成，不要求拖曳。",
            "動畫尊重prefers-reduced-motion；彩帶不閃爍、2秒內結束；音效有靜音鍵且不自動長時間播放。",
            "所有生成插圖不承載關鍵文字；站名、時間、錯誤訊息與按鈕全部以HTML真實文字呈現。",
        ],
    )
    add_heading(doc, marker, "六、離線與故障替代", level=3)
    add_bullets(
        doc,
        marker,
        [
            "網站核心檔案與模擬資料可做簡易PWA快取；無網路時仍能完成六步，但頁面要顯示「離線練習模式」。",
            "若iPad無法開啟，使用同版面的A4六步紙卡、三張車次卡與SAMPLE確認票；學習目標不變。",
            "教師課前在實際教室Wi-Fi用至少兩台iPad測試QR、觸控、音量、旋轉與重新整理；準備短網址與紙本備援。",
        ],
    )
    add_heading(doc, marker, "七、驗收測試案例", level=3)
    add_numbered(
        doc,
        marker,
        [
            "標準流程：二水→車埕→2026-10-08→車次B→確認，顯示正確摘要、SAMPLE票、彩帶與2秒鈴。",
            "起訖顛倒：車埕→二水，能使用交換或回上一步修正，不遺失日期。",
            "起訖相同：二水→二水，禁止進入查詢並以文字說明原因。",
            "缺漏欄位：未選日期或車次，焦點移到缺漏欄位，螢幕閱讀器可讀提示。",
            "重整頁面：當次狀態依sessionStorage合理恢復；使用「重新開始」可完全清空。",
            "無網路：已快取版本可完成；未快取時顯示紙本替代指示，不出現空白頁。",
            "隱私：開發者工具與網路請求中不含姓名、身分證、電話、電子郵件、付款或真實訂票資料。",
            "版面：iPad直／橫向、Safari與Chrome無橫向溢出，所有按鈕可觸控，文字不被鍵盤遮住。",
        ],
    )
    add_label_text(doc, marker, "建議網站檔案｜", "index.html、styles.css、app.js、config.js、manifest.webmanifest、service-worker.js、assets/、README_TEACHER.md、QA_CHECKLIST.md。此階段只完成規劃，不建立或部署網站。")


def add_lesson5(doc, marker):
    add_heading(doc, marker, "第5節自製教材｜站務角色、六步搭車流程與危險動作圖卡", level=2, page_break=True)
    add_heading(doc, marker, "一、教學流程重新分析", level=3)
    add_table(
        doc,
        marker,
        [
            ("Warm-up 5分鐘", "並列安全／不安全月臺圖，學生先以手勢判斷，再說可見證據；不播放事故影像。"),
            ("Presentation 10分鐘", "用4張角色卡建立職責，再把6張流程卡貼成看票→看板→排隊→先下後上→入座→保管物品。教師示範一次正確流程與一次可修正錯誤。"),
            ("Production 20分鐘", "四人輪流扮station worker、conductor、passenger、safety monitor；每輪抽1張危險動作卡，安全員先說Stop／Wait，再把行為改成安全版本。"),
            ("Wrap-up 5分鐘", "學生抽角色卡回答Who are you?並說一項任務；全班用Line up.／Wait, please.／Let people get off first.完成安全口號。"),
        ],
    )
    add_heading(doc, marker, "二、四張站務角色圖卡", level=3)
    add_table(
        doc,
        marker,
        [
            ("規格", "A5直式、同一月臺背景、同一畫風與視角；人物占65%，上方英文職稱、下方中文與一項任務後製區。"),
            ("共通句型", "Who are you?　I’m a _____.／What do you do?　I help／check／wait／remind…"),
            ("安全", "所有人物在安全線後或已停穩列車旁；不仿製真實制服徽章；安全員明確是學生角色而非正式職稱。"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "四卡共通提示詞",
        "製作四張同系列A5直式角色圖卡，溫暖3D日式繪本風、清楚全身姿勢、臺灣火車站月臺背景，人物表情友善且角色任務一眼可辨。"
        "上方與下方保留文字安全區，文字後製。人物不站軌道、不越過安全線、不使用真實臺鐵商標或制服徽章；多元性別與外觀，"
        "不要把列車長畫成司機，不要讓學生安全員看起來像警察。",
    )
    for code, english, chinese, scene in ROLE_CARDS:
        add_prompt(
            doc,
            marker,
            f"{code}個別提示詞｜{english} {chinese}",
            f"{scene} 後製精確文字「{english}」「{chinese}」；背面任務短句依角色填入。建議檔名L5_{code}_{english.replace(' ','_')}.png。",
        )
    add_label_text(doc, marker, "角色卡驗收｜", "遮住文字仍能從動作辨識角色；四卡任務不重複；安全員明確為課堂分工；英語拼法正確；無軌道站人與商標誤用。")

    add_heading(doc, marker, "三、六張搭車流程卡", level=3)
    add_table(
        doc,
        marker,
        [
            ("規格", "A5橫式或16:9小卡，六卡可在黑板排成一列；左上放1–6序號，中央圖像，底部英文大字＋中文。"),
            ("共同視覺", "同一組四年級學生、同一車站與同一列車；流程時間連續，列車只在第4步前停穩並開門。"),
            ("課堂使用", "先打亂排序，再由學生用First／Next／Then／Finally說明；教師以箭頭卡連接。"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "六卡共通提示詞",
        "製作六張連續故事式搭車流程卡，A5橫式、清楚3D兒童動畫風；固定同一組學生、同一月臺、同一紅黃列車、同一光線與服裝。"
        "畫面以動作為中心，序號與文字均後製；安全線、車門與行進方向前後一致。禁止跳步、列車未停即開門、人物站在軌道、"
        "行李擋住走道、簡體字、浮水印。",
    )
    for code, english, chinese, scene in BOARDING_CARDS:
        add_prompt(
            doc,
            marker,
            f"{code}個別提示詞｜{english} {chinese}",
            f"{scene} 後製左上序號{int(code[-2:])}，底部精確文字「{english}」「{chinese}」。建議檔名L5_{code}_{english.replace(' ','_')}.png。",
        )
    add_label_text(doc, marker, "流程卡驗收｜", "單看圖能排出唯一順序；第4步清楚呈現先下後上；人物與物品位置連續；英文與中文一致；列車和車門狀態合理。")

    add_heading(doc, marker, "四、四張危險動作圖卡", level=3)
    add_table(
        doc,
        marker,
        [
            ("規格", "A5橫式、紅色外框；正面只呈現待修行為與What’s wrong?，背面呈現安全改法與Do this.。"),
            ("呈現原則", "不出現事故結果、血、受傷、跌落或驚恐特寫；危險點靠姿勢、箭頭與驚嘆號辨認。"),
            ("教學操作", "抽卡者描述可見行為，安全員說Stop／Wait, please.，小組用角色演出安全改法。"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "四卡共通提示詞",
        "製作四張A5橫式月臺安全判斷卡，同一車站背景、清楚3D兒童漫畫風。正面用紅框與一個驚嘆號指出待修行為，"
        "但不畫事故後果；背面同構圖改成安全行為，以綠框和勾號呈現。保留What’s wrong?與Do this.後製區。"
        "禁止血腥、受傷、掉落軌道、碰撞、恐怖表情、過度誇張、人物臉部污名化。",
    )
    for code, title, unsafe_scene, safe_fix in DANGER_CARDS:
        add_prompt(
            doc,
            marker,
            f"{code}個別提示詞｜{title}",
            f"正面：{unsafe_scene} 背面使用相同人物與鏡位改成「{safe_fix}」。"
            f"後製文字「What’s wrong?」「{title}」「Do this: {safe_fix}」。建議檔名L5_{code}_unsafe_front.png與_safe_back.png。",
        )
    add_label_text(doc, marker, "危險卡驗收｜", "危險行為可辨識但沒有事故傷害；背面改法與正面一一對應；安全線方向一致；學生能用一句話指出問題；不把任何角色畫成壞人。")


def add_lesson6(doc, marker):
    add_heading(doc, marker, "第6節數位教材｜教室數位集集線之旅：觀察地方景觀", level=2, page_break=True)
    add_heading(doc, marker, "一、教學流程重新分析", level=3)
    add_table(
        doc,
        marker,
        [
            ("Warm-up 5分鐘", "用二水、集集、水里、車埕照片各取一個局部，學生先分類natural／human-made，再說圖像證據。"),
            ("Presentation 10分鐘", "播放教師預錄地圖導覽，明確示範「整體看→局部找→資料核」：先看站點位置，再放大景觀，最後讀官方短卡，避免只憑照片猜歷史。"),
            ("Production 20分鐘", "四站輪站，每站4分鐘：攝影師框景觀、資料員讀短卡、英文員說station／train／old／new、記錄員完成以前／現在；輪站換角色。"),
            ("Wrap-up 5分鐘", "每組選一站說「我看見＿＿；資料告訴我＿＿」，再提出一個可查證問題，教師把未能確認的問題列入後續查證。"),
        ],
    )
    add_heading(doc, marker, "二、官方沿線照片包", level=3)
    add_table(
        doc,
        marker,
        [
            ("資料夾結構", "L6_PhotoPack/01_Ershui、02_Jiji、03_Shuili、04_Checheng；每站3張：站體／月臺或軌道／地方景觀。"),
            ("照片來源", "優先臺鐵、交通部觀光署、南投縣政府旅遊網；每張另存source.csv：檔名、標題、發布單位、網址、查詢日期、授權／使用方式。"),
            ("選圖標準", "能支持學生找自然／人文、交通／聚落、以前／現在的證據；避免純美照、濾鏡、無地點說明、人物臉部特寫。"),
            ("使用邊界", "不確定授權時不下載重發，改用教師投影原網頁或連結；不得把AI生成圖標成官方照片。"),
            ("尺寸", "教室投影長邊至少1600px；卡片裁切4:3但保留原圖；右下加來源編號而非網址長字串。"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "照片包整理提示詞",
        "你是國小地方課程教材編輯。只使用我提供、已確認來源與授權的官方照片，不能生成或改造地標。"
        "為二水、集集、水里、車埕各建立3張4:3觀察卡：A站體全景、B月臺／軌道交通線索、C地方自然或人文景觀。"
        "只做亮度、水平與安全裁切，不移除或新增物件；右上保留站名後製區，右下放來源編號。"
        "輸出教師投影版與學生無站名猜圖版。若照片無法證明站名或年代，標為「待查證」，不要推測。",
    )
    for idx, (zh, en, focus, then_now, photos) in enumerate(STATION_CARDS, start=1):
        add_prompt(
            doc,
            marker,
            f"照片包{idx}｜{zh} {en}",
            f"選圖焦點：{focus} 必備三張：{photos}。後製檔名L6_{idx:02d}_{en}_A_station／B_track／C_landscape.jpg；"
            f"學生版只顯示照片編號，教師版加「{zh} {en}」與來源編號。",
        )
    add_label_text(doc, marker, "照片包驗收｜", "每站正好3張；每張有可追溯來源；站名與照片一致；裁切不移除關鍵證據；四站視覺可比較；沒有把生成圖冒充照片。")

    add_heading(doc, marker, "三、教師預錄地圖導覽", level=3)
    add_table(
        doc,
        marker,
        [
            ("建議長度", "2分30秒：0–15秒任務；15–35秒全線與方向；35–60秒二水；60–85秒集集；85–110秒水里；110–135秒車埕；135–150秒三步觀察回顧。"),
            ("畫面", "同一張正確集集線簡圖＋四站官方照片；路線依序點亮，不跳過源泉、濁水、龍泉，但四站用較大圓點。"),
            ("旁白重點", "只說照片與官方短卡支持的內容；明確區分「我看見」與「資料告訴我」。"),
            ("輸出", "1920×1080 MP4、字幕版與無字幕版；字幕繁體中文，站名中英並列；另輸出6張投影片PDF作紙本備援。"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "預錄導覽統一提示詞",
        "製作2分30秒四年級地圖導覽，使用已查證集集線簡圖與已授權官方照片。路線固定依二水→源泉→濁水→龍泉→集集→水里→車埕，"
        "四個學習站二水、集集、水里、車埕使用大圓點，其餘小圓點。每站畫面依序示範：①整體看站點在路線哪裡；"
        "②局部找一個自然與一個人文線索；③資料核對官方短卡。旁白使用「我看見……；資料告訴我……」句框。"
        "所有站名、字幕與箭頭後製；不要把照片動畫化成不存在的事件，不推測年代，不加入未查證地標。",
    )
    add_prompt(
        doc,
        marker,
        "六張導覽投影片提示詞",
        "第1張任務與三步觀察；第2張完整路線與北向箭頭；第3–6張依序二水、集集、水里、車埕，左側官方照片、右側兩欄"
        "「我看見」「資料告訴我」，底部英語標籤station／train／old／new。每張標來源編號與查詢日期。",
    )
    add_label_text(doc, marker, "導覽驗收｜", "路線順序正確；四站各有時間；旁白與字幕一致；照片和資料卡來源可追溯；自然／人文例子能從圖中指出；總長不超過2分30秒±2秒。")

    add_heading(doc, marker, "四、集集線官方圖文短卡", level=3)
    add_table(
        doc,
        marker,
        [
            ("卡面", "每站A5直式；上方站名中英與路線位置，中段官方照片，下方70–90字四年級短文。"),
            ("短文結構", "第1句位置／功能；第2句可見景觀；第3句以前／現在；第4句證據來源提示。"),
            ("閱讀支架", "關鍵詞加粗但不超過4個；附「我看見＿＿」「資料告訴我＿＿」兩個句框。"),
            ("不可做", "不直接複製長篇官方文字、不加入無來源年代、不把觀光宣傳語當歷史事實。"),
        ],
    )
    short_texts = [
        ("二水 Ershui", "二水站是集集線的起點，也是認識地方鐵路的重要入口。從車站出發，列車沿著濁水溪谷北岸前進，窗外可見田野和山景。早期鐵路曾協助運送工程材料與農產品，現在也服務交通與旅遊。"),
        ("集集 Jiji", "集集站位在集集線中段，車站和小鎮生活緊密相連。照片中可以找站體、月臺、街道和綠意。鐵路過去服務居民與物資運輸，現在仍連接地方，也讓旅客從車站認識集集。"),
        ("水里 Shuili", "水里站位在山谷聚落，是集集線沿線的重要車站與轉乘地點。照片中可以找月臺、軌道、聚落和山景。鐵路過去協助連接山區交通，現在仍讓居民與旅客往來各站。"),
        ("車埕 Checheng", "車埕站是集集線的終點。照片中可見木造車站、軌道終點、山林與木業文化線索。鐵路過去和地方工程、產業發展有關；現在，人們也從車站與木造景觀認識地方故事。"),
    ]
    for name, text in short_texts:
        add_prompt(doc, marker, f"學生短卡文字｜{name}", text)
    add_label_text(doc, marker, "短卡驗收｜", "每卡70–90字左右；詞彙適合四年級；以前／現在在同一比較面向；每項歷史敘述有來源；學生能找到至少兩個照片證據。")

    add_heading(doc, marker, "五、四站「以前／現在」資料卡", level=3)
    add_table(
        doc,
        marker,
        [
            ("規格", "每站A4橫式一張，左圈「以前」、右圈「現在」、中間「一直都有／共同功能」；照片與短句分離，可遮答案。"),
            ("來源", "以前照片優先國家文化記憶庫、國家檔案、地方文化館或官方典藏；現在照片用授權官方照片。兩者都要記錄年代或標「未註明」。"),
            ("任務", "學生先只看兩張圖說變化，再翻開資料文字核對；不能從黑白顏色直接判斷年代。"),
            ("英文支架", "old／new／station／train；口說不要求比較級：It is old／new.／This is a station."),
        ],
    )
    add_prompt(
        doc,
        marker,
        "四站資料卡共通提示詞",
        "使用我提供、已查證來源的以前與現在照片製作A4橫式比較卡。左側米褐色「以前」、右側淺藍色「現在」、中央白色共同點，"
        "每側保留照片、來源編號與兩行短句；照片不做AI重繪、不把黑白照自動上色。下方放四個證據圖示：建築、軌道、載運、地方功能。"
        "若照片日期不明，精確標「日期未註明」，不得猜測。",
    )
    for idx, (zh, en, focus, then_now, photos) in enumerate(STATION_CARDS, start=1):
        add_prompt(
            doc,
            marker,
            f"資料卡{idx}｜{zh} {en}",
            f"比較文字：{then_now} 學生版遮住文字，只保留兩張圖、來源編號與「我看到的證據」空格；"
            f"教師版顯示短句與參考答案。建議檔名L6_{idx:02d}_{en}_ThenNow_student／teacher.pdf。",
        )
    add_label_text(doc, marker, "資料卡驗收｜", "以前與現在照片確為同站或同地區；日期有來源或標未註明；比較面向一致；學生版無答案；教師版不把推測寫成事實。")

    add_heading(doc, marker, "六、《旅行包6｜數位車站觀察》學習單", level=3)
    add_table(
        doc,
        marker,
        [
            ("版面", "A4直式雙面；正面二水／集集，背面水里／車埕。每站四欄：照片編號、我看見、資料告訴我、以前／現在。"),
            ("英語", "每站圈選station／train／old／new至少1詞，完成This is a _____.；不把英語拼寫量拉高。"),
            ("角色紀錄", "頁首四角色輪換勾選：攝影師、資料員、英文員、記錄員。"),
            ("教師版", "同版面加入每站2個可見線索、1項資料訊息與可接受英語詞；開放答案以「答案示例」標示。"),
            ("檔名", "L6_WS_DigitalStationObservation_student.pdf；L6_WS_DigitalStationObservation_teacher.pdf"),
        ],
    )
    add_prompt(
        doc,
        marker,
        "學習單提示詞",
        "A4直式雙面、相機觀察框與證據放大鏡風格，四年級易讀。二水、集集、水里、車埕各一個固定色框，"
        "每框依序放「照片編號」「我看見」「資料告訴我」「以前／現在」；每欄至少兩行手寫空間。"
        "頁側放station／train／old／new可圈選詞卡；頁尾放句框This is a _____."
        "學生版刪除示例；教師版複製同底稿加半透明示例答案。不要預先替學生判斷所有景觀。",
    )
    add_label_text(doc, marker, "學習單驗收｜", "四站名稱完整；輪站4分鐘內可填一格；『看見』與『資料』欄明確分開；學生版無答案；教師版保留開放答案彈性；英語負荷適合四年級。")


def update_remaining_worksheet_summaries(doc):
    replacements = {
        "第1張《旅行包1｜路線護照》：上方二水—集集—車埕三站空白路線，中段三站照片貼格與自然／人文分類，下方What’s this/that?互簽。提示詞：A4直式、紅黃小火車護照風、三站箭頭清楚、留大面積貼卡與圈選區。":
            "第1張《旅行包1｜路線護照》：沿用本文件第1節詳規；三站空白路線、照片貼格、自然／人文分類及What’s this/that?互簽。路線須保留中間站小圓點並標示教學示意圖，不能畫成三站彼此相鄰。",
        "第2張《旅行包2｜時刻偵探》：三時鐘、三站三車次表、三題任務、四步作答及at...口說勾核。提示詞：A4直式、藍黃綠偵探風、所有時間與站名大字、模擬教材警語醒目。":
            "第2張《旅行包2｜時刻偵探》：沿用第2節詳規；六時鐘（兩組各3張）、三站三車次模擬表、三題唯一答案任務、證據時間與口說紀錄；學生版與教師答案版使用同一底稿。",
        "第3張《旅行包3｜車票解碼》：模擬票、from→to→date→time→car→seat六格、錯票診斷、隱私勾選。提示詞：A4直式、可愛驗票員風、六色欄位與SAMPLE假條碼。":
            "第3張《旅行包3｜車票解碼》：沿用第3節詳規；旅程條件卡、SAMPLE票、六欄解碼、兩張單一錯誤票與隱私判斷；學生版與教師答案版使用同一底稿。",
        "第4張《旅行包4｜安全購票流程》：六步網站截圖空格、每步箭頭、起訖顛倒修正題、確認摘要核對表。提示詞：A4直式、iPad介面與大按鈕視覺、不得要求任何個資、最後寫「到確認頁先停一下」。":
            "第4張《旅行包4｜安全購票流程》：配合第4節網站規劃；六步截圖／圖示空格、起訖顛倒修正、確認摘要檢核；固定文字「到確認頁即停止」「練習完成，沒有真的訂票」，不得要求個資。",
        "第5張《旅行包5｜搭車安全檢核》：六步排序區、四角色任務表、四張危險動作改正欄、Line up.口說勾核。提示詞：A4橫式、月臺安全漫畫風、綠色正確／紅色待修、人物遠離軌道。":
            "第5張《旅行包5｜搭車安全檢核》：配合第5節14張圖卡；六步排序、四角色任務、四種危險行為與安全改法、Line up.／Wait, please.口說勾核；不呈現事故傷害。",
        "第6張《旅行包6｜數位車站觀察》：照片編號、看見的事、資料告訴我的事、我的感受三欄，附station／train／old／new圖詞。提示詞：A4直式、相機觀察框與證據放大鏡風、四格照片編號、欄位清楚。":
            "第6張《旅行包6｜數位車站觀察》：沿用第6節詳規；二水、集集、水里、車埕各含照片編號、我看見、資料告訴我、以前／現在，附station／train／old／new圖詞；學生版與教師版同底稿。",
    }
    replace_text_everywhere(doc, replacements)


def main():
    doc = Document(DOCX_PATH)
    marker = remove_old_appendix_section(doc)

    flow_replacements = {
        "可擦寫時刻表、三色透明片、任務卡、時鐘圖卡。":
            "可擦寫時刻表、三色透明片、任務卡、時鐘圖卡兩組（整點組09:00／10:00／11:00；半點挑戰組09:00／09:30／10:00）。",
        "《旅行包6｜數位觀察紀錄》：四站各一格，畫／貼景觀、記證據、填一個英文標籤。":
            "《旅行包6｜數位觀察紀錄》：二水、集集、水里、車埕各一格，畫／貼景觀、分開記錄「我看見」與「資料告訴我」，填一個英文標籤。",
        "教師帶看四站，示範觀察三步":
            "教師帶看二水、集集、水里、車埕，示範觀察三步",
        "四站輪站觀察與同伴互教":
            "二水、集集、水里、車埕四站輪站觀察與同伴互教",
        "四站各4分鐘，任務分為「找景觀」「找變化」「英語標籤」。":
            "二水、集集、水里、車埕各4分鐘，任務分為「找景觀」「找變化」「英語標籤」。",
    }
    replace_text_everywhere(doc, flow_replacements)

    add_source_register(doc, marker)
    add_lesson1(doc, marker)
    add_lesson2(doc, marker)
    add_lesson3(doc, marker)
    add_lesson4(doc, marker)
    add_lesson5(doc, marker)
    add_lesson6(doc, marker)
    update_remaining_worksheet_summaries(doc)

    # Improve document-wide typography for newly inserted headings without disturbing existing content.
    styles = doc.styles
    for name, size, color in (
        ("Heading 1", 18, (31, 78, 121)),
        ("Heading 2", 16, (31, 78, 121)),
        ("Heading 3", 13, (47, 84, 150)),
    ):
        if name in styles:
            style = styles[name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor(*color)

    core = doc.core_properties
    core.subject = "第3–10週坐火車趣集集：第1–6節自製教材製作規格、提示詞與網站規劃"
    core.comments = f"依RDQ確認規格更新；資料查詢日期{QUERY_DATE}；第7–8節未變更。"
    doc.save(DOCX_PATH)

    # Read-back checks.
    check = Document(DOCX_PATH)
    all_text = "\n".join(p.text for p in check.paragraphs)
    all_text += "\n" + "\n".join(cell.text for t in check.tables for row in t.rows for cell in row.cells)
    required = [
        "分鏡12｜三站回顧與出口說準備",
        "個別Google Flow動態提示詞",
        "基礎整點組",
        "半點挑戰組",
        "iPad模擬線上購票網站設計規劃書（完善版）",
        "STATION WORKER",
        "LET PEOPLE GET OFF FIRST",
        "二水、集集、水里、車埕",
        "模擬教材，不可作為真實乘車資訊",
        "不拍照、不分享",
    ]
    missing = [item for item in required if item not in all_text]
    print(f"saved={DOCX_PATH}")
    print(f"paragraphs={len(check.paragraphs)} tables={len(check.tables)}")
    print(f"missing_required={missing}")
    print(f"flow_prompt_count={all_text.count('個別Google Flow動態提示詞')}")
    print(f"storyboard_count={sum(1 for p in check.paragraphs if p.text.startswith('分鏡') and '｜' in p.text)}")
    print(f"lesson7_present={'第7節' in all_text} lesson8_present={'第8節' in all_text}")
    if missing:
        raise SystemExit(2)


if __name__ == "__main__":
    main()
