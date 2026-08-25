from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "在地課程4年級上學期教案" / "02_第3-10週_坐火車趣集集.docx"

ZH_LABEL = "中文版 Google Flow 動態提示詞："
EN_LABEL = "English Google Flow motion prompt: "
SUBTITLE_LABEL = "後製字幕："
NARRATION_LABEL = "後製旁白："
NEW_NARRATION_LABEL = "中文旁白逐字稿（已嵌入中英文 Flow 提示詞）："


ZH_OLD_AUDIO = [
    "只加入柔和列車鈴、低音量木琴與輕微行車聲，不要人聲、旁白或廣播。",
    "只加入鳥鳴、微風與低音量車站環境聲，不要可辨識廣播、人聲或旁白。",
    "只加入輕微紙張展開聲與柔和提示音，不要人聲。",
    "只加入柔和輪軌聲、風聲與田野蟲鳴，不要人聲、字幕、浮水印或亂碼。",
    "只加入遠方水聲、柔和行車聲與低音量音樂，不要旁白、字幕、浮水印、亂碼或額外圖示。",
    "只加入樹葉聲、柔和輪軌聲與兩聲木塊提示音，不要人聲、字幕、浮水印或亂碼。",
    "只加入減速輪軌聲、短鈴與低音量人群環境聲，不要可辨識廣播、旁白、字幕、浮水印或亂碼。",
    "只加入一聲自行車鈴、鳥鳴與低音量街道環境聲，不要旁白、字幕、分類文字、浮水印或亂碼。",
    "只加入山谷風聲、柔和輪軌聲與低音量車站環境聲，不要旁白、字幕、浮水印、亂碼或額外站名。",
    "只加入減速聲、短鈴、鳥鳴與柔和完成音，不要旁白、字幕、浮水印、亂碼或其他站名。",
    "只加入林間風聲、鳥鳴與柔和木琴，不要電鋸聲、旁白、字幕、浮水印或亂碼。",
    "加入四次柔和閃卡音與最後一聲完成鈴即可，不要人聲。",
]

EN_OLD_AUDIO = [
    "Use only a soft train bell, quiet xylophone music, and gentle rail ambience. No speech, narration, or announcement.",
    "Use birds, light wind, and quiet station ambience only. No intelligible announcement, speech, or narration.",
    "Use only a soft paper-opening sound and gentle cue tones. No speech.",
    "Use soft rail rhythm, wind, and field insects only. No speech, subtitles, watermark, or gibberish.",
    "Use distant water, soft rail sounds, and quiet music only. No narration, subtitles, watermark, gibberish, or extra icons.",
    "Use leaves, soft rail sounds, and two gentle wooden cue sounds only. No speech, subtitles, watermark, or gibberish.",
    "Use slowing rail sounds, one short bell, and quiet crowd ambience only. No intelligible announcement, narration, subtitles, watermark, or gibberish.",
    "Use one bicycle bell, birds, and quiet street ambience only. No narration, subtitles, category labels, watermark, or gibberish.",
    "Use valley wind, soft rail sounds, and quiet station ambience only. No narration, subtitles, watermark, gibberish, or additional station names.",
    "Use slowing rail sounds, one short bell, birds, and a soft completion tone only. No narration, subtitles, watermark, gibberish, or other station names.",
    "Use forest wind, birds, and gentle xylophone only. No chainsaw sound, narration, subtitles, watermark, or gibberish.",
    "Use four soft flash-card sounds and one final completion bell only. No speech.",
]

ZH_SFX = [
    "音效與配樂：開場先有一聲柔和列車鈴；全程以極低音量木琴和輕微模型列車行進聲襯底，旁白開始時自動壓低，結尾以柔和完成音收束；不要汽笛、月臺廣播或其他人聲。",
    "音效與配樂：使用低音量鳥鳴、微風與自然車站環境聲；不要可辨識廣播、列車汽笛或其他人聲；旁白開始時將環境聲明顯壓低。",
    "音效與配樂：開場使用輕微紙張展開聲，每亮一個節點加入一聲短促柔和的木琴提示音；不加火車汽笛，旁白期間所有提示音維持低音量。",
    "音效與配樂：使用柔和輪軌聲、微風、田野蟲鳴及少量溫暖木吉他和弦；不要尖銳汽笛或急促節奏，旁白開始時將環境聲和配樂壓低。",
    "音效與配樂：使用遠方流水、山谷微風、柔和輪軌聲及極低音量弦樂墊底；不要巨大水聲、強烈鼓點或汽笛，旁白保持最清楚並置於聲音中央。",
    "音效與配樂：使用樹葉沙沙聲與柔和輪軌聲；材料及農產品圖示微動時各加入一聲輕柔木塊提示音，並以極低音量木琴襯底；旁白開始時自動壓低。",
    "音效與配樂：使用列車減速輪軌聲、一聲短鈴及低音量月臺環境聲；不得出現可辨識廣播、喧鬧談話或其他人聲，列車聲在旁白開始時降低。",
    "音效與配樂：開場使用一聲柔和自行車鈴，配合鳥鳴與低音量小鎮街道聲；不要汽車喇叭、叫賣或可辨識交談，旁白始終最清楚。",
    "音效與配樂：使用山谷微風、遠方流水、柔和輪軌聲和低音量車站環境聲；不要廣播、汽笛或其他人聲，旁白期間背景聲降低。",
    "音效與配樂：使用列車減速聲、一聲短鈴、鳥鳴及柔和抵達完成音；不要尖銳煞車聲或汽笛，旁白開始後壓低列車聲。",
    "音效與配樂：使用林間微風、鳥鳴和柔和木琴；禁止電鋸、機械轟鳴或其他人聲，木琴只在旁白停頓處輕響。",
    "音效與配樂：說出四個站名時，各配一個音高不同但音色一致的柔和閃卡提示音，最後加入一聲完成鈴；不加背景音樂、汽笛或其他人聲，確保學生能清楚聽見站名與提問。",
]

EN_SFX = [
    "Sound design: begin with one soft train bell. Keep quiet xylophone and very light miniature-train movement underneath, automatically ducking both under the narration, then end with a gentle completion tone. No horn, platform announcement, or other human voice.",
    "Sound design: use low birds, light wind, and natural station ambience. No intelligible announcement, train horn, or other human voice. Duck the ambience clearly when the narration begins.",
    "Sound design: begin with a soft paper-opening sound and add one short gentle xylophone cue whenever a node lights up. No train horn. Keep every cue quiet under the narration.",
    "Sound design: use gentle rail rhythm, breeze, field insects, and a few warm acoustic-guitar chords. No sharp train horn or fast beat. Duck all ambience and music when the narration begins.",
    "Sound design: use distant flowing water, valley breeze, soft rail sounds, and a very quiet string pad. No loud water, strong drums, or horn. Keep the narration centered and clearly dominant.",
    "Sound design: use rustling leaves and soft rail sounds. Add one gentle wooden-block cue for the material icon and one for the farm-product icon, with very quiet xylophone underneath. Duck all sounds under the narration.",
    "Sound design: use slowing rail sounds, one short bell, and quiet platform ambience. No intelligible announcement, noisy conversation, or other human voice. Reduce the train sound when the narration begins.",
    "Sound design: begin with one soft bicycle bell, then use birds and quiet small-town street ambience. No car horn, vendor call, or intelligible conversation. Keep the narration clearly dominant.",
    "Sound design: use valley wind, distant water, soft rail sounds, and quiet station ambience. No announcement, horn, or other human voice. Duck all background sound under the narration.",
    "Sound design: use slowing rail sounds, one short bell, birds, and a soft arrival-completion tone. No sharp braking sound or train horn. Reduce the train sound after the narration begins.",
    "Sound design: use forest breeze, birds, and gentle xylophone. No chainsaw, machine roar, or other human voice. Let the xylophone sound only in brief narration pauses.",
    "Sound design: pair the four spoken station names with four soft flash-card tones of different pitch but matching timbre, followed by one completion bell. Use no background music, horn, or other human voice so every station name and the final question remain clear.",
]

TIMINGS = [
    ("0.3", "7.8"),
    ("0.5", "6.6"),
    ("0.5", "6.8"),
    ("0.5", "6.6"),
    ("0.5", "6.8"),
    ("0.3", "7.8"),
    ("0.5", "6.9"),
    ("0.4", "7.3"),
    ("0.5", "7.1"),
    ("0.8", "6.0"),
    ("0.5", "6.9"),
    ("0.3", "7.7"),
]


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None


def set_run_font(run, name: str, size: float, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = name
    run.font.size = Pt(size)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)


def set_labeled_paragraph(paragraph, label: str, body: str, body_font: str) -> None:
    paragraph.clear()
    label_run = paragraph.add_run(label)
    set_run_font(label_run, "Microsoft JhengHei", 10.5, bold=True)
    body_run = paragraph.add_run(body)
    set_run_font(body_run, body_font, 10.5, bold=False)


def set_body_paragraph(paragraph, text: str, font: str = "Microsoft JhengHei") -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, font, 10.5, bold=False)


def exact_replace(text: str, old: str, new: str, scene: int, language: str) -> str:
    if old not in text:
        raise RuntimeError(
            f"Storyboard {scene:02d} {language} audio sentence was not found."
        )
    return text.replace(old, new, 1)


def zh_voiceover(narration: str, start: str, end: str) -> str:
    return (
        f"中文旁白與語音規格：在第 {start} 秒開始、第 {end} 秒前自然說完；"
        "使用溫暖親切、適合國小四年級學生的臺灣華語教師聲線，語速稍慢、"
        "咬字清楚、自然有精神，不要機械聲、新聞播報腔或誇張戲劇語氣。"
        f"逐字說：「{narration}」不得翻譯、改寫、增刪或重複。"
        "只准出現這一段中文旁白，不要英語旁白、站內廣播或其他人聲。"
        "不要生成字幕、隱藏式字幕、逐字稿、卡拉 OK 字幕、標題卡、對話框，"
        "也不要讓任何文字隨語音出現；首幀原本已烘焙的站名與必要標示仍須完全鎖定不變。"
    )


def en_voiceover(narration: str, start: str, end: str) -> str:
    return (
        f"Mandarin voice-over: start at {start} seconds and finish naturally before "
        f"{end} seconds. Use a warm, friendly Taiwan Mandarin teacher voice suitable "
        "for Grade 4 children, with slightly slow pacing, clear articulation, and "
        "natural energy. Avoid a robotic tone, newsreader delivery, or exaggerated "
        f"acting. Speak exactly in Chinese: \"{narration}\" Do not translate, "
        "paraphrase, omit, add, or repeat any word. This exact Mandarin narration "
        "must be the only human voice. No English narration, station announcement, "
        "or other speech. Generate no subtitles, closed captions, transcript, "
        "karaoke text, title card, speech bubble, or any text synchronized with the "
        "voice. Keep only the text already baked into the input frame perfectly "
        "frozen and unchanged."
    )


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    doc = Document(DOCX)
    zh_paragraphs = [p for p in doc.paragraphs if p.text.startswith(ZH_LABEL)]
    en_paragraphs = [p for p in doc.paragraphs if p.text.startswith(EN_LABEL)]
    subtitle_paragraphs = [
        p for p in doc.paragraphs if p.text.startswith(SUBTITLE_LABEL)
    ]
    narration_paragraphs = [
        p for p in doc.paragraphs if p.text.startswith(NARRATION_LABEL)
    ]

    counts = (
        len(zh_paragraphs),
        len(en_paragraphs),
        len(subtitle_paragraphs),
        len(narration_paragraphs),
    )
    if counts != (12, 12, 12, 12):
        raise RuntimeError(f"Unexpected prompt counts: {counts}")

    narrations = [
        p.text[len(NARRATION_LABEL) :].strip() for p in narration_paragraphs
    ]

    for index in range(12):
        scene = index + 1
        start, end = TIMINGS[index]

        zh_body = zh_paragraphs[index].text[len(ZH_LABEL) :].strip()
        zh_audio = (
            zh_voiceover(narrations[index], start, end)
            + ZH_SFX[index]
            + "混音要求：中文旁白音量始終高於配樂與環境聲，語音清楚、無破音、無回音；整段維持 8 秒，單一連續鏡頭。"
        )
        zh_body = exact_replace(
            zh_body, ZH_OLD_AUDIO[index], zh_audio, scene, "Chinese"
        )
        set_labeled_paragraph(
            zh_paragraphs[index], ZH_LABEL, zh_body, "Microsoft JhengHei"
        )

        en_body = en_paragraphs[index].text[len(EN_LABEL) :].strip()
        en_audio = (
            en_voiceover(narrations[index], start, end)
            + EN_SFX[index]
            + "Mixing requirement: keep the Mandarin narration clearly louder than "
            "music and ambience, with no clipping or echo. Keep the complete clip "
            "eight seconds long as one continuous shot."
        )
        en_body = exact_replace(
            en_body, EN_OLD_AUDIO[index], en_audio, scene, "English"
        )
        set_labeled_paragraph(en_paragraphs[index], EN_LABEL, en_body, "Comic Sans MS")

        set_labeled_paragraph(
            narration_paragraphs[index],
            NEW_NARRATION_LABEL,
            narrations[index],
            "Microsoft JhengHei",
        )

    for paragraph in subtitle_paragraphs:
        remove_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        text = paragraph.text
        replacement = None
        if text == (
            "每張分鏡皆提供中文版與英文版 Flow 動態提示詞。Flow 只負責動態，"
            "不生成站名、字幕、旁白、對話框或鏡間轉場；需要的文字先正確烘焙於"
            "靜態首幀，精確字幕與旁白依各分鏡所列內容後製加入。"
        ):
            replacement = (
                "每張分鏡皆提供中文版與英文版 Flow 動態提示詞。Flow 負責畫面動態、"
                "符合場景的背景音效及指定的中文旁白；每組提示詞均內嵌逐字旁白、"
                "臺灣華語聲線、開始與結束時間及混音要求。影片不加入字幕、隱藏式字幕、"
                "逐字稿或隨語音出現的文字；必要站名先正確烘焙於靜態首幀並全程鎖定。"
            )
        elif text == (
            "輸出1920×1080、H.264 MP4、25或30fps；旁白峰值清楚高於環境聲，"
            "字幕至少42px、白字深色描邊，置於下方安全區。"
        ):
            replacement = (
                "輸出1920×1080、H.264 MP4、25或30fps；使用溫暖親切、咬字清楚的"
                "臺灣華語教師旁白，旁白峰值清楚高於環境聲。成片不建立字幕軌，"
                "畫面也不得生成字幕、逐字稿、標題卡或卡拉 OK 文字。"
            )
        elif text.startswith("Flow一致性Ingredient規格："):
            replacement = text.replace(
                "天空方向、路線圖符號與字幕安全區一致",
                "天空方向、路線圖符號、旁白聲線與視覺安全區一致",
            )
        elif text.startswith("個別分鏡靜態提示詞："):
            replacement = text.replace(
                "站名、路線線條、字幕與箭頭只保留乾淨安全區，文字後製，不讓生成模型自行亂寫。",
                "站名、路線線條與箭頭只保留乾淨安全區；影片不預留或生成字幕區，不讓生成模型自行亂寫文字。",
            ).replace(
                "天空方向、路線圖符號與字幕安全區一致",
                "天空方向、路線圖符號與視覺安全區一致",
            )
        elif text.startswith("單鏡驗收｜"):
            replacement = text.replace(
                "字幕無錯字；聲音不蓋過旁白",
                "未生成字幕、逐字稿或隨語音出現的文字；中文旁白逐字正確、聲線溫暖親切；背景音效不蓋過旁白",
            )

        if replacement is not None and replacement != text:
            set_body_paragraph(paragraph, replacement)

    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup = DOCX.with_name(
        f"{DOCX.stem}_整合中文旁白前備份_{timestamp}{DOCX.suffix}"
    )
    shutil.copy2(DOCX, backup)

    temp = DOCX.with_name(f"{DOCX.stem}.voiceover-temp{DOCX.suffix}")
    doc.save(temp)
    Document(temp)
    temp.replace(DOCX)

    print(f"Updated: {DOCX}")
    print(f"Backup: {backup}")


if __name__ == "__main__":
    main()
