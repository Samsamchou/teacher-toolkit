from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "在地課程4年級上學期教案" / "02_第3-10週_坐火車趣集集.docx"
START = "第4節數位教材｜iPad模擬線上購票網站設計規劃書"
END = "第5節自製教材｜站務角色、六步搭車流程與危險動作圖卡"

FONT_LATIN = "Arial"
FONT_CJK = "Microsoft JhengHei"
NAVY = "17324D"
BLUE = "1E6FA8"
LIGHT_BLUE = "EAF3F8"
LIGHT_CREAM = "FFF8E8"
LIGHT_GRAY = "F4F6F8"
WHITE = "FFFFFF"
TEXT = "24313D"


def set_run_font(run, size=10.5, bold=False, color=TEXT, latin=FONT_LATIN):
    run.bold = bold
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.color.rgb = None
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    color_el = rpr.find(qn("w:color"))
    if color_el is None:
        color_el = OxmlElement("w:color")
        rpr.append(color_el)
    color_el.set(qn("w:val"), color)


def set_paragraph_text(paragraph, text, size=10.5, bold=False, color=TEXT):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def keep_with_next(paragraph, value=True):
    ppr = paragraph._p.get_or_add_pPr()
    element = ppr.find(qn("w:keepNext"))
    if value and element is None:
        ppr.append(OxmlElement("w:keepNext"))
    elif not value and element is not None:
        ppr.remove(element)


def cant_split(row):
    trpr = row._tr.get_or_add_trPr()
    if trpr.find(qn("w:cantSplit")) is None:
        trpr.append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    if trpr.find(qn("w:tblHeader")) is None:
        trpr.append(OxmlElement("w:tblHeader"))


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tcmar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, inches):
    width = int(Inches(inches))
    cell.width = width
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(round(inches * 1440)))
    tcw.set(qn("w:type"), "dxa")


def format_cell(cell, text, header=False, size=9.2, align=None, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size=size,
        bold=header,
        color=WHITE if header else TEXT,
    )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if fill:
        shade_cell(cell, fill)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tablepr = table._tbl.tblPr
    tblw = tablepr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tablepr.append(tblw)
    tblw.set(qn("w:w"), str(round(sum(widths) * 1440)))
    tblw.set(qn("w:type"), "dxa")
    tblind = tablepr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tablepr.append(tblind)
    tblind.set(qn("w:w"), "0")
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)
    for row in table.rows:
        cant_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])


def add_heading(doc, text, level=3):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_text(paragraph, text, size=12 if level == 3 else 14, bold=True, color=NAVY)
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    keep_with_next(paragraph)
    return paragraph


def add_body(doc, text, style="Normal", bold_prefix=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2.5)
    paragraph.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest)
    else:
        set_paragraph_text(paragraph, text)
    return paragraph


def add_table(doc, headers, rows, widths, body_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        format_cell(
            table.rows[0].cells[idx],
            header,
            header=True,
            size=9.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=NAVY,
        )
    repeat_header(table.rows[0])
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for col_index, value in enumerate(values):
            fill = LIGHT_BLUE if row_index % 2 else WHITE
            align = WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            format_cell(
                cells[col_index],
                value,
                header=False,
                size=body_size,
                align=align,
                fill=fill,
            )
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def remove_between(start_el, end_el):
    current = start_el.getnext()
    while current is not None and current is not end_el:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt


def move_new_elements_before(doc, end_el, original_body_children):
    original_ids = {id(x) for x in original_body_children}
    new_elements = [
        child
        for child in list(doc.element.body)
        if id(child) not in original_ids and child is not doc.element.body.sectPr
    ]
    for element in new_elements:
        end_el.addprevious(element)


def main():
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    doc = Document(DOCX)
    start_p = next(p for p in doc.paragraphs if p.text.startswith(START))
    end_p = next(p for p in doc.paragraphs if p.text.startswith(END))
    start_el = start_p._p
    end_el = end_p._p
    original_body_children = list(doc.element.body)

    set_paragraph_text(
        start_p,
        "第4節數位教材｜iPad模擬線上購票網站設計規劃書（ChatGPT Sites完善版）",
        size=14,
        bold=True,
        color=NAVY,
    )
    start_p.style = "Heading 2"
    keep_with_next(start_p)
    remove_between(start_el, end_el)

    add_heading(doc, "一、教學定位、平台與完成成果")
    add_table(
        doc,
        ["項目", "完整規格"],
        [
            (
                "課堂目標",
                "學生依序完成起點、目的地、日期、查詢、選車次與確認摘要六步，能發現起訖顛倒並修正；到最終SAMPLE車票即停止，不進入真實訂票。",
            ),
            (
                "使用情境",
                "教師投影示範1次；兩人共用1台iPad，Operator操作、Checker逐項說Check from／to／date／train／time；第2輪交換角色。",
            ),
            (
                "學生端成果",
                "每一步正確後取得過關畫面；六步加最終成功車票共七張，由系統合成一份七頁PDF。",
            ),
            (
                "教師端成果",
                "依日期與學號開啟紀錄，查看摘要、依時間順序重播完整操作事件，並預覽或下載七頁PDF。",
            ),
            (
                "平台",
                "網站規劃部署於ChatGPT Sites；以Cloudflare D1保存結構化紀錄，以R2保存PDF檔案。D1與R2只由伺服器端存取。",
            ),
            (
                "資料最小化",
                "學生只輸入五位數學號，例如40100；不輸入姓名、身分證、電話、電子郵件、信用卡、付款或真實訂票代碼，也不保存IP位址。",
            ),
            (
                "非目標",
                "不連接臺鐵真實訂票、不付款、不建立學生會員、不錄製真實螢幕影片、麥克風或系統聲音。",
            ),
            (
                "本階段邊界",
                "本次只完善Word規劃書；尚不建立ChatGPT Sites專案、不生成正式網站素材，也不部署網站。",
            ),
        ],
        [1.35, 5.15],
    )

    add_heading(doc, "二、上層網站「4年級上學期 在地課程」")
    add_body(
        doc,
        "首頁上方使用深藍色導覽列，左側顯示「4年級上學期 在地課程」，右上角固定設置「教師後台」；學生不需登入即可進入四個單元，教師後台則必須完成ChatGPT登入與電子郵件白名單檢核。",
        style="List Bullet",
    )
    add_table(
        doc,
        ["單元入口", "首頁卡片內容", "本次資料接入狀態"],
        [
            ("扇形車庫", "扇形軌道、轉車臺與車庫門意象。", "預留單元路由與成果類型，日後另訂。"),
            ("坐火車趣集集", "集集小火車、田野、綠色隧道與山谷意象。", "本次完整定義學生流程、事件重播與七頁PDF。"),
            ("閱覽鐵道風華", "山線、海線、路線圖與閱讀卡意象。", "預留單元路由與成果類型，日後另訂。"),
            (
                "介紹五分車與認識小火車鐵道",
                "五分車、窄軌與糖業運輸意象。",
                "預留單元路由與成果類型，日後另訂。",
            ),
        ],
        [1.65, 2.35, 2.5],
    )
    add_body(
        doc,
        "共用教師後台必須能用單元、日期與學號篩選紀錄；各單元共用attempt ID、保存期限、教師授權與刪除機制，但可擁有不同的成果檔案及重播規則。",
        style="List Bullet",
    )

    add_heading(doc, "三、購票網站首頁、版面與ImageGen素材")
    add_body(
        doc,
        "網站首頁名稱：火車線上購票網站—Buy Train Tickets Online。標題下方放置大尺寸學號輸入框，placeholder為「請輸入學號，例如40100」；只接受五位數字，驗證通過後由系統建立不重複的attempt ID。",
        style="List Bullet",
    )
    add_body(
        doc,
        "首頁固定顯示警語：「模擬教材，不可作為真實乘車或訂票資訊。」開始按鈕使用「Start／開始練習」，不使用「購買」「付款」或「送出訂票」字樣。",
        style="List Bullet",
    )
    add_table(
        doc,
        ["ImageGen原創場景", "靜態畫面內容", "網站輕量動畫方式"],
        [
            ("1 集集小火車", "可愛紅黃小火車沿支線前進，人物位於安全區。", "CSS讓火車緩慢橫向移動；減少動態模式改為靜止。"),
            ("2 綠色隧道", "枝葉形成拱形綠色隧道，陽光柔和穿過樹葉。", "分層樹葉小幅搖晃，不使用連續閃爍。"),
            ("3 二水田野", "稻田、水圳、遠山及小火車，符合二水出發情境。", "雲朵慢移、稻葉輕擺，低耗能循環。"),
            ("4 車埕木造站", "木造站體、山林與終點氣氛，不畫都市高樓。", "光點及鳥群緩慢移動，避免影片檔造成iPad載入負擔。"),
        ],
        [1.3, 2.9, 2.3],
    )
    add_table(
        doc,
        ["色彩角色", "教學版建議值", "使用規則"],
        [
            ("深藍導覽", "#17324D", "承接臺鐵官方網站的穩重交通服務感，用於上方導覽與主要標題。"),
            ("票務藍", "#1E6FA8", "主要按鈕、目前步驟與可操作欄位；不得作為唯一狀態線索。"),
            ("暖橘強調", "#F4A62A", "路線、提示箭頭與重點數字；錯誤訊息不使用羞辱性的大片紅色。"),
            ("成功綠", "#2E8B57", "「You’re right!」、勾號及通過邊框。"),
            ("背景與文字", "#F7FAFC／#24313D", "保持高對比與投影可讀性，搭配兒童感圓角卡片。"),
        ],
        [1.3, 1.35, 3.85],
    )
    add_body(
        doc,
        "配色查證來源（教師端）：國營臺灣鐵路股份有限公司「個人訂票」官方頁面，查詢日期2026-07-29：https://www.railway.gov.tw/tra-tip-web/tip/tip001/tip121/query。上述色碼是兒童教學版建議值，不宣稱為臺鐵官方色碼；實作時只參考其票務層級與交通服務氛圍，不複製商標、真實票面或完整介面。",
        style="Normal",
        bold_prefix="配色查證來源（教師端）：",
    )

    add_heading(doc, "四、學生六步流程、正確回饋與七張證據畫面")
    add_body(
        doc,
        "學生輸入學號後才建立attempt；首頁登入畫面不計入PDF。每一步只有在必要欄位與條件完全正確時才通過，系統依序寫入step_passed事件並擷取網站內容區。",
        style="List Bullet",
    )
    add_table(
        doc,
        ["步驟／畫面", "學生操作", "通過條件、回饋與證據"],
        [
            (
                "1 選起點",
                "選擇二水 Ershui。",
                "二水節點亮起；播放「You’re right!」、顯示勾號與短動畫；擷取PDF第1頁。",
            ),
            (
                "2 選目的地",
                "可選集集、水里或車埕；任務主線選車埕。",
                "目的地與起點不同且符合任務卡；路線亮到目的地；擷取第2頁。",
            ),
            (
                "3 選日期",
                "從三個虛構教學日期選擇指定日期。",
                "醒目標示「模擬日期」；日期正確才通過；擷取第3頁。",
            ),
            (
                "4 查詢",
                "檢核起點、目的地與日期後按查詢。",
                "起訖相同或顛倒時只給友善修正提示；全部正確才慶祝並擷取第4頁。",
            ),
            (
                "5 選車次",
                "從A／B／C模擬班次選擇任務指定車次。",
                "卡片顯示出發、抵達與所需時間及SAMPLE；正確後擷取第5頁。",
            ),
            (
                "6 確認摘要",
                "逐項核對from、to、date、train、depart、arrive六欄。",
                "六欄全勾且資料一致才完成；擷取第6頁，不出現付款或真實訂票按鈕。",
            ),
            (
                "7 最終成功",
                "閱讀SAMPLE車票及「練習完成，沒有真的訂票」。",
                "播放最後完成鈴與較完整但不閃爍的慶祝動畫；擷取第7頁並開始合成PDF。",
            ),
        ],
        [1.1, 1.95, 3.45],
        body_size=8.8,
    )
    add_body(
        doc,
        "每次正確回饋同時包含可見文字「You’re right!」、綠色勾號、1至1.5秒慶祝動畫及短音效；音效有靜音鍵。啟用prefers-reduced-motion時，以靜態勾號與淡入取代彩帶或位移。",
        style="List Bullet",
    )
    add_body(
        doc,
        "錯誤時不播放成功音效；以橘色文字說明可修正的欄位，焦點移到該欄並透過aria-live朗讀。錯誤、交換起訖及回上一步都要寫入事件紀錄，以呈現從頭到尾的真實學習歷程。",
        style="List Bullet",
    )

    add_heading(doc, "五、七頁PDF的擷取、合成與上傳")
    add_body(
        doc,
        "每次step_passed後等待版面與慶祝動畫穩定，再擷取購票網站主要內容容器；不得擷取瀏覽器網址列、其他分頁、通知、聲音或裝置畫面。",
        style="List Bullet",
    )
    add_body(
        doc,
        "每頁採A4橫式，放一張過關畫面，頁首標示單元、日期、學號、attempt ID末六碼、步驟名稱與第幾頁；畫面內保留「模擬教材」警語。七頁順序固定，不可缺頁或重複。",
        style="List Bullet",
    )
    add_body(
        doc,
        "七張壓縮圖片先在當次操作暫存，完成後於學生端合成單一PDF並送至受保護的伺服器端點；伺服器確認attempt、頁數、大小與雜湊後寫入R2，再把R2索引寫回D1。PDF成功後清除個別暫存圖。",
        style="List Bullet",
    )
    add_body(
        doc,
        "若網路暫時中斷，學生仍可完成離線練習；頁面顯示「等待同步」，恢復連線後重試事件與PDF上傳。未成功同步前不得顯示「教師已收到」。",
        style="List Bullet",
    )

    add_heading(doc, "六、操作事件錄製與無聲動畫重播")
    add_body(
        doc,
        "本功能是「操作事件錄製」，不是iPad真實螢幕錄影：網站保存點擊、選項、修正、交換起訖、過關與完成等結構化事件，再由教師後台重建相同畫面並依時間播放。",
        style="List Bullet",
    )
    add_table(
        doc,
        ["事件欄位", "內容與限制"],
        [
            ("attempt_id", "伺服器產生的不重複識別碼；不可只用學號當主鍵。"),
            ("seq", "由1開始的連續流水號；伺服器拒絕重複或倒序事件。"),
            ("step", "student_id、origin、destination、date、search、train、summary、success。"),
            ("action", "attempt_started、field_selected、validation_failed、swap、back、step_passed、attempt_completed。"),
            ("payload_json", "只接受欄位白名單與模擬值，不接受任意文字、姓名或真實票務資料。"),
            ("client_elapsed_ms", "從attempt開始計算的相對時間，用於重播節奏，不依賴裝置時鐘。"),
            ("server_received_at", "伺服器時間，作為排序、稽核與一年保存期限的依據。"),
            ("before／after", "保存操作前後的允許狀態，教師能看見學生如何修正。"),
        ],
        [1.55, 4.95],
    )
    add_body(
        doc,
        "教師重播播放器提供播放／暫停、0.5×／1×／2×、時間軸及步驟跳轉；使用游標光圈與欄位高亮呈現操作，不播放學生端音效，也不生成影片檔。",
        style="List Bullet",
    )
    add_body(
        doc,
        "attempt完成後鎖定既有事件，不允許前端改寫；若需重做，建立新的attempt。相同日期與學號可有多筆，左欄另顯示完成時間和attempt末六碼。",
        style="List Bullet",
    )

    add_heading(doc, "七、ChatGPT Sites、D1與R2資料契約")
    add_table(
        doc,
        ["D1資料表", "主要欄位", "用途"],
        [
            (
                "teacher_allowlist",
                "email、role、active、created_at、updated_at",
                "允許使用教師後台的ChatGPT帳號；只由伺服器端查詢。",
            ),
            (
                "units",
                "slug、display_name、version、evidence_policy",
                "四單元共用設定；本次tickets單元使用event_replay＋seven_page_pdf。",
            ),
            (
                "attempts",
                "attempt_id、student_id、unit_slug、started_at、completed_at、status、expires_at、event_count、pdf_key",
                "一次完整練習的主紀錄與一年到期日。",
            ),
            (
                "attempt_events",
                "attempt_id、seq、step、action、payload_json、client_elapsed_ms、server_received_at",
                "從頭到尾的操作與修正歷程；依attempt_id＋seq建立唯一索引。",
            ),
            (
                "evidence_manifest",
                "attempt_id、page_no、step_key、captured_at、checksum、upload_status",
                "驗證七頁順序、完整性及PDF建立狀態。",
            ),
            (
                "deletion_log",
                "attempt_id、reason、requested_by、deleted_at、r2_result",
                "記錄一年到期或教師手動刪除結果，不保留被刪除的學習內容。",
            ),
        ],
        [1.35, 3.0, 2.15],
        body_size=8.4,
    )
    add_table(
        doc,
        ["R2項目", "規格"],
        [
            (
                "物件路徑",
                "evidence/{unit_slug}/{attempt_id}/proof.pdf；路徑使用不透明attempt ID，不直接放學號。",
            ),
            (
                "物件metadata",
                "attempt_id、page_count=7、checksum、created_at、expires_at；PDF不設公開網址。",
            ),
            (
                "讀取",
                "只有通過教師身分及白名單檢查的伺服器端點可串流預覽或下載。",
            ),
            (
                "刪除",
                "完成日起365天到期；先刪R2物件並確認成功，再清除D1事件、manifest與attempt內容。",
            ),
        ],
        [1.35, 5.15],
    )
    add_body(
        doc,
        "ChatGPT Sites專案在.openai/hosting.json宣告邏輯D1綁定DB與R2綁定EVIDENCE_BUCKET；正式Cloudflare資源與部署連線由Sites管理。D1查詢使用參數化prepared statements，R2與D1不得直接暴露給瀏覽器。",
        style="List Bullet",
    )

    add_heading(doc, "八、教師登入、權限與後台版面")
    add_body(
        doc,
        "學生路由公開；教師進入「教師後台」時使用ChatGPT登入。伺服器取得已驗證的帳號電子郵件，再查詢D1 teacher_allowlist；未登入或不在白名單者一律拒絕，不能以?teacher=1、隱藏網址或前端密碼代替授權。",
        style="List Bullet",
    )
    add_table(
        doc,
        ["區域", "畫面內容", "教師操作"],
        [
            (
                "頂部工具列",
                "四單元篩選、日期範圍、學號搜尋、同步狀態、登入帳號。",
                "切換單元、搜尋、匯出清單及安全登出。",
            ),
            (
                "左欄紀錄列",
                "依日期群組，顯示學號、完成時間、完成／待同步狀態及attempt末六碼。",
                "點選一筆後在右欄顯示；同一學號多次作答不可互相覆蓋。",
            ),
            (
                "右欄摘要",
                "起訖、日期、車次、開始／完成時間、錯誤與修正次數、七頁狀態。",
                "快速確認是否完整及是否需要重做。",
            ),
            (
                "右欄重播",
                "無聲事件動畫、時間軸、播放速度、事件清單。",
                "播放、暫停、跳到某一步，查看前後狀態。",
            ),
            (
                "右欄PDF",
                "七頁縮圖、頁碼與PDF預覽。",
                "預覽、下載；刪除前二次確認並同步清除D1／R2。",
            ),
        ],
        [1.25, 3.0, 2.25],
        body_size=8.8,
    )

    add_heading(doc, "九、一年保存、個資與安全規則")
    for text in [
        "保存期限由attempt完成時間起算365天；未完成attempt由最後事件時間起算365天。教師可在期限內下載PDF或匯出紀錄，也可提前手動刪除。",
        "到期清理由平台可用的週期性機制執行；若Sites未提供排程，教師後台開啟時執行有數量上限的到期掃描，並保留「立即清理到期資料」按鈕作為備援。",
        "學生學號屬於教學紀錄，畫面與PDF只在必要位置顯示；R2物件名稱不含學號。教師電子郵件只用於後台授權與刪除稽核。",
        "學生端只可建立當次attempt及追加具有短效attempt token的事件，不能列出、查詢或下載其他學生紀錄；所有教師查詢、PDF讀取、匯出及刪除都在伺服器端再次授權。",
        "限制事件種類、欄位長度、請求頻率與PDF大小；拒絕任意HTML、檔案名稱、外部網址及真實票務資料，並使用參數化D1查詢。",
    ]:
        add_body(doc, text, style="List Bullet")

    add_heading(doc, "十、iPad、無障礙、離線與故障替代")
    for text in [
        "iPad直式與橫式皆可用；主要內容寬度不超過960px，按鈕觸控區至少48×48px、間距至少8px、學生正文至少18px，Safari與Chrome不得橫向溢出。",
        "每一步使用h1標題、可見鍵盤焦點、正確label及aria-live；可用Tab／Enter完成，不要求拖曳，顏色同時搭配站名、圖示、形狀與文字。",
        "You’re right!音效提供全程靜音鍵；慶祝動畫不閃爍並尊重prefers-reduced-motion。ImageGen圖像不承載站名、時間、錯誤訊息或按鈕文字。",
        "PWA快取核心頁面與模擬資料；離線時允許練習及暫存事件／七張圖，但清楚顯示「離線練習，尚未同步」。恢復連線後依seq重試，不重複建立attempt。",
        "若iPad無法開啟，改用同版面A4六步紙卡、三張車次卡與SAMPLE確認票；教師課前以至少兩台iPad測試QR、音量、旋轉、重新整理、離線及重新連線。",
    ]:
        add_body(doc, text, style="List Bullet")

    add_heading(doc, "十一、驗收測試案例")
    tests = [
        "標準流程：輸入40100，完成二水→車埕→模擬日期→查詢→車次B→六欄確認；每一步皆有You're right!，最後產生完整七頁PDF。",
        "正確回饋：每個step_passed只播放一次成功音效與動畫；靜音及減少動態效果開啟後仍有文字和勾號。",
        "錯誤修正：起訖相同、起訖顛倒、錯誤日期或車次不會慶祝；事件重播能看見提示、交換及重新選擇。",
        "事件完整：seq連續、前後狀態一致、attempt完成後不可改寫；教師重播順序與實際操作一致。",
        "相同學號：同一天以40100完成兩次，後台顯示兩筆不同attempt與時間，不互相覆蓋。",
        "七頁證據：PDF依六步加最終成功排序，無重複、無缺頁、無瀏覽器UI；R2 metadata的page_count為7。",
        "教師授權：未登入、非白名單帳號及直接呼叫API都被拒絕；白名單教師可查看但學生端不能列出紀錄。",
        "後台版面：左欄依日期及學號顯示；點選後右欄可切換摘要、重播與PDF，iPad與桌機皆可操作。",
        "保存與刪除：到期測試能先刪R2再清D1；手動刪除需二次確認，刪除後連結無法再讀取。",
        "離線恢復：離線完成後顯示未同步；恢復網路可補傳事件及PDF，不產生第二份attempt或重複頁面。",
        "隱私與模擬邊界：網路請求、D1與R2不含姓名、IP、電話、付款或真實訂票資訊；所有頁面均顯示模擬教材警語。",
    ]
    for text in tests:
        add_body(doc, text, style="List Number")

    add_heading(doc, "十二、未來ChatGPT Sites實作檔案與路由")
    add_table(
        doc,
        ["項目", "規劃"],
        [
            ("/", "四單元首頁與右上角教師後台入口。"),
            ("/units/train-tickets", "學生學號首頁、六步購票與最終SAMPLE票。"),
            ("/teacher", "ChatGPT登入、D1白名單檢核及左右欄教師後台。"),
            ("/api/attempts／events／evidence", "學生建立attempt、追加事件與上傳PDF的受限伺服器端點。"),
            ("/api/teacher/*", "教師查詢、重播資料、PDF串流、匯出及刪除端點；每次請求皆重新授權。"),
            ("db/schema.ts／migrations", "D1資料表、索引及遷移；每個prepared statement只執行一個SQL statement。"),
            (".openai/hosting.json", "宣告ChatGPT Sites專案、D1綁定DB與R2綁定EVIDENCE_BUCKET。"),
            ("public/assets/home", "四張ImageGen場景圖及網站所需音效；精確文字全部由HTML呈現。"),
            ("文件", "README_TEACHER.md、QA_CHECKLIST.md、PRIVACY_RETENTION.md及素材來源紀錄。"),
        ],
        [2.0, 4.5],
        body_size=8.8,
    )
    add_body(
        doc,
        "本規劃通過後，未來才進入ChatGPT Sites建置、ImageGen素材生成、iPad實測與正式部署；建立Sites專案時必須沿用平台產生的project ID，D1與R2的正式資源由Sites配置，不自行臆造識別碼。",
        style="Normal",
        bold_prefix="本規劃通過後，",
    )

    move_new_elements_before(doc, end_el, original_body_children)

    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup = DOCX.with_name(
        f"{DOCX.stem}_更新ChatGPT_Sites購票規劃前備份_{timestamp}{DOCX.suffix}"
    )
    shutil.copy2(DOCX, backup)
    temp = DOCX.with_name(f"{DOCX.stem}.sites-plan-temp{DOCX.suffix}")
    doc.save(temp)
    Document(temp)
    temp.replace(DOCX)
    print(f"Updated: {DOCX}")
    print(f"Backup: {backup}")


if __name__ == "__main__":
    main()
