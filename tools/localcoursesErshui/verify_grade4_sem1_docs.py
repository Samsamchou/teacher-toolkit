from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parent
FOLDER = ROOT / "在地課程4年級上學期教案"
EXPECTED = {
    "01_第1-2週_扇形車庫.docx": 2,
    "02_第3-10週_坐火車趣集集.docx": 8,
    "03_第11-14週_閱覽鐵道風華.docx": 4,
    "04_第15-20週_介紹五分車與認識小火車鐵道.docx": 6,
}


def main():
    failures = []
    for name, expected_lessons in EXPECTED.items():
        path = FOLDER / name
        if not path.exists():
            failures.append(f"missing: {name}")
            continue
        with ZipFile(path) as archive:
            bad = archive.testzip()
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        text += "\n" + "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        lesson_titles = [
            p.text for p in doc.paragraphs if p.style.name == "Lesson Title"
        ]
        stage_rows = sum(
            1
            for table in doc.tables
            for row in table.rows
            if row.cells
            and row.cells[0].text.startswith(
                ("Warm-up", "Presentation", "Production", "Wrap-up")
            )
        )
        links = {
            rel.target_ref
            for rel in doc.part.rels.values()
            if "hyperlink" in rel.reltype
        }
        checks = {
            "zip_ok": bad is None,
            "lesson_count": len(lesson_titles) == expected_lessons,
            "stage_count": stage_rows == expected_lessons * 4,
            "prep_table_terms": all(
                term in text
                for term in (
                    "主要英語口說／句型",
                    "生活用語",
                    "核心英文字詞",
                    "數位教材",
                    "手作教材",
                    "本節學習單",
                )
            ),
            "flow_table_terms": all(
                term in text
                for term in (
                    "教師如何教與如何提問",
                    "學生如何學與如何互動",
                    "形成性評量／成果",
                )
            ),
            "curriculum_sections": all(
                term in text
                for term in (
                    "單元定位與核心成果",
                    "學習目標",
                    "英語聽說評量手冊融入",
                    "建議影片、繪本、照片與資料",
                    "領域課綱對應",
                    "逐節教學流程",
                )
            ),
            "curriculum_columns": all(
                term in text
                for term in (
                    "學習表現（代碼＋完整敘述）",
                    "學習內容（代碼＋完整敘述）",
                    "在本單元中的具體證據",
                )
            ),
            "links_present": len(links) > 0,
            "grade3_correction": (
                "扇形車庫" not in name
                or ("What’s this?" in text and "What’s that?" in text)
            ),
            "no_wrong_main_pattern": (
                "扇形車庫" not in name or "What is it? It’s a" not in text
            ),
            "one_day_change": (
                "閱覽鐵道風華" not in name
                or ("只完成一日行程" in text and "完成三天旅遊" not in text)
            ),
        }
        failed = [key for key, value in checks.items() if not value]
        if failed:
            failures.append(f"{name}: {', '.join(failed)}")
        print(
            f"{name}\n"
            f"  bytes={path.stat().st_size} paragraphs={len(doc.paragraphs)} "
            f"tables={len(doc.tables)} lessons={len(lesson_titles)} "
            f"flow_rows={stage_rows} links={len(links)}\n"
            f"  checks={'PASS' if not failed else 'FAIL ' + ','.join(failed)}"
        )
    if failures:
        print("\nFAILURES")
        for item in failures:
            print(item)
        raise SystemExit(1)
    print("\nALL STRUCTURAL CHECKS PASSED")


if __name__ == "__main__":
    main()
