from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path.cwd() / "二水國小四年級在地課程_Animals融入教學詳案_教室資料探究版.docx"
WORKSHEET_IMAGE = Path.cwd() / "assets" / "animals_worksheet" / "學習單一_我的動物朋友搭上二水小火車_3D日式動漫.png"
FONT = "Microsoft JhengHei"
NAVY = "1F4E79"
TEAL = "148A8A"
AMBER = "F4B183"
PALE_BLUE = "EAF3F8"
PALE_TEAL = "E7F5F4"
PALE_AMBER = "FFF1E5"
GREY = "F3F5F7"
DARK = "263238"
WHITE = "FFFFFF"


def set_east_asia(element, font=FONT):
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="B7C9D6", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=95, start=105, bottom=95, end=105):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_col_width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run(run, 8, color="6B7A86")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 頁")
    set_run(run, 8, color="6B7A86")


def set_run(run, size=10.5, bold=False, color=DARK, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    set_east_asia(run._element, FONT)
    return run


def set_paragraph(paragraph, before=0, after=4, line=1.25, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep:
        fmt.keep_with_next = True
    return paragraph


def add_text(doc, text, size=10.5, bold=False, color=DARK, align=None, before=0, after=4, line=1.25, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph(p, before, after, line, keep)
    set_run(p.add_run(text), size, bold, color)
    return p


def add_bullet(doc, text, level=0, color=DARK):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_paragraph(p, after=2, line=1.2)
    set_run(p.add_run(text), 10.2, color=color)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph(p, after=2, line=1.2)
    set_run(p.add_run(text), 10.2)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph(p, before=11 if level == 1 else 7, after=5, line=1.0, keep=True)
    if level == 1:
        run = p.add_run(text)
        set_run(run, 15, True, NAVY)
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "14")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), TEAL)
        borders.append(bottom)
        p_pr.append(borders)
    elif level == 2:
        run = p.add_run(text)
        set_run(run, 12.5, True, TEAL)
    else:
        run = p.add_run(text)
        set_run(run, 11, True, NAVY)
    return p


def add_callout(doc, title, body, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "B7D8D8")
    set_cell_margins(cell, 115, 165, 115, 165)
    p = cell.paragraphs[0]
    set_paragraph(p, after=2, line=1.2)
    set_run(p.add_run(title + "　"), 10.5, True, NAVY)
    set_run(p.add_run(body), 10.2)
    return table


def add_key_value_table(doc, rows, widths=(3.2, 13.9)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for cell in cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_shading(cells[0], PALE_BLUE)
        set_col_width(cells[0], widths[0])
        set_col_width(cells[1], widths[1])
        p = cells[0].paragraphs[0]
        set_paragraph(p, after=0)
        set_run(p.add_run(key), 9.8, True, NAVY)
        p = cells[1].paragraphs[0]
        set_paragraph(p, after=0, line=1.15)
        set_run(p.add_run(value), 9.7)
    return table


def add_grid_table(doc, headers, rows, widths=None, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    prevent_row_split(hdr)
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, NAVY)
        set_cell_margins(cell, 85, 90, 85, 90)
        if widths:
            set_col_width(cell, widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, after=0)
        set_run(p.add_run(value), font_size, True, WHITE)
    for row in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_border(cell)
            set_cell_margins(cell, 75, 85, 75, 85)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if widths:
                set_col_width(cell, widths[i])
            if len(table.rows) % 2 == 1:
                set_cell_shading(cell, "F8FBFC")
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.12)
            set_run(p.add_run(str(value)), font_size)
    return table


def add_assessment_table(doc, title, rows):
    add_heading(doc, title, 2)
    return add_grid_table(doc, ["面向", "可觀察證據", "達成標準"], rows, (3.0, 10.0, 4.0), 9.0)


def add_question_table(doc, title, rows):
    add_heading(doc, title, 2)
    add_text(doc, "使用方式：教師可口頭讀題，或作為紙本／數位檢核；每題均為三選一。", 9.2, color="60717C", after=3)
    data = []
    for number, question, a, b, c, answer in rows:
        data.append([number, question, "A. " + a, "B. " + b, "C. " + c, answer])
    return add_grid_table(doc, ["題號", "題目", "選項 A", "選項 B", "選項 C", "答案"], data,
                          (0.9, 5.9, 3.1, 3.1, 3.1, 1.2), 8.1)


def add_lesson_overview(doc, number, weeks, title, local_focus, language_focus, outcome):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_TEAL)
    set_cell_border(cell, TEAL, "14")
    set_cell_margins(cell, 140, 150, 130, 150)
    p = cell.paragraphs[0]
    set_paragraph(p, after=3)
    set_run(p.add_run(f"第 {number} 節｜{weeks}　"), 12.2, True, NAVY)
    set_run(p.add_run(title), 12.2, True, TEAL)
    for label, text in [("在地課程焦點", local_focus), ("英語焦點", language_focus), ("可見成果", outcome)]:
        p = cell.add_paragraph()
        set_paragraph(p, after=1, line=1.15)
        set_run(p.add_run(label + "："), 9.8, True, NAVY)
        set_run(p.add_run(text), 9.8)
    return table


def add_steps_table(doc, steps):
    rows = []
    for stage, mins, teacher, students, evidence in steps:
        rows.append([stage, mins, teacher, students, evidence])
    return add_grid_table(doc, ["流程", "時間", "教師活動", "學生任務", "形成性評量／成果"], rows,
                          (2.0, 1.2, 5.4, 5.4, 3.0), 8.6)


def add_worksheet_plan(doc, title, focus, visual, student_work, teacher_prep, output):
    add_heading(doc, title, 2)
    return add_key_value_table(doc, [
        ("學習單名稱", focus),
        ("圖片／版面設計", visual),
        ("學生填寫內容", student_work),
        ("教師準備", teacher_prep),
        ("繳交成果", output),
    ], widths=(3.2, 13.9))


def add_worksheet_header(doc, title, subtitle):
    add_text(doc, title, 16, True, NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    add_text(doc, subtitle, 10, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, text in enumerate(["班級：________", "座號：________", "姓名：________"]):
        cell = table.cell(0, i)
        set_cell_shading(cell, PALE_BLUE)
        set_cell_border(cell)
        set_cell_margins(cell, 70, 95, 70, 95)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0)
        set_run(p.add_run(text), 9.5, True, NAVY)


def add_image_placeholder(doc, title, note, height_lines=4):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, GREY)
    set_cell_border(cell, "9AAEBB", "10")
    set_cell_margins(cell, 120, 140, 120, 140)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=2)
    set_run(p.add_run("【圖像區】" + title), 10.5, True, NAVY)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=0, line=1.1)
    set_run(p.add_run(note), 9.0, color="60717C")
    for _ in range(height_lines):
        p = cell.add_paragraph()
        set_paragraph(p, after=0, line=0.65)
        set_run(p.add_run(" "), 6)
    return table


def add_line(doc, text=""):
    p = doc.add_paragraph()
    set_paragraph(p, after=3)
    set_run(p.add_run(text + "＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿"), 10)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.65)
    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "List Bullet 2", "List Number"]:
        style = styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(10.5)
        set_east_asia(style._element, FONT)
    # Header and footer.
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(p, after=0)
    set_run(p.add_run("二水國小｜四年級在地課程英語融入"), 8.5, True, TEAL)
    footer = section.footer
    p = footer.paragraphs[0]
    set_paragraph(p, after=0)
    set_run(p.add_run("Animals × 戀戀火車 × 二八水水公園　　"), 8, color="6B7A86")
    add_page_number(p)
    return doc


def add_cover(doc):
    for _ in range(4):
        add_text(doc, "", 4, after=0)
    stripe = doc.add_table(rows=1, cols=1)
    cell = stripe.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_border(cell, NAVY)
    cell.height = Cm(0.36)
    p = cell.paragraphs[0]
    set_paragraph(p, after=0)
    set_run(p.add_run(" "), 3)
    add_text(doc, "彰化縣二水國小四年級在地課程", 20, True, NAVY, WD_ALIGN_PARAGRAPH.CENTER, before=14, after=2)
    add_text(doc, "英語融入教學詳案", 20, True, NAVY, WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_text(doc, "Topic 2 Lesson 4: Animals", 17, True, TEAL, WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(doc, "下學期第 5–8 週｜戀戀火車 × 二八水水公園資料探究", 13, True, "A25D33", WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_callout(doc, "核心任務", "學生以 dog、cat、turtle 與 Do you have a …?／Is it a …? 完成「二八水水公園友善動物資料卡」：以教師提供的真實在地資料、正確英語及友善公共空間行動介紹二水。", PALE_AMBER)
    add_text(doc, "", 6, after=6)
    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["設計年級", "節數", "在地課程單元", "主要產出"]
    values = ["四年級", "3 節，每節 40 分鐘；另含課後檢核", "下學期第 5–8 週：戀戀火車、二八水水公園資料探究（對應校定課程「走訪二八水水公園」）", "個人學習單 3 張＋小組雙語友善動物資料卡"]
    for r in range(4):
        for c in range(2):
            cell = info.cell(r, c)
            set_cell_border(cell)
            set_cell_margins(cell, 120, 135, 120, 135)
            set_col_width(cell, 3.8 if c == 0 else 13.3)
            if c == 0:
                set_cell_shading(cell, PALE_BLUE)
            p = cell.paragraphs[0]
            set_paragraph(p, after=0)
            set_run(p.add_run(labels[r] if c == 0 else values[r]), 10.5, c == 0, NAVY if c == 0 else DARK)
    add_text(doc, "依據：使用者提供的三、四年級何嘉仁英語部定／彈性課程計畫、四年級校定課程，以及國民中小學課程綱要英語文與社會領域。", 8.8, color="60717C", align=WD_ALIGN_PARAGRAPH.CENTER, before=14, after=0)
    doc.add_page_break()


def add_intro(doc):
    add_heading(doc, "一、課程定位與設計依據", 1)
    add_key_value_table(doc, [
        ("主軸定位", "本案是「在地課程」的教學活動，不另設英語課；英語是學生完成教室內在地資料探究、訪問與成果介紹時必須使用的表達工具。"),
        ("在地課程連結", "四年級下學期第 5–8 週接續「戀戀火車」與校定課程「走訪二八水水公園」：改以教師提供的老火車照片、短片、地圖、解說牌圖像或資料卡，在教室認識二水鐵道與公共空間。"),
        ("英語範圍", "完全依四年級部定計畫 Topic 2 Lesson 4 的已列內容設計：dog、cat、turtle；Do you have a …?；Is it a …?。回應句採 Yes, I do.／No, I don’t. 與 Yes, it is.／No, it isn’t.。"),
        ("彈性課程精神", "參考四年級彈性課程「彰化全球通：從職人生活到家鄉探索」的在地扎根、探究、合作與簡易英語溝通取向；活動採觀察—提問—整理—發表循環。"),
        ("三年級資料使用", "已檢視三年級部定與彈性課程計畫，作為銜接與避免重複的參考；本詳案的正式目標與評量均以四年級第二學習階段為準。"),
    ])
    add_callout(doc, "真實性原則", "動物字卡是英語學習素材，不預設二八水水公園必定有某種動物。所有在地內容均取自教師事先篩選的照片、短片、地圖、解說牌圖像或資料卡；學生須標示資料來源，不以圖卡或想像取代資料。", PALE_AMBER)

    add_heading(doc, "二、學習目標與總結性成果", 1)
    add_grid_table(doc, ["面向", "學生能夠做到"], [
        ["英語", "辨識、說出、讀出 dog／cat／turtle；以 Do you have a …? 與 Is it a …? 進行一問一答。"],
        ["在地知識", "說出二八水水公園與二水鐵道學習主題的關聯，並以資料卡整理老火車與公共空間的線索。"],
        ["社會探究", "區分「資料告訴我／我還不知道／我猜想」，用教師提供的照片、短片、地圖與資料卡支持自己的說法。"],
        ["友善行動", "在公共空間以安全、尊重、不干擾的方式看待動物與他人攜帶的寵物，完成一項小組倡議。"],
    ], (3.0, 13.9), 9.5)
    add_text(doc, "總結性成果：每組完成一張「二八水水公園友善動物資料卡」，卡上至少有 1 項可標示來源的在地資料、3 個動物詞彙辨識、1 組英語問答及 1 條中文友善行動。", 10.2, True, TEAL, before=6, after=5)

    add_heading(doc, "三、三節課進程總覽", 1)
    add_grid_table(doc, ["週次／課次", "課名", "在地任務", "英語任務", "階段成果"], [
        ["第 5 週／第 1 節", "我的動物朋友搭上二水小火車", "以「戀戀火車」導入鐵道與二水生活的連結，建立公共空間觀察規則。", "辨識 dog、cat、turtle；問答 Do you have a …?。", "動物字卡小車票＋同儕訪問記錄。"],
        ["第 7 週／第 2 節", "認識二八水水公園：資料探究", "在教室閱讀老火車、公園空間的照片、短片與資料卡，練習標示資料來源。", "用 Is it a …? 辨認教師圖卡；用 Do you have a …? 進行同儕訪問。", "公園資料卡＋一項照片／短片／資料卡來源。"],
        ["第 8 週／第 3 節", "二八水水公園友善動物資料卡", "整理資料，設計對使用公共空間者有用的友善行動與地方介紹。", "完成並口頭呈現 1 組動物英語問答。", "小組雙語資料卡與 30 秒分享。"],
    ], (2.2, 3.8, 4.9, 4.9, 3.1), 8.6)
    doc.add_page_break()


def add_curriculum_mapping(doc):
    add_heading(doc, "四、領域課綱對應", 1)
    add_heading(doc, "英語文領域（第二學習階段）", 2)
    add_grid_table(doc, ["類別", "代碼與對應說明"], [
        ["核心素養", "英-E-B1：運用所學字詞及句型進行簡易日常溝通；英-E-C2：積極參與小組學習活動，培養合作精神。"],
        ["學習表現", "1-Ⅱ-7 聽懂課堂所學字詞；2-Ⅱ-3 說出課堂所學字詞；2-Ⅱ-6 以正確發音及適切語調說出簡易句型句子；3-Ⅱ-3 看懂課堂所學句子；4-Ⅱ-4 臨摹抄寫課堂所學句子；6-Ⅱ-2 積極參與課堂練習；9-Ⅱ-1 將所學字詞作簡易歸類。"],
        ["學習內容", "Ac-Ⅱ-3 第二學習階段所學字詞；B-Ⅱ-1 第二學習階段所學字詞及句型的生活溝通；D-Ⅱ-1 所學字詞的簡易歸類。"],
        ["本案具體證據", "學生可辨認並說出 dog、cat、turtle；可完成 Do you have a …?／Is it a …? 一問一答；可在資料卡上抄寫 1 組句型，並與同伴合作分類圖卡。"],
    ], (3.1, 13.8), 9.2)

    add_heading(doc, "社會領域（第二學習階段）", 2)
    add_grid_table(doc, ["類別", "代碼與對應說明"], [
        ["核心素養", "社-E-A2：敏覺居住地方的人文環境變遷，關注生活問題並思考解決方法；社-E-C2：建立良好的人際互動關係，養成尊重差異、關懷他人及團隊合作的態度。"],
        ["學習表現", "1b-Ⅱ-1 解釋社會事物與環境之間的關係；2a-Ⅱ-2 表達對居住地方社會事物與環境的關懷；3a-Ⅱ-1 透過日常觀察與省思，對社會事物與環境提出感興趣的問題；3b-Ⅱ-2 摘取相關資料中的重點；3c-Ⅱ-2 透過同儕合作進行體驗、探究與實作；3d-Ⅱ-3 將問題解決的過程與結果，進行報告分享或實作展演。"],
        ["學習內容", "Ab-Ⅱ-1 居民的生活方式與空間利用，和其居住地方的自然、人文環境相互影響；Ca-Ⅱ-1 居住地方的環境隨著社會與經濟發展而改變；Cb-Ⅱ-1 居住地方的歷史人物、事件與文物古蹟，反映當地的歷史變遷。"],
        ["本案具體證據", "以教師提供的老火車、公園空間照片、短片、地圖與解說牌圖像為資料，區分資料事實與猜想；提出「如何友善使用公園」建議，並以小組資料卡公開分享。"],
    ], (3.1, 13.8), 9.2)

    add_heading(doc, "生活課程融入說明", 2)
    add_callout(doc, "正式課綱採用原則", "生活課程是第一學習階段（一、二年級）課程；四年級屬第二學習階段，因此本詳案不以生活課程的學習表現或學習內容代碼作為正式對應。為回應跨域需求，仍融入其探索、感受、創作與實踐精神：觀看在地影像、比較資料、繪製、材料整理、友善行動與分享。", PALE_AMBER)
    doc.add_page_break()


def add_lesson_one(doc):
    add_heading(doc, "五、逐節教學活動與檢核", 1)
    add_lesson_overview(doc, 1, "下學期第 5 週", "我的動物朋友搭上二水小火車",
                        "接續「戀戀火車」：從二水鐵道如何連結地方生活出發，建立公共空間中尊重他人與動物的觀察規則。",
                        "dog／cat／turtle；Do you have a …?；Yes, I do.／No, I don’t.；Is it a …?。",
                        "完成「動物字卡小車票」及 2 位同儕的簡易寵物訪問。")
    add_heading(doc, "教學流程（40 分鐘）", 2)
    add_steps_table(doc, [
        ["Warm-up", "5 分", "投影二水火車／公園照片與 dog、cat、turtle 圖卡。以五題快問快答診斷先備知識，不先宣稱公園中有這些動物。", "舉 A／B／C 卡作答；跟讀三個字詞。", "教師記錄可辨識字詞者；完成課前 5 題。"],
        ["Presentation", "10 分", "以「小車票」設計字卡：車票正面是動物圖，背面是英文。示範 Do you have a dog?—Yes, I do.／No, I don’t.；再以遮住圖片示範 Is it a cat?。", "以手勢與圖卡聽辨、跟讀、兩人輪流問答。", "能以圖卡正確說出至少 2 個字詞。"],
        ["Guided practice", "8 分", "說明在地資料探究規則：圖卡用來學英語；地方資訊必須來自教師提供的照片、短片、地圖、解說牌圖像或資料卡。", "在「資料告訴我／我還不知道」兩欄分類：老火車是本次要認識的地方物件；動物是否出現不預設，不能由圖卡推論。", "能說出「有資料再記錄」規則。"],
        ["Production", "14 分", "發下學習單一，示範訪問一位同學並記錄 Yes／No；巡迴協助發音與 a 的使用。", "完成三張動物小車票中的詞彙辨識，訪問 2 位同學：Do you have a …?。", "個人學習單一；至少一組正確問答。"],
        ["Wrap-up", "3 分", "請兩組展示一張小車票，帶全班口說一次問答；預告第 7 週將在教室閱讀二八水水公園的照片、短片與資料卡。", "舉起一張圖卡並說出字詞；口頭自評。", "Exit sentence：Do you have a …?。"],
    ])
    add_assessment_table(doc, "本節成果與評量", [
        ["英語口說", "能看圖說 dog／cat／turtle；能說出一個問句。", "三詞中正確 2 詞以上，且能完成一問或一答。"],
        ["探究態度", "不把圖卡內容誤當成公園的資料事實。", "能說出「要有教師資料才記錄」。"],
        ["合作", "輪流擔任提問者、回答者與記錄者。", "完成 2 人訪問記錄，無代答。"],
    ])
    add_worksheet_plan(doc, "學習單一規劃", "《我的動物朋友搭上二水小火車》",
                       "上方置入一張教師自行拍攝的二水火車／公園照片；中段以三張黑白線條「小車票」框呈現 dog、cat、turtle（圖像在上、英文在下）；下方為兩欄訪問紀錄。黑白列印仍可清楚使用。",
                       "圈選／抄寫三個動物詞；依圖完成 Is it a …?；訪問 2 位同學並勾選 Yes／No；寫下「有資料再記錄」的中文提醒。",
                       "二水鐵道或二八水水公園照片、短片截圖或資料卡 1 組；dog、cat、turtle 圖卡；A／B／C 作答卡。",
                       "個人學習單一，作為第 2 節資料卡的英語字詞工具。")

    warm = [
        [1, "哪一個是本課要學的動物英文？", "dog", "火車", "公園", "A"],
        [2, "教師讀出 cat，應選哪一張圖卡？", "貓的圖卡", "狗的圖卡", "烏龜的圖卡", "A"],
        [3, "哪一個問句是在問「你有一隻狗嗎？」？", "Do you have a dog?", "Is it a dog?", "Dog, cat, turtle.", "A"],
        [4, "教師指著烏龜圖卡問 Is it a turtle?，合適回答是？", "Yes, it is.", "Yes, I do.", "Do you have a turtle?", "A"],
        [5, "dog、cat、turtle 可以一起歸類為？", "animals", "火車種類", "公園設施", "A"],
    ]
    wrap = [
        [1, "哪一個字是「狗」？", "dog", "cat", "turtle", "A"],
        [2, "哪一個字是「貓」？", "turtle", "cat", "dog", "B"],
        [3, "哪一個字是「烏龜」？", "cat", "dog", "turtle", "C"],
        [4, "Do you have a cat? 的合適肯定回答是？", "Yes, I do.", "Yes, it is.", "Is it a cat?", "A"],
        [5, "Do you have a turtle? 的合適否定回答是？", "No, it isn’t.", "No, I don’t.", "It is a turtle.", "B"],
        [6, "看圖猜動物時，應使用哪一個問句？", "Is it a dog?", "Do you have a dog?", "Dog is it?", "A"],
        [7, "要把一張動物圖卡記成「二八水水公園有這種動物」，還需要什麼？", "教師提供的可靠資料", "只要圖卡漂亮", "同學猜一猜", "A"],
        [8, "在教室閱讀動物相關影像時，較合適的作法是？", "依資料判讀，不自行補故事", "把想像當事實", "只看圖卡就下結論", "A"],
        [9, "本節的在地課程焦點是？", "二水火車與公共空間", "認識外國動物園", "背誦所有動物名字", "A"],
        [10, "完成訪問時，最能展現合作的是？", "輪流問、答、記錄", "只讓一人完成", "不聽同學回答", "A"],
    ]
    add_question_table(doc, "課前 Warm-up Questions（5 題）", warm)
    add_question_table(doc, "課後 Wrap-up Questions（10 題，紙本或數位，約 8–10 分鐘）", wrap)
    doc.add_page_break()


def add_lesson_two(doc):
    add_lesson_overview(doc, 2, "下學期第 7 週", "認識二八水水公園：資料探究",
                        "不需外出；在教室閱讀教師事先準備的老火車、公園空間照片、短片、地圖與解說牌圖像，整理可追溯的在地資料。",
                        "複習 dog／cat／turtle；Is it a …?；Do you have a …?。",
                        "完成「公園資料卡」：一項老火車／空間資料、資料來源標記及一組英語圖卡問答。")
    add_callout(doc, "教室資料使用提醒", "教師先選用可確認來源的二八水水公園照片、短片、地圖或解說牌圖像；學生不需外出，也不從圖卡推論公園必有某種動物。每張資料卡應標示來源、拍攝者或網頁名稱與使用日期。", PALE_AMBER)
    add_heading(doc, "教學流程（40 分鐘；全程在教室內進行）", 2)
    add_steps_table(doc, [
        ["Warm-up", "5 分", "回顧第 1 節字卡，出示二八水水公園老火車照片，說明今日任務是從教師資料中摘取可證明的在地訊息。", "完成五題快問快答；說出一種可用的資料來源。", "課前 5 題；資料來源口說。"],
        ["Presentation", "8 分", "以「資料告訴我／我還不知道／我猜想」三色標籤示範記錄；用 dog、cat、turtle 圖卡帶 Is it a …? 猜圖遊戲。", "辨認圖卡並用 Yes, it is.／No, it isn’t. 回答。", "能區分英語圖卡練習與在地資料事實。"],
        ["資料探究站", "15 分", "設置照片、短片截圖、地圖、解說牌圖像四個資料站。每組只選 1 項可確認的老火車或公園空間資訊，在資料卡寫下資料來源。", "分工：資料閱讀員、記錄員、英語圖卡員、時間員；完成資料卡。", "一項附資料來源的在地資料。"],
        ["Production", "9 分", "安排「圖卡問答站」：一位學生抽圖，一位問 Is it a …?；再以 Do you have a …? 訪問組員。", "完成一組英語問答，將正確字詞貼／寫在資料卡英語區。", "每組至少一組可聽見的英語問答。"],
        ["Wrap-up", "3 分", "各組以一句中文說明「資料告訴我們……，來源是……」，並以一句英語做圖卡問答。", "分享並互評是否「有資料再記錄」。", "資料卡初稿；同儕回饋 1 則。"],
    ])
    add_assessment_table(doc, "本節成果與評量", [
        ["在地資料", "資料卡有老火車／公園空間的一項具體記錄。", "有教師提供的照片、短片、地圖或解說牌圖像來源標記，不以臆測充當事實。"],
        ["英語溝通", "依圖使用 Is it a …? 並回答；可用 Do you have a …? 問組員。", "問答句型有 1 組正確且可理解。"],
        ["公共空間行動", "能根據資料提出安全與友善規則。", "能提出不進軌道、不攀爬、不干擾或餵食動物等合理建議。"],
    ])
    add_worksheet_plan(doc, "學習單二規劃", "《認識二八水水公園：資料探究》",
                       "上半部預留「老火車／公園空間」大圖框，教師放入已確認來源的照片、短片截圖、地圖或解說牌圖像；中間有三色資料來源標籤；右下是 dog、cat、turtle 小圖卡辨認區。圖像全部以粗黑線框與大字設計，方便黑白影印。",
                       "寫下教師資料中指出的物件、資料來源與一個問題；在英語區完成 Is it a …? 問答；勾選「資料提到／資料未提到」動物資訊。",
                       "教師篩選的老火車／公園照片、短片截圖、地圖或解說牌圖像；三張動物圖卡。",
                       "個人資料卡一張，作為第 3 節友善動物資料卡的證據。")
    warm = [
        [1, "本節要用哪一種方式認識二八水水公園？", "在教室閱讀教師資料", "全班外出走訪", "只靠想像畫圖", "A"],
        [2, "在地課程中，閱讀老火車資料時最重要的是？", "標示資料來源再記錄", "先寫想像故事", "把圖卡當公園事實", "A"],
        [3, "教師指著狗圖卡問 Is it a dog?，合適回答是？", "Yes, it is.", "Yes, I do.", "Do you have a dog?", "A"],
        [4, "若教師資料沒有提到動物，資料卡最誠實的寫法是？", "資料未提到", "一定有 dog", "一定有 turtle", "A"],
        [5, "在教室進行資料探究時哪一項作法正確？", "閱讀資料後寫下來源", "把想像當答案", "任意改變資料內容", "A"],
    ]
    wrap = [
        [1, "二八水水公園在本案中的主要資料探究主題是？", "老火車與公園空間", "大型水族箱", "機場跑道", "A"],
        [2, "哪一項可以當作資料卡的資料來源？", "教師提供的照片或短片截圖", "沒有根據的猜測", "同學隨口說", "A"],
        [3, "Is it a cat? 是在做什麼？", "依圖猜動物", "詢問是否擁有寵物", "介紹火車種類", "A"],
        [4, "Do you have a turtle? 的合適回答是？", "Yes, I do.／No, I don’t.", "Yes, it is.／No, it isn’t.", "Is it a turtle?", "A"],
        [5, "從教師資料找到一項公園物件後，下一步較合適的是？", "標記來源並簡要記錄", "把它改成別的物件", "不需要留下證據", "A"],
        [6, "遇到不認識的動物時，最友善的作法是？", "安靜、保持距離", "追趕牠", "餵牠零食", "A"],
        [7, "dog、cat、turtle 圖卡在第 2 節的主要功能是？", "練習英語辨認與問答", "證明公園一定有牠們", "取代教師資料", "A"],
        [8, "哪一組最符合「有資料再記錄」？", "有來源標記的老火車資料卡", "想像的烏龜紀錄", "沒有來源的動物清單", "A"],
        [9, "小組分工的好處是？", "共同完成更完整的記錄", "讓一人做全部", "不用溝通", "A"],
        [10, "本節完成的主要成果是？", "公園資料卡初稿", "英文課本考卷", "火車模型", "A"],
    ]
    add_question_table(doc, "課前 Warm-up Questions（5 題）", warm)
    add_question_table(doc, "課後 Wrap-up Questions（10 題，紙本或數位，約 8–10 分鐘）", wrap)
    doc.add_page_break()


def add_lesson_three(doc):
    add_lesson_overview(doc, 3, "下學期第 8 週", "二八水水公園友善動物資料卡",
                        "整理第 1、2 節的教師資料探究結果，設計對使用公共空間者有用、對動物與他人友善的地方行動卡。",
                        "dog／cat／turtle；Do you have a …?；Is it a …?；Yes／No 回應。",
                        "每組發表一張雙語友善動物資料卡，完成 30 秒地方介紹與英語問答。")
    add_heading(doc, "教學流程（40 分鐘）", 2)
    add_steps_table(doc, [
        ["Warm-up", "5 分", "以第 2 節資料卡的三種資料來源圖示快速複習，進行五題快問快答。", "檢查個人資料卡；以 A／B／C 作答。", "課前 5 題；帶齊資料卡。"],
        ["Presentation", "8 分", "展示完成範例：左側為在地資料與來源，中間為 dog／cat／turtle 英語圖卡，右側為中文友善行動。示範 30 秒分享架構。", "聽辨範例中 Is it a …?／Do you have a …?，找出各區塊功能。", "能說出「資料、英語、行動」三區用途。"],
        ["Planning", "8 分", "引導各組從所有資料卡中挑一項有來源的地方內容；提醒不把英語圖卡或資料未提到的動物寫成事實。", "分工選資料、校對字詞、完成圖像與行動。", "草稿有資料來源與一項中文友善行動。"],
        ["Production", "14 分", "提供圖卡、色紙與學習單三，巡迴協助學生以正確大小寫抄寫英文，並練習問答語調。", "完成資料卡；練習：A: Is it a cat? B: Yes, it is.／A: Do you have a dog? B: No, I don’t.。", "一張小組資料卡；每人有口說角色。"],
        ["Wrap-up", "5 分", "安排 Gallery Walk：每組 30 秒發表；同儕用「有證據／英語正確／行動友善」三點勾選回饋。", "發表、聆聽、給一組正向回饋。", "完成同儕回饋與課後 10 題檢核。"],
    ])
    add_assessment_table(doc, "本節成果與評量", [
        ["在地內容", "資料卡具 1 項二水公園／鐵道的在地資料與來源。", "資料具體、可追溯，且未將資料未提到的動物寫成事實。"],
        ["英語", "正確使用至少 2 個動物詞與 1 組問答。", "字詞可辨認、句型可理解；小組至少一組問答正確。"],
        ["創作與行動", "有圖像區、英語區、中文友善行動區。", "版面清楚；行動具體且不干擾動物與他人。"],
        ["分享合作", "每人有發表或支援角色，能回應同儕。", "完成 30 秒發表與 1 則同儕回饋。"],
    ])
    add_worksheet_plan(doc, "學習單三規劃", "《二八水水公園友善動物資料卡》",
                       "採 A4 橫式三欄：左欄「資料告訴我」放老火車或公園空間的教師照片、短片截圖、地圖或解說牌圖像；中欄「Animals English」放三個圓角圖卡框與對話泡泡；右欄「友善行動」以愛心與腳印圖示框呈現。色彩建議深藍＋湖水綠＋橘黃，黑白印刷仍保留粗框與圖示。",
                       "填寫一項在地資料與來源；圈選／寫下 dog、cat、turtle；完成一組英語問答；寫下一項中文友善行動；完成口說分工。",
                       "第 1、2 節學習單；色紙或厚紙板；動物圖卡；公園／老火車照片（可選）；彩色筆。",
                       "小組成品一張，張貼或拍照留存；每人保留學習單三草稿作個人評量。")
    warm = [
        [1, "製作資料卡前，應先選哪一種內容？", "有資料來源的在地資料", "沒有證據的想像", "網路上不明的圖片", "A"],
        [2, "資料卡的英語區應使用本課已學的哪些字？", "dog、cat、turtle", "所有動物名字", "火車零件名字", "A"],
        [3, "要詢問同學是否有一隻貓，應說？", "Do you have a cat?", "Is it a cat?", "Yes, it is.", "A"],
        [4, "要依圖片猜「這是一隻狗嗎？」應說？", "Is it a dog?", "Do you have a dog?", "No, I don’t.", "A"],
        [5, "友善動物資料卡中，哪一項最重要？", "不打擾並尊重公共空間", "追趕動物拍照", "把資料未提到的動物寫上去", "A"],
    ]
    wrap = [
        [1, "一張完整的友善動物資料卡應有？", "在地資料、英語、友善行動", "只有英語單字", "只有圖片", "A"],
        [2, "若資料顯示的是老火車，正確的表達是？", "把老火車記為在地資料", "把老火車寫成 turtle", "不需要資料來源", "A"],
        [3, "哪一句是在問是否擁有寵物？", "Do you have a dog?", "Is it a dog?", "Yes, it is.", "A"],
        [4, "哪一句是在依圖辨認動物？", "Is it a turtle?", "Do you have a turtle?", "No, I don’t.", "A"],
        [5, "Is it a cat? 的合適否定回答是？", "No, it isn’t.", "No, I don’t.", "Do you have a cat?", "A"],
        [6, "Do you have a cat? 的合適否定回答是？", "No, I don’t.", "No, it isn’t.", "Is it a cat?", "A"],
        [7, "小組發表時，哪一種做法最好？", "每人負責一小部分", "只讓一人說完", "不看同組資料", "A"],
        [8, "給別組回饋時，可以優先檢查？", "是否有資料來源與正確英語", "作品顏色是否最多", "誰講得最大聲", "A"],
        [9, "根據二八水水公園資料討論公共空間時，較友善的行動是？", "不攀爬展示物並保持整潔", "任意餵食動物", "跨越安全範圍", "A"],
        [10, "本單元最終成果是？", "二八水水公園友善動物資料卡與分享", "背完整本字典", "把圖卡當作公園事實", "A"],
    ]
    add_question_table(doc, "課前 Warm-up Questions（5 題）", warm)
    add_question_table(doc, "課後 Wrap-up Questions（10 題，紙本或數位，約 8–10 分鐘）", wrap)
    doc.add_page_break()


def add_worksheets(doc):
    add_heading(doc, "六、可直接列印的學生學習單", 1)
    add_text(doc, "下列三張學習單依序對應第 1–3 節。所有二水在地素材均由教師在教室內提供，可使用已確認來源的照片、短片截圖、地圖、解說牌圖像或資料卡；學生不需外出。", 10.2, color="60717C", after=5)
    doc.add_page_break()
    if WORKSHEET_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(WORKSHEET_IMAGE), width=Cm(16.0))
    else:
        add_callout(doc, "圖像版學習單一", "請先執行 compose_lesson1_worksheet_png.py 產生圖像版學習單。", PALE_AMBER)
    doc.add_page_break()

    add_worksheet_header(doc, "學習單一｜我的動物朋友搭上二水小火車", "Animals × 戀戀火車｜第 1 節")
    add_image_placeholder(doc, "二水火車／公園照片或學生素描", "教師放入自己拍攝的在地照片；請勿使用未授權網路圖片。學生也可畫一節車廂和窗外景色。", 3)
    add_heading(doc, "A. 我的動物小車票（看圖／教師讀字，圈出正確英文）", 3)
    t = doc.add_table(rows=2, cols=3)
    t.style = "Table Grid"
    animals = [("狗的圖卡", "dog"), ("貓的圖卡", "cat"), ("烏龜的圖卡", "turtle")]
    for i, (image, word) in enumerate(animals):
        cell = t.cell(0, i)
        set_cell_shading(cell, GREY); set_cell_border(cell); set_cell_margins(cell, 100, 95, 100, 95)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph(p, after=1)
        set_run(p.add_run("【圖】" + image), 10, True, NAVY)
        p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph(p, after=0)
        set_run(p.add_run("（教師貼圖／學生描圖）"), 8.5, color="60717C")
        cell = t.cell(1, i)
        set_cell_border(cell); set_cell_margins(cell, 95, 95, 95, 95)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph(p, after=0)
        set_run(p.add_run("英文：________________\n提示：" + word), 10, True, TEAL)
    add_heading(doc, "B. 問問同學：Do you have a …?", 3)
    add_grid_table(doc, ["我問誰？", "我的問題（圈一個）", "同學回答（勾選）", "我記錄的動物"], [
        ["同學 1：________", "dog / cat / turtle", "□ Yes, I do.　□ No, I don’t.", "________"],
        ["同學 2：________", "dog / cat / turtle", "□ Yes, I do.　□ No, I don’t.", "________"],
    ], (3.1, 4.0, 6.0, 3.7), 9.0)
    add_heading(doc, "C. 我知道：有資料再記錄", 3)
    add_text(doc, "我會先閱讀老師提供的照片、短片或資料卡，再寫下二水公園的事情。　□ 我做得到", 10.2, True, NAVY, after=2)
    doc.add_page_break()

    add_worksheet_header(doc, "學習單二｜認識二八水水公園：資料探究", "Animals × 二八水水公園｜第 2 節")
    add_image_placeholder(doc, "老火車／公園空間資料圖", "請閱讀教師提供的照片、短片截圖、地圖或解說牌圖像，並在框內圈選、描繪或貼上其中一項資料。請加上資料來源。", 5)
    add_heading(doc, "A. 我的在地資料", 3)
    add_grid_table(doc, ["我記錄的物件／地方", "我從哪裡知道？", "我的一個問題"], [
        ["________________________", "□ 教師照片　□ 短片截圖　□ 地圖／解說牌圖像　□ 資料卡", "________________________"],
    ], (5.0, 6.6, 5.2), 9.0)
    add_heading(doc, "B. Animals English 圖卡站", 3)
    add_text(doc, "同伴抽一張圖卡，我問：Is it a __________________?　同伴答：□ Yes, it is.　□ No, it isn’t.", 10.2, True, NAVY, after=3)
    add_text(doc, "教師資料是否提到動物？　□ 有（寫下資料來源：____________）　□ 資料未提到", 10.2, True, TEAL, after=3)
    add_heading(doc, "C. 公園友善行動", 3)
    add_grid_table(doc, ["我會做到（勾選）", "我想提醒大家的一句中文話"], [
        ["□ 不進軌道　□ 不攀爬展示物　□ 不餵食、不追逐動物　□ 保持整潔", "______________________________________"],
    ], (8.8, 8.0), 9.0)
    doc.add_page_break()

    add_worksheet_header(doc, "學習單三｜二八水水公園友善動物資料卡", "Animals × 在地課程成果發表｜第 3 節")
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    labels = [
        ("① 在地資料探究", "請畫出／貼上教師提供的老火車或公園空間資料。\n資料來源：________________\n\n\n\n\n\n"),
        ("② Animals English", "我會說：□ dog　□ cat　□ turtle\n\nA: Is it a __________?\nB: □ Yes, it is.　□ No, it isn’t.\n\nA: Do you have a __________?\nB: □ Yes, I do.　□ No, I don’t."),
        ("③ 友善行動", "我們的提醒：\n________________________\n\n□ 不干擾動物\n□ 尊重他人\n□ 愛護公園\n\n我的角色：____________"),
    ]
    fills = [PALE_BLUE, PALE_TEAL, PALE_AMBER]
    for i, (title, body) in enumerate(labels):
        cell = t.cell(0, i)
        set_cell_shading(cell, fills[i]); set_cell_border(cell, NAVY if i == 0 else TEAL); set_cell_margins(cell, 120, 115, 120, 115)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph(p, after=5)
        set_run(p.add_run(title), 11.2, True, NAVY)
        p = cell.add_paragraph(); set_paragraph(p, after=0, line=1.25)
        set_run(p.add_run(body), 9.7, color=DARK)
    add_text(doc, "發表前自我檢查：□ 有一項在地資料　□ 有資料來源　□ 有至少 2 個動物詞　□ 有一組問答　□ 有友善行動", 10.2, True, NAVY, before=8, after=4)
    add_text(doc, "同儕回饋：我喜歡第 ______ 組，因為他們的資料／英語／友善行動是：____________________________________________", 10, color="43515B", after=0)
    doc.add_page_break()


def add_rubric_sources(doc):
    add_heading(doc, "七、教師評量規準與備課提醒", 1)
    add_grid_table(doc, ["評量面向", "4：表現穩定", "3：達成", "2：需提示", "1：需支持"], [
        ["英語字詞與句型", "三詞正確，能主動完成 2 種問答。", "三詞多數正確，能完成 1 組問答。", "需圖卡或同伴提示才可完成。", "尚無法辨認多數字詞／句型。"],
        ["在地資料品質", "記錄具體、有來源，能分辨事實與未知。", "有一項可追溯的記錄。", "有記錄但來源不完整。", "以想像取代觀察，需重新引導。"],
        ["友善公共空間", "能提出具體友善行動並身體力行。", "能遵守並說出一項規則。", "需教師提醒才遵守。", "尚未能遵守安全／友善規則。"],
        ["合作與分享", "分工明確，能傾聽並建設性回饋。", "能完成分工與分享。", "參與不穩定，需提醒。", "難以參與，需個別支持。"],
    ], (3.0, 3.9, 3.9, 3.9, 3.9), 8.5)
    add_heading(doc, "差異化支持", 2)
    add_bullet(doc, "需要支持者：保留 dog／cat／turtle 三張大圖卡與句型條；採指認、跟讀、二選一回答，再逐步過渡到三選一。")
    add_bullet(doc, "進階者：擔任「資料來源檢核員」或「英語主持人」，帶領小組練習兩種問答，並說明為何某項內容屬於「資料未提到」。")
    add_bullet(doc, "教室資料探究：使用教師自攝二八水水公園照片、官方介紹頁截圖、短片截圖、地圖或解說牌圖像；每一份資料均標示來源與使用日期。")
    add_heading(doc, "備課檢核", 2)
    add_bullet(doc, "確認照片、素描與資料來源只用於教學且符合肖像、授權與校內規範；優先使用教師自攝圖。")
    add_bullet(doc, "課前確認數位設備、投影、資料卡列印與資料來源標示；本詳案全程可在一般教室內完成。")
    add_bullet(doc, "英語目標保持聚焦：不因在地內容增加超出本課範圍的英文動物詞；地方知識可用中文深度討論。")

    add_heading(doc, "八、參考與使用資料", 1)
    add_text(doc, "一、使用者提供課程文件（本案已檢視）", 10.2, True, NAVY, after=2)
    for item in [
        "何嘉仁四年級英語部定課程計畫.docx：Topic 2 Lesson 4: Animals，列示 dog、cat、turtle、Do you have a …?、Is it a …? 與 B-Ⅱ-1／D-Ⅱ-1。",
        "何嘉仁四年級彈性課程計畫.docx：彰化全球通：從職人生活到家鄉探索，提供在地扎根、探究、合作與英語溝通的設計取向。",
        "何嘉仁三年級英語部定課程計畫.docx、何嘉仁三年級彈性課程計畫.docx：作為學習銜接與避免重複之參考。",
        "4年級 校定課程.pdf：下學期「戀戀火車」與「走訪二八水水公園」的教學脈絡。",
        "語文領域-英語文.pdf、課程綱要國民中小學暨普通型高級中等學校－社會領域.pdf：第二學習階段學習表現與學習內容。",
    ]:
        add_bullet(doc, item)
    add_text(doc, "二、在地資料（教師備課可查核）", 10.2, True, NAVY, before=4, after=2)
    for item in [
        "彰化縣政府城市暨觀光發展處：二八水水公園。https://tourism.chcg.gov.tw/AttractionsContent.aspx?chk=4df749d6-c9d8-423f-b867-b6239f6e1d09&id=392&l=TW",
        "教育部：十二年國民基本教育課程綱要。https://www.naer.edu.tw/upload/1/19/doc/192475292.pdf",
        "國家教育研究院：生活課程課程綱要（用於確認生活課程適用第一學習階段）。https://www.naer.edu.tw/upload/1/16/doc/813/%28%E7%99%BC%E5%B8%83%E7%89%88%29%E7%94%9F%E6%B4%BB%E8%AA%B2%E7%A8%8B%E8%AA%B2%E7%A8%8B%E7%B6%B1%E8%A6%81.pdf",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "版本註記", "本詳案以使用者提供之四年級課程計畫為準。若實際教材圖卡、單字表或校定課程週次後續調整，請優先依學校正式版本更新圖卡與進度；教室資料探究原則及第二學習階段課綱對應可維持使用。", PALE_BLUE)


def main():
    doc = setup_document()
    add_cover(doc)
    add_intro(doc)
    add_curriculum_mapping(doc)
    add_lesson_one(doc)
    add_lesson_two(doc)
    add_lesson_three(doc)
    add_worksheets(doc)
    add_rubric_sources(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
