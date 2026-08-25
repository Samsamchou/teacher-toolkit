from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "_course_source_md"
OUT.mkdir(exist_ok=True)


def clean(value: str) -> str:
    return " ".join((value or "").replace("\u3000", " ").split())


def docx_to_md(path: Path) -> Path:
    doc = Document(path)
    lines = [f"# {path.stem}", ""]
    paragraph_count = 0
    for para in doc.paragraphs:
        text = clean(para.text)
        if not text:
            continue
        paragraph_count += 1
        style = para.style.name if para.style else ""
        if style.startswith("Heading"):
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 2
            lines.extend([f"{'#' * max(1, min(level, 6))} {text}", ""])
        elif para.style and "List" in style:
            lines.append(f"- {text}")
        else:
            lines.extend([text, ""])

    for table_no, table in enumerate(doc.tables, 1):
        lines.extend([f"## 表格 {table_no}", ""])
        for row_no, row in enumerate(table.rows, 1):
            cells = [clean(cell.text) for cell in row.cells]
            lines.append(f"- 第 {row_no} 列： " + " | ".join(cells))
        lines.append("")

    out_path = OUT / f"{path.stem}.md"
    out_path.write_text("\n".join(lines), encoding="utf-8")
    print(
        f"DOCX\t{path.name}\tparagraphs={paragraph_count}\t"
        f"tables={len(doc.tables)}\toutput={out_path.name}"
    )
    return out_path


def docx_tables_to_tsv(path: Path) -> Path:
    doc = Document(path)
    lines = []
    for table_no, table in enumerate(doc.tables, 1):
        lines.append(f"### TABLE {table_no} {len(table.rows)}x{len(table.columns)}")
        for row_no, row in enumerate(table.rows, 1):
            values = []
            for cell in row.cells:
                value = clean(cell.text)
                if not values or value != values[-1]:
                    values.append(value)
            lines.append(f"{row_no}\t" + "\t".join(values))
        lines.append("")
    out_path = OUT / f"{path.stem}.tables.tsv"
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return out_path


def pdf_to_md(path: Path) -> Path:
    pdf = PdfReader(path)
    lines = [f"# {path.stem}", ""]
    nonempty_pages = 0
    for page_no, page in enumerate(pdf.pages, 1):
        text = (page.extract_text() or "").strip()
        lines.extend([f"## 第 {page_no} 頁", ""])
        if text:
            nonempty_pages += 1
            lines.extend([text, ""])
        else:
            lines.extend(["[本頁未擷取到文字]", ""])
    out_path = OUT / f"{path.stem}.md"
    out_path.write_text("\n".join(lines), encoding="utf-8")
    print(
        f"PDF\t{path.name}\tpages={len(pdf.pages)}\tnonempty_pages={nonempty_pages}\t"
        f"output={out_path.name}"
    )
    return out_path


for source in sorted(ROOT.iterdir()):
    if source.suffix.lower() == ".docx" and (
        "課程計畫" in source.name or source.name.startswith("何嘉仁")
    ):
        docx_to_md(source)
        docx_tables_to_tsv(source)
    elif source.suffix.lower() == ".pdf":
        pdf_to_md(source)
