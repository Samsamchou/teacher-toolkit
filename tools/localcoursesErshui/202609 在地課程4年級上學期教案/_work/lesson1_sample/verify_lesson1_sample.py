from pathlib import Path
import hashlib
import json
import re

from docx import Document
from docx.oxml.ns import qn


SCRIPT_DIR = Path(__file__).resolve().parent
OUT_DIR = SCRIPT_DIR.parents[1]
ROOT = SCRIPT_DIR.parents[2]
FINAL = OUT_DIR / "04_扇形車庫_第1節四階段詳案樣稿_清楚活動版_20260825.docx"
REFERENCE = ROOT / "在地課程4年級上學期教案" / "01_第1-2週_扇形車庫.docx"
EXPECTED_REFERENCE_SHA = "3779FB5E553287DF55027EBD64A26DB1F6809997520CE639373FF043FA0EB759"


def sha256(path):
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        parts.extend(p.text for p in section.header.paragraphs)
        parts.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(parts)


def table_geometry(table):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    width = int(tbl_w.get(qn("w:w"))) if tbl_w is not None else None
    grid = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
    return {"width": width, "grid_sum": sum(grid), "columns": len(grid)}


def main():
    assert FINAL.exists(), FINAL
    doc = Document(FINAL)
    text = all_text(doc)
    required = [
        "Warm-up",
        "Presentation",
        "Production",
        "Wrap-up",
        "交通、安全、時間及經費因素",
        "What are these/those?",
        "原始來源—設計決定—教學活動—評量證據",
        "不評鐵道專有詞拼寫",
        "看四張局部圖，說出你看見的線索",
        "看影片認識扇形車庫的四個主要部分",
        "兩人把四張名稱卡放到正確位置",
        "看扇形車庫全景圖，指出三個主要構件",
        "每2人1張A3扇形車庫全景圖",
        "每人1張A5無文字全景圖",
    ]
    for token in required:
        assert token in text, token
    assert "理由待補" not in text
    assert "原因尚待" not in text
    for stale_title in [
        "停—看—指—說：影片證據示範",
        "四人構件偵探：配置、證據、輪換",
        "無標示全景出口票",
        "學習單縮圖記錄停格編號",
    ]:
        assert stale_title not in text, stale_title

    flow_table = None
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == "階段／時間":
            flow_table = table
            break
    assert flow_table is not None
    assert len(flow_table.rows) == 5
    flow_font_sizes = []
    for row in flow_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        assert run.font.size is not None, run.text
                        flow_font_sizes.append(round(run.font.size.pt, 1))
    assert flow_font_sizes
    assert set(flow_font_sizes) == {12.0}, sorted(set(flow_font_sizes))
    minutes = []
    for row in flow_table.rows[1:]:
        match = re.search(r"(\d+)分鐘", row.cells[0].text)
        assert match, row.cells[0].text
        minutes.append(int(match.group(1)))
    assert minutes == [5, 10, 20, 5]
    assert sum(minutes) == 40

    section = doc.sections[0]
    geometry = {
        "page_width_in": round(section.page_width.inches, 2),
        "page_height_in": round(section.page_height.inches, 2),
        "left_margin_in": round(section.left_margin.inches, 2),
        "right_margin_in": round(section.right_margin.inches, 2),
        "top_margin_in": round(section.top_margin.inches, 2),
        "bottom_margin_in": round(section.bottom_margin.inches, 2),
    }
    assert geometry == {
        "page_width_in": 8.5,
        "page_height_in": 11.0,
        "left_margin_in": 1.0,
        "right_margin_in": 1.0,
        "top_margin_in": 1.0,
        "bottom_margin_in": 1.0,
    }

    table_checks = [table_geometry(table) for table in doc.tables]
    assert all(item["width"] == 9360 for item in table_checks)
    assert all(item["grid_sum"] == 9360 for item in table_checks)

    reference_sha = sha256(REFERENCE)
    assert reference_sha == EXPECTED_REFERENCE_SHA
    result = {
        "final": str(FINAL),
        "bytes": FINAL.stat().st_size,
        "sha256": sha256(FINAL),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "flow_minutes": minutes,
        "flow_total": sum(minutes),
        "flow_font_sizes_pt": sorted(set(flow_font_sizes)),
        "page_geometry": geometry,
        "table_geometry": table_checks,
        "reference_sha256": reference_sha,
        "required_tokens": "PASS",
        "stale_pending_wording": "PASS",
        "clear_activity_titles": "PASS",
    }
    print(json.dumps(result, ensure_ascii=True, indent=2))


if __name__ == "__main__":
    main()
