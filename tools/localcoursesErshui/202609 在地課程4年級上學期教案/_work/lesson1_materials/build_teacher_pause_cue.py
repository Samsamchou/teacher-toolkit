from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\User\.codex\visualizations\2026\08\25\01a03887-2ca6-7233-8e1e-37a3f0b56a27\06_teacher_pause_cue_20260825.docx")

QUESTIONS = [
    {
        "number": 1,
        "time": "00:15",
        "evidence": "主持人站在轉車台旁，背景可見車庫；字幕說明火車頭需要維修或保養時會進入車庫。",
        "question": "火車頭需要維修或保養時，會開到哪裡？",
        "options": ["A. 車庫", "B. 車站月台", "C. 平交道"],
        "answer": "A. 車庫",
        "explanation": "車庫是火車頭休息、保養與維修的地方。",
    },
    {
        "number": 2,
        "time": "00:21",
        "evidence": "畫面從高處拍到圓弧車庫和扇狀軌道；字幕說明外形像一把扇子。",
        "question": "外形看起來像一把扇子的火車庫，稱為什麼？",
        "options": ["A. 圓形車庫", "B. 扇形車庫", "C. 三角形車庫"],
        "answer": "B. 扇形車庫",
        "explanation": "軌道和車庫向外展開，看起來像打開的扇子。",
    },
    {
        "number": 3,
        "time": "00:29",
        "evidence": "畫面先出現轉車台機構，再拉到轉車台和車庫；字幕說明轉車台位在扇狀配置中心。",
        "question": "扇形車庫以哪一個構件為中心？",
        "options": ["A. 車庫門", "B. 轉車台", "C. 火車頭"],
        "answer": "B. 轉車台",
        "explanation": "轉車台像扇子的扇軸，軌道從它向各車庫門展開。",
    },
    {
        "number": 4,
        "time": "00:35",
        "evidence": "畫面顯示一格一格的車庫門及停放其中的火車頭；字幕說明可停放12個火車頭。",
        "question": "扇形車庫最多可以停放幾個火車頭？",
        "options": ["A. 11個", "B. 10個", "C. 12個"],
        "answer": "C. 12個",
        "explanation": "12股道通往12個車庫位置，因此可停放12個火車頭。",
    },
    {
        "number": 5,
        "time": "00:59",
        "evidence": "畫面顯示火車頭停在轉車台上，轉車台轉向不同股道；字幕說明可依維修安排轉入第1至第12股。",
        "question": "轉車台的主要工作是什麼？",
        "options": ["A. 把火車頭轉到需要進入的股道", "B. 載旅客到車站月台", "C. 清洗火車頭外殼"],
        "answer": "A. 把火車頭轉到需要進入的股道",
        "explanation": "轉車台先轉動火車頭，再對準指定軌道，讓火車頭進入相應車庫。",
    },
]


# compact_reference_guide preset, with named overrides for readable teacher use:
# - body 12 pt and line spacing 1.15
# - dark teal accent palette
# Page geometry remains Letter portrait with 1-inch margins and 9360-DXA tables.
FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft JhengHei"
NAVY = "17365D"
TEAL = "0F6B6D"
LIGHT_TEAL = "E7F3F2"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "5A6573"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C7D0D9", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    assert sum(widths) == TABLE_WIDTH_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_run_font(run, size=12, bold=False, color="000000", italic=False):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before=0, after=4, line=1.15, keep=False, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_together = keep
    if align is not None:
        paragraph.alignment = align


def add_text(paragraph, text, size=12, bold=False, color="000000", italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return run


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MID_GRAY)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def add_question_card(doc, item):
    table = doc.add_table(rows=4, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table)
    mark_header_row(table.rows[0])
    for row in table.rows:
        set_row_cant_split(row)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    header = table.cell(0, 0)
    set_cell_shading(header, TEAL)
    header.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = header.paragraphs[0]
    clear_paragraph(p)
    style_paragraph(p, after=0, keep=True)
    add_text(p, f"題目 {item['number']}　暫停時間 {item['time']}", size=13, bold=True, color=WHITE)

    evidence_cell = table.cell(1, 0)
    set_cell_shading(evidence_cell, LIGHT_GRAY)
    p = evidence_cell.paragraphs[0]
    clear_paragraph(p)
    style_paragraph(p, after=0, keep=True)
    add_text(p, "畫面／旁白依據：", bold=True, color=NAVY)
    add_text(p, item["evidence"])

    question_cell = table.cell(2, 0)
    p = question_cell.paragraphs[0]
    clear_paragraph(p)
    style_paragraph(p, after=4, keep=True)
    add_text(p, "問題：", bold=True, color=NAVY)
    add_text(p, item["question"], bold=True)
    for option in item["options"]:
        op = question_cell.add_paragraph()
        style_paragraph(op, before=0, after=1, line=1.1, keep=True)
        op.paragraph_format.left_indent = Inches(0.18)
        add_text(op, option)

    answer_cell = table.cell(3, 0)
    set_cell_shading(answer_cell, LIGHT_TEAL)
    p = answer_cell.paragraphs[0]
    clear_paragraph(p)
    style_paragraph(p, after=2, keep=True)
    add_text(p, "正確答案：", bold=True, color=TEAL)
    add_text(p, item["answer"], bold=True)
    p2 = answer_cell.add_paragraph()
    style_paragraph(p2, after=0, keep=True)
    add_text(p2, "教師簡短解析：", bold=True, color=TEAL)
    add_text(p2, item["explanation"])

    # Keep each four-row question card together on one page. Every paragraph
    # except the final paragraph chains to the next paragraph/row.
    chained = []
    for row in table.rows:
        for cell in row.cells:
            chained.extend(cell.paragraphs)
    for paragraph in chained[:-1]:
        paragraph.paragraph_format.keep_with_next = True

    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=3)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(12)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    header = section.header
    hp = header.paragraphs[0]
    clear_paragraph(hp)
    style_paragraph(hp, after=0, line=1.0)
    add_text(hp, "二水國小｜四年級在地課程", size=9, bold=True, color=MID_GRAY)
    add_text(hp, "　　　　　　　　　　　　　　　　　〈扇形車庫〉第1節", size=9, color=MID_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    style_paragraph(fp, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(fp, "教師播放提示｜2026-08-25｜第 ", size=9, color=MID_GRAY)
    add_page_field(fp)
    add_text(fp, " 頁", size=9, color=MID_GRAY)

    title = doc.add_paragraph()
    style_paragraph(title, after=3)
    add_text(title, "影片播放與 Kahoot 暫停提示表", size=23, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, after=10)
    add_text(subtitle, "單元：扇形車庫｜正式5題版（原第4題已移除）", size=12, bold=True, color=TEAL)

    strip = doc.add_table(rows=1, cols=5)
    set_table_geometry(strip, [1872, 1872, 1872, 1872, 1872])
    set_table_borders(strip, color="AFCBC9", size=6)
    mark_header_row(strip.rows[0])
    for idx, item in enumerate(QUESTIONS):
        cell = strip.cell(0, idx)
        set_cell_shading(cell, LIGHT_TEAL if idx % 2 == 0 else LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        clear_paragraph(p)
        style_paragraph(p, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, f"題{item['number']}\n{item['time']}", size=11.5, bold=True, color=NAVY)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    use_box = doc.add_table(rows=1, cols=1)
    set_table_geometry(use_box, [TABLE_WIDTH_DXA])
    set_table_borders(use_box, color="B8C4D1", size=6)
    mark_header_row(use_box.rows[0])
    set_cell_shading(use_box.cell(0, 0), LIGHT_BLUE)
    p = use_box.cell(0, 0).paragraphs[0]
    clear_paragraph(p)
    style_paragraph(p, after=0)
    add_text(p, "操作順序：", bold=True, color=NAVY)
    add_text(p, "播放至指定時間並暫停 → 切換對應 Kahoot 題目 → 學生作答 → 教師用一句話解析 → 回到影片繼續播放。")

    note = doc.add_paragraph()
    style_paragraph(note, before=5, after=8)
    add_text(note, "教學提醒：", bold=True, color=TEAL)
    add_text(note, "若切換畫面耗時，可先完整播放一次，再於第二次播放時依五個時間點作答。題目不要求學生拼寫鐵道英文專有詞。")

    for item in QUESTIONS:
        add_question_card(doc, item)

    source_heading = doc.add_paragraph()
    style_paragraph(source_heading, before=6, after=2, keep=True)
    add_text(source_heading, "影片來源與用字", size=12, bold=True, color=NAVY)
    source = doc.add_paragraph()
    style_paragraph(source, after=2)
    add_text(source, "小公視｜花路米去哪裡－彰化火車扇形車庫：", bold=True)
    add_text(source, "https://www.youtube.com/watch?v=y-1dJzGhwOk")
    wording = doc.add_paragraph()
    style_paragraph(wording, after=0)
    add_text(wording, "影片畫面曾出現「扇型車庫」；本教材依官方名稱統一使用「扇形車庫」。", italic=True, color=MID_GRAY)

    doc.core_properties.title = "扇形車庫第1節｜影片Kahoot教師暫停提示表"
    doc.core_properties.subject = "四年級在地課程教學教材"
    doc.core_properties.author = "二水國小"
    doc.core_properties.keywords = "扇形車庫,Kahoot,教師提示,在地課程"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"questions={len(QUESTIONS)}")


if __name__ == "__main__":
    build_document()
