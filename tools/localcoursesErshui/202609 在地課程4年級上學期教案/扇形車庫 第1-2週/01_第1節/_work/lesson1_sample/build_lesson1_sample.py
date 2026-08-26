from pathlib import Path
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


SCRIPT_DIR = Path(__file__).resolve().parent
OUT_DIR = SCRIPT_DIR.parents[1]
ROOT = SCRIPT_DIR.parents[4]
REFERENCE = ROOT / "在地課程4年級上學期教案" / "01_第1-2週_扇形車庫.docx"
FINAL = Path(
    os.environ.get(
        "LESSON1_SAMPLE_OUTPUT",
        str(OUT_DIR / "04_扇形車庫_第1節四階段詳案樣稿_清楚活動版_20260825.docx"),
    )
)

sys.path.insert(0, str(ROOT))

from build_grade4_sem1_docs import (  # noqa: E402
    BORDER,
    HEADER_FILL,
    MUTED,
    NAVY,
    PALE_AMBER,
    PALE_BLUE,
    PALE_TEAL,
    TEAL,
    WHITE,
    add_body,
    add_grid_table,
    add_heading,
    add_kv_table,
    add_lesson_title,
    add_real_bullet,
    add_resource_table,
    add_text_cell,
    format_curriculum_entries,
    prevent_row_split,
    set_cell_shading,
    set_paragraph_format,
    set_repeat_table_header,
    set_run_font,
    set_table_geometry,
)
from grade4_sem1_lesson_data import UNITS  # noqa: E402


def clear_body(doc):
    body = doc._body._element
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)


def add_metadata_table(doc):
    rows = [
        ["適用單元", "第1–2週〈扇形車庫〉", "本檔範圍", "第1節完整樣稿"],
        ["授課時間", "40分鐘", "對象", "二水國小四年級"],
        ["骨架狀態", "2026-08-25教師已確認", "版本", "第1節清楚活動比較版"],
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for idx, value in enumerate(("項目", "內容", "項目", "內容")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        add_text_cell(cell, value, 9.2, True, WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(row):
            if idx in (0, 2):
                set_cell_shading(cells[idx], HEADER_FILL)
                add_text_cell(cells[idx], value, 9.3, True, NAVY)
            else:
                add_text_cell(cells[idx], value, 9.3)
    set_table_geometry(table, [1350, 3330, 1350, 3330])
    doc.add_paragraph()


def stage(name, minutes, title, teacher, students, assessment):
    return {
        "name": name,
        "minutes": minutes,
        "title": title,
        "teacher": teacher,
        "students": students,
        "assessment": assessment,
    }


def lesson_payload():
    return {
        "title": "第1節｜影片觀賞：找到扇形車庫的關鍵構件",
        "focus": "以原課程的影片觀賞為主軸；因交通、安全、時間及經費因素取消踏查，改從影片與經查證的圖片觀察四個主要部分。先會看、會指、會說出判斷線索，再使用英語指認。",
        "patterns": "What’s this/that? It’s a turntable／roundhouse door／locomotive.／What are these/those? They’re tracks.",
        "pattern_source": "三年級部定計畫TSV第42–43行；四年級部定計畫TSV第43–44行。本節只依畫面單複數使用，不另教完整文法。",
        "daily": "Look here.／Point to the track(s).／Your turn.／Good job.",
        "core": "turntable 轉車台、track／tracks 軌道、door／roundhouse door 車庫門、locomotive 火車頭",
        "related": "roundhouse 扇形車庫（整體名稱）。鐵道專有詞以看圖辨識和口說標籤為主，不要求拼寫；位置句型 Where’s the turntable? 留到第2節。",
        "digital": "教師指定的扇形車庫影片；4張構件局部圖；1張沒有文字標示的扇形車庫全景圖。影片題目與暫停時間須等本節樣稿確認後，另行讀取影片內容再製作。",
        "handmade": "每2人1張A3扇形車庫全景圖；轉車台turntable、放射狀軌道tracks、車庫門roundhouse door、火車頭locomotive中英名稱卡各1張；每人1張A5無文字全景圖；單數／複數口說支架卡。",
        "worksheet": "《扇形車庫四個主要部分》：以同一張無文字全景圖供學生指認；中英文名稱由教師在學習單版面疊加，不另生成有文字的圖片。",
        "stages": [
            stage(
                "Warm-up",
                5,
                "看四張局部圖，說出你看見的線索",
                "依序顯示轉車台、放射狀軌道、車庫門與火車頭的局部圖，每張保留10秒觀察時間。先問「你看見什麼形狀或線條？」讓學生說中文線索，再指著畫面示範一個英文名稱。桌前小圖使用 What’s this?，遠處投影使用 What’s that?。",
                "兩人一起看圖；A先說一個中文線索，B指出畫面中的位置，再一起跟讀英文名稱。下一張圖交換先說的人，讓兩人都有觀察和口說機會。",
                "個人能說出1項可見線索，並正確指出或口說至少1個核心詞；需要時可使用圖卡。",
            ),
            stage(
                "Presentation",
                10,
                "看影片認識扇形車庫的四個主要部分",
                "播放教師指定影片。第一次播放時請學生專心找出可能出現的四個部分；第二次播放時，教師在看得到轉車台、放射狀軌道、車庫門或火車頭的畫面暫停。每次暫停都依序做四件事：詢問畫面中看見什麼、請學生用手指出位置、把中英名稱卡貼到全景圖旁、帶全班口說。看到多條軌道時示範 What are these? They’re tracks.。提醒學生這是觀看影片，不是親自踏查。正式暫停秒數等樣稿確認後再依影片內容核對。",
                "第一次看片只觀察，不寫答案；第二次停下來時，學生先在自己的無文字全景圖上指出相同位置，再和搭檔輪流說出看到的名稱。聽到 Point to the turntable／tracks. 時，直接用手指出正確位置。",
                "教師隨機出示一個暫停畫面；學生能指出並說出對應名稱。至少能分辨一個turntable和多條tracks；中文功能說明可由搭檔補充。",
            ),
            stage(
                "Production",
                20,
                "兩人把四張名稱卡放到正確位置",
                "每2人發1張A3無文字全景圖和4張中英名稱卡。先用3分鐘示範：看名稱卡、在圖中找位置、把卡放在物件旁、用手指出並說名稱。接著給學生7分鐘共同放卡；教師巡視時只問「你在圖上看到什麼？」和「這張卡應該靠近哪裡？」。再用6分鐘進行兩人口說：A指圖提問，B回答；完成兩張卡後交換。最後4分鐘和鄰近一組比較位置，有不同時回看全景圖再調整。名稱卡只放在圖上，不黏死，方便修正。",
                "兩人先一起完成4張卡的位置。口說時，A依序指出兩個構件並問 What’s this／that?，B用 It’s a ... 回答；遇到多條軌道改問 What are these?，回答 They’re tracks.。完成兩張卡後交換提問與回答。最後向鄰組說明一個相同處或不同處。",
                "兩人能把轉車台、放射狀軌道、車庫門與火車頭4張名稱卡放在合理位置；每人至少完成1次提問和1次回答，並能用中文說出一項判斷線索。教師不評專有詞拼寫。",
            ),
            stage(
                "Wrap-up",
                5,
                "看扇形車庫全景圖，指出三個主要構件",
                "每人桌上放1張A5無文字全景圖。教師從4個名稱中說出3個，一次說1個並留3秒；全班各自在自己的圖上指出位置，教師巡看並快速勾記。接著抽點學生，請他指著投影完成1次英文指認。最後提醒：今天是看影片和圖片認識扇形車庫，不是親自到現場。",
                "聽到教師說出名稱後，在自己的全景圖上指出正確位置；完成3次指認後，和鄰座互相檢查。被抽點時，以 It’s a ... 或 They’re tracks. 回答。",
                "個人能在3次指令中至少正確指出2個構件，並能口說至少1個英文名稱。教師依勾記分成「能獨立辨識／需要名稱卡支援／第2節先複習」；不評鐵道專有詞拼寫。",
            ),
        ],
    }


def add_lesson_without_page_break(doc, lesson):
    add_lesson_title(doc, lesson["title"])
    add_body(doc, lesson["focus"], bold_label="本節焦點｜", fill=PALE_TEAL)
    add_heading(doc, "課前教學配置", 2)
    add_kv_table(
        doc,
        [
            ("主要英語口說／句型", lesson["patterns"]),
            ("句型來源與使用界線", lesson["pattern_source"]),
            ("生活用語", lesson["daily"]),
            ("核心英文字詞", lesson["core"]),
            ("相關英文字詞＋中文", lesson["related"]),
            ("數位教材", lesson["digital"]),
            ("手作教材", lesson["handmade"]),
            ("本節學習單", lesson["worksheet"]),
        ],
    )
    add_heading(doc, "40分鐘詳細教學流程", 2)
    flow_rows = [
        [
            f"{item['name']}\n{item['minutes']}分鐘",
            item["title"],
            item["teacher"],
            item["students"],
            item["assessment"],
        ]
        for item in lesson["stages"]
    ]
    add_grid_table(
        doc,
        ["階段／時間", "主要活動標題", "教師如何教與如何提問", "學生如何學與如何互動", "形成性評量／成果"],
        flow_rows,
        [1100, 1400, 2900, 2500, 1460],
        font_size=12.0,
    )
    total = sum(item["minutes"] for item in lesson["stages"])
    if total != 40:
        raise ValueError(f"Lesson stages total {total} minutes")
    add_body(
        doc,
        "本節英語口說以『能在地方學習任務中聽懂、指出、輪流說』為達成標準；不把專有名詞拼寫或尚未正式教過的文法列為必要通過條件。",
        bold_label="四年級適切性檢核｜",
        fill=PALE_AMBER,
    )


def build():
    if not REFERENCE.exists():
        raise FileNotFoundError(REFERENCE)
    doc = Document(REFERENCE)
    clear_body(doc)
    doc.core_properties.title = "〈扇形車庫〉第1節四階段詳案樣稿"
    doc.core_properties.subject = "四年級在地課程｜第1節清楚活動比較版"
    doc.core_properties.author = ""

    p = doc.add_paragraph()
    set_paragraph_format(p, before=10, after=8, line=1.0)
    set_run_font(p.add_run("LOCAL RAILWAY × ENGLISH"), 10, True, TEAL)
    p = doc.add_paragraph()
    set_paragraph_format(p, before=22, after=10, line=1.0)
    set_run_font(p.add_run("四年級在地課程教案\n〈扇形車庫〉第1節樣稿"), 26, True, NAVY)
    p = doc.add_paragraph()
    set_paragraph_format(p, after=18, line=1.2)
    set_run_font(p.add_run("活動先行 × 語言融入 × 證據評量"), 14, True, TEAL)
    add_metadata_table(doc)
    add_body(
        doc,
        "此檔是骨架確認後的第1節完整樣稿，不是整個單元完成版。現行四單元20節Word仍作基線，未被覆寫。",
        bold_label="樣稿界線｜",
        fill=PALE_BLUE,
    )
    add_body(
        doc,
        "本版將難以想像的活動術語改為教師可直接執行的步驟：影片中找四個部分、兩人放置四張名稱卡、個別在全景圖上指出三個構件。",
        bold_label="本版修改｜",
        fill=PALE_TEAL,
    )
    add_body(
        doc,
        "因交通、安全、時間及經費因素取消實地踏查，改用官方影片、官方照片與俯視簡圖。影像用於觀察與找證據，不等同親身現場經驗。",
        bold_label="已確認替換｜",
        fill=PALE_AMBER,
    )

    add_heading(doc, "一、已確認的設計決定", 1)
    decisions = [
        ["活動拆分", "第1節影片找構件；第2節討論、學習單與發表。"],
        ["替代教材", "使用教師指定影片、4張構件局部圖與1張無文字全景圖；名稱在學習單或Kahoot版面疊加。"],
        ["活動操作", "Production改為兩人合作放置4張中英名稱卡，不設四人角色、證據編號或輪換任務。"],
        ["英語融入", "單數指認取自三年級計畫；複數指認與位置問答取自四年級計畫。語言只在活動功能需要時使用。"],
        ["評量順序", "在地構件辨識、觀察與證據為主；英語看圖指認及可理解口說為輔。"],
        ["課綱結構", "英語、社會、綜合活動課綱只在單元層級集中列入，不在四階段逐列代碼。"],
    ]
    add_grid_table(doc, ["決定面向", "確認內容"], decisions, [1900, 7460], font_size=9.1)

    add_heading(doc, "二、單元層級課綱", 1)
    unit = UNITS[0]
    curriculum_rows = [
        [
            "英語文",
            format_curriculum_entries(unit["curriculum"]["english"]["performance"]),
            format_curriculum_entries(unit["curriculum"]["english"]["content"]),
            "聽懂並回應指認指令；依單數／複數畫面完成簡易問答；以圖卡、手勢與情境非語言訊息協助理解。",
        ],
        [
            "社會",
            format_curriculum_entries(unit["curriculum"]["social"]["performance"]),
            format_curriculum_entries(unit["curriculum"]["social"]["content"]),
            "從影片與官方圖文摘取構件及其空間關係，說明鐵道設施如何反映地方交通歷史，並在小組中聆聽與表達。",
        ],
        [
            "綜合活動",
            format_curriculum_entries(unit["curriculum"]["comprehensive"]["performance"]),
            format_curriculum_entries(unit["curriculum"]["comprehensive"]["content"]),
            "運用看圖、指出位置、兩人放置名稱卡及互相檢查完成共同任務，體會合作及鐵道文化與地方生活的關係。",
        ],
    ]
    add_grid_table(
        doc,
        ["領域", "學習表現（代碼＋完整敘述）", "學習內容（代碼＋完整敘述）", "本單元活動與評量證據"],
        curriculum_rows,
        [1000, 3100, 2800, 2460],
        font_size=8.4,
    )

    add_heading(doc, "三、第1節學習目標與可見成果", 1)
    goals = [
        "能從影片、官方照片或俯視簡圖指出至少3項主要構件，並說出至少1項可見證據。",
        "能在兩人活動中把轉車台、放射狀軌道、車庫門與火車頭4張中英名稱卡放到全景圖的合理位置。",
        "能看圖口說至少2個核心詞，並完成至少1組單數或複數指認問答。",
        "能區分影像觀察與親自踏查，不宣稱自己到過現場。",
    ]
    for goal in goals:
        add_real_bullet(doc, goal)
    add_body(
        doc,
        "兩人A3名稱卡配置圖＋個人A5無文字全景圖指認紀錄。",
        bold_label="本節成果｜",
        fill=PALE_TEAL,
    )

    add_heading(doc, "四、第1節40分鐘四階段詳案", 1)
    add_lesson_without_page_break(doc, lesson_payload())

    add_heading(doc, "五、評量判準、差異化與備援", 1)
    assessment_rows = [
        ["主要：在地內容", "指出至少3項構件；4張名稱卡位置合理；能說1項可見線索。", "兩人名稱卡配置圖、個人全景圖指認"],
        ["主要：觀察判斷", "能依影片或全景圖中的形狀、線條及位置說明判斷，不只憑印象猜測。", "教師巡視勾記與學生中文口頭說明"],
        ["輔助：英語功能", "看圖說至少2詞；完成1組單數或複數指認，語意可理解。", "個人口說勾核；不評專有詞拼寫"],
    ]
    add_grid_table(doc, ["層級", "通過表現", "蒐集證據"], assessment_rows, [1800, 4700, 2860], font_size=9.0)
    for item in [
        "需要支援：提供圖卡、中文詞與英文口說音節提示；可先指認再口說。",
        "一般任務：依圖片數量選用單數或複數問答，並說一項中文可見線索。",
        "延伸任務：用中文解釋中央轉車台與放射狀軌道的空間關係；不增加未確認的歷史敘述。",
        "設備備援：影片無法播放時，依同一順序使用預先下載的官方停格圖；學習目標與證據要求不變。",
    ]:
        add_real_bullet(doc, item)

    add_heading(doc, "六、原始來源—設計決定—教學活動—評量證據", 1)
    trace_rows = [
        ["原目標『構成要件』＋活動『影片觀賞』", "D-02、D-03、D-04、D-11", "局部圖、分段影片、兩人名稱卡配置", "指出至少3項構件＋1項可見線索"],
        ["三年級計畫 What’s this/that?", "D-06", "近距圖卡與遠距投影指認", "完成1組單數問答"],
        ["四年級計畫 What are these/those? They’re...", "D-12", "全景辨識多條放射狀軌道", "依複數畫面完成問答"],
        ["專有詞口說標籤原則", "D-09", "圖卡、手勢、指出位置", "聽說可理解；不評拼寫"],
    ]
    add_grid_table(doc, ["原始來源", "設計決定", "教學活動", "評量證據"], trace_rows, [2400, 1200, 3000, 2760], font_size=8.8)

    add_heading(doc, "七、影片、照片與資料來源", 1)
    add_resource_table(doc, unit["source_keys"])
    add_body(
        doc,
        "授課前重新開啟來源確認內容與網址；只使用教學所需片段或照片，保留來源標示。俯視簡圖須依可信資料人工核對構件位置，不以AI圖取代事實查證。",
        bold_label="來源使用界線｜",
        fill=PALE_AMBER,
    )

    add_heading(doc, "八、教師樣稿審查點", 1)
    for item in [
        "Warm-up的標題與步驟是否清楚，並能在5分鐘內完成『中文線索＋英文名稱』。",
        "Presentation是否能讓教師直接依『播放、暫停、提問、指位置、貼名稱卡、全班口說』操作；正式秒數留待教材階段核對。",
        "Production的兩人名稱卡活動是否適合班級人數與座位；每人是否都有放卡、提問及回答的機會。",
        "Wrap-up是否能在5分鐘內讓每位學生指出3個構件，並留下教師可快速勾記的結果。",
        "形成性評量是否先看在地構件與觀察線索，再看英語功能表現。",
        "確認本樣稿後，才依同一原則展開第2節與30秒小組導覽。",
    ]:
        add_real_bullet(doc, item)

    FINAL.parent.mkdir(parents=True, exist_ok=True)
    doc.save(FINAL)
    print(FINAL)


if __name__ == "__main__":
    build()
