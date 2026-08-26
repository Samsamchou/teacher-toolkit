from pathlib import Path
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


SCRIPT_DIR = Path(__file__).resolve().parent
UNIT_DIR = SCRIPT_DIR.parents[1]
ROOT = SCRIPT_DIR.parents[4]
REFERENCE = ROOT / "在地課程4年級上學期教案" / "01_第1-2週_扇形車庫.docx"
FINAL = Path(
    os.environ.get(
        "LESSON2_SAMPLE_OUTPUT",
        str(UNIT_DIR / "12_扇形車庫_第2節四階段詳案樣稿_審查標題與七類教材草圖版_20260826.docx"),
    )
)
SKETCH_DIR = Path(
    os.environ.get(
        "LESSON2_SKETCH_DIR",
        str(UNIT_DIR / "13_扇形車庫_第2節七類教材草圖與風格比較_20260826"),
    )
)

sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(SCRIPT_DIR))

from build_grade4_sem1_docs import (  # noqa: E402
    HEADER_FILL,
    NAVY,
    PALE_AMBER,
    PALE_BLUE,
    PALE_TEAL,
    TEAL,
    WHITE,
    add_body,
    add_grid_table,
    add_heading,
    add_kv_table,
    add_lesson_title,
    add_real_bullet,
    add_text_cell,
    format_curriculum_entries,
    prevent_row_split,
    set_cell_shading,
    set_paragraph_format,
    set_repeat_table_header,
    set_run_font,
    set_table_geometry,
)
from grade4_sem1_lesson_data import UNITS  # noqa: E402
from lesson2_data import DECISIONS, GOALS, LESSON, SOURCES, TRACE_ROWS  # noqa: E402


def clear_body(doc):
    body = doc._body._element
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)


def reset_page_number_footer(doc):
    """Replace the inherited simple PAGE field with a stable complex field."""
    paragraph = doc.sections[0].footer.paragraphs[0]
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(paragraph, before=0, after=0, line=1.0)
    set_run_font(paragraph.add_run("第 "), 8.5, color="5B6B77")

    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    set_run_font(paragraph.add_run("1"), 8.5, color="5B6B77")

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    set_run_font(paragraph.add_run(" 頁"), 8.5, color="5B6B77")


def add_metadata_table(doc):
    rows = [
        ["適用單元", "第1–2週〈扇形車庫〉", "本檔範圍", "第2節完整樣稿"],
        ["授課時間", "40分鐘", "對象", "二水國小四年級"],
        ["前序成果", "第1節詳案與教材已完成QA", "本稿狀態", "草稿待教師確認"],
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for idx, value in enumerate(("項目", "內容", "項目", "內容")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        add_text_cell(cell, value, 9.2, True, WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(row):
            if idx in (0, 2):
                set_cell_shading(cells[idx], HEADER_FILL)
                add_text_cell(cells[idx], value, 9.3, True, NAVY)
            else:
                add_text_cell(cells[idx], value, 9.3)
    set_table_geometry(table, [1350, 3330, 1350, 3330])
    doc.add_paragraph()


def add_lesson(doc):
    add_lesson_title(doc, LESSON["title"])
    add_body(doc, LESSON["focus"], bold_label="本節焦點｜", fill=PALE_TEAL)
    add_heading(doc, "課前教學配置", 2)
    add_kv_table(
        doc,
        [
            ("主要英語口說／句型", LESSON["patterns"]),
            ("句型來源與使用界線", LESSON["pattern_source"]),
            ("四色圖例英語", LESSON["color_pattern"]),
            ("生活用語", LESSON["daily"]),
            ("核心英文字詞", LESSON["core"]),
            ("相關英文字詞＋中文", LESSON["related"]),
            ("數位教材需求", LESSON["digital"]),
            ("手作教材需求", LESSON["handmade"]),
            ("本節學習單需求", LESSON["worksheet"]),
        ],
    )
    add_body(
        doc,
        "以上只定義本節所需教材及課堂功能；詳案獲教師確認後才另建教材草稿，不在本檔生成正式學習單、資料卡或投影片。",
        bold_label="教材關卡｜",
        fill=PALE_AMBER,
    )
    add_heading(doc, "40分鐘詳細教學流程", 2)
    rows = [
        [
            f"{item['name']}\n{item['minutes']}分鐘",
            f"{item['activity_heading']}\n{item['activity_description']}",
            item["teacher"],
            item["students"],
            item["assessment"],
        ]
        for item in LESSON["stages"]
    ]
    table = add_grid_table(
        doc,
        ["階段／時間", "主要活動標題", "教師如何教與如何提問", "學生如何學與如何互動", "形成性評量／成果"],
        rows,
        [1100, 1400, 2900, 2500, 1460],
        font_size=11.5,
    )
    for row, item in zip(table.rows[1:], LESSON["stages"]):
        cell = row.cells[1]
        cell.text = ""
        p = cell.paragraphs[0]
        set_paragraph_format(p, after=3, line=1.12)
        set_run_font(p.add_run(item["activity_heading"]), 11.5, True, NAVY)
        p = cell.add_paragraph()
        set_paragraph_format(p, after=0, line=1.16)
        set_run_font(p.add_run(item["activity_description"]), 11.5)
    total = sum(item["minutes"] for item in LESSON["stages"])
    if total != 40:
        raise ValueError(f"Lesson stages total {total} minutes")
    add_body(
        doc,
        "本節在地知識與官方證據是主要通過條件；英語只要求能在看圖、指位置及輪流導覽時完成可理解的簡短口說，不評鐵道專有詞拼寫。",
        bold_label="四年級適切性檢核｜",
        fill=PALE_AMBER,
    )


def add_material_sketches(doc):
    sketches = [
        (
            "01_個人解密單_詳案理解草圖.png",
            "圖1｜A4個人《扇形車庫解密單》",
            "四區依序是由來填空、三組構件—功能連線、依四色圖例定位，以及Where’s the turntable?句框。",
            3.75,
        ),
        (
            "02_四色圖例卡_詳案理解草圖.png",
            "圖2｜A5四色圖例卡",
            "紅、藍、綠、黃分別對應轉車台、軌道、車庫及火車頭；顏色只用來協助讀圖。",
            5.2,
        ),
        (
            "03_三格出口票_詳案理解草圖.png",
            "圖3｜A6三格出口票",
            "三格分別蒐集由來、轉車台功能與位置答案；各格都有語意圖示與不洩漏完整答案的句首提示。",
            5.2,
        ),
        (
            "04_何時構件功能三格圖示板_詳案理解草圖.png",
            "圖4｜A4橫式『何時—構件—功能』三格圖示板",
            "時鐘、放大鏡與齒輪圖示分別代表何時、構件與功能；小組把8張中文資料句條放入對應區。",
            6.25,
        ),
        (
            "05_中文資料句條8張_詳案理解草圖.png",
            "圖5｜A4裁切頁中文資料句條8張",
            "何時2張、構件3張、功能3張；學生版正面不印分類，教師代碼可放背面。",
            5.2,
        ),
        (
            "06_英文句型卡4張_詳案理解草圖.png",
            "圖6｜A5裁切頁英文句型卡4張",
            "兩張問句卡與兩張回答卡只支援構件辨認和位置表達，功能仍以中文說明。",
            5.2,
        ),
        (
            "07_30秒導覽順序卡_詳案理解草圖.png",
            "圖7｜A5的30秒導覽順序卡",
            "四人依『由來—英文提問—英文回答—功能』順序接力，附8、6、6、10秒的時間提示。",
            5.2,
        ),
    ]
    doc.add_page_break()
    add_heading(doc, "六、七類手作教材內容與操作草圖（詳案確認用）", 1)
    add_body(
        doc,
        "以下七張草圖用來說明各教材的尺寸、人數、欄位、題目、圖像鷹架、學生操作順序及教師答案；"
        "角色仍是中性結構占位，不是可列印正式教材。教師確認內容並選定視覺風格後，"
        "仍須另做學生版與教師解答版教材草稿並再次確認。",
        bold_label="草圖界線｜",
        fill=PALE_AMBER,
    )
    for index, (filename, title, note, width_inches) in enumerate(sketches):
        heading = add_heading(doc, title, 2)
        if index:
            heading.paragraph_format.page_break_before = True
        add_body(doc, note)
        path = SKETCH_DIR / filename
        if not path.exists():
            raise FileNotFoundError(path)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(path), width=Inches(width_inches))
        shape._inline.docPr.set("descr", title)
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(caption, before=3, after=4, line=1.0)
        set_run_font(caption.add_run(f"{title}｜詳案理解用／非正式教材"), 9.0, True, TEAL)


def add_style_board(doc):
    heading = add_heading(doc, "七、火車小偵探五風格比較板（教師選擇用）", 1)
    heading.paragraph_format.page_break_before = True
    add_body(
        doc,
        "比較板使用同一位原創火車小偵探及同一個放大鏡動作，依序比較：3D Q版溫暖手繪、"
        "3D Q版剪紙、3D家庭動畫電影感、溫暖日式手繪動畫感、日式Q版動漫。"
        "只採一般化視覺特徵，不模仿特定公司角色或既有畫面；教師選定其中一種後，才製作正式教材美術版。",
        bold_label="選擇關卡｜",
        fill=PALE_AMBER,
    )
    path = SKETCH_DIR / "08_火車小偵探五風格比較板_教師選擇用.png"
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(path), width=Inches(6.25))
    shape._inline.docPr.set("descr", "同一位原創火車小偵探的五種一般化視覺風格比較板")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(caption, before=3, after=4, line=1.0)
    set_run_font(caption.add_run("圖8｜火車小偵探五風格比較板｜教師選擇用，非正式教材"), 9.0, True, TEAL)


def build():
    if not REFERENCE.exists():
        raise FileNotFoundError(REFERENCE)
    doc = Document(REFERENCE)
    clear_body(doc)
    reset_page_number_footer(doc)
    doc.core_properties.title = "〈扇形車庫〉第2節四階段詳案樣稿｜審查標題與七類教材草圖版"
    doc.core_properties.subject = "四年級在地課程｜第2節草稿待教師確認"
    doc.core_properties.author = ""

    p = doc.add_paragraph()
    set_paragraph_format(p, before=10, after=8, line=1.0)
    set_run_font(p.add_run("LOCAL RAILWAY × ENGLISH"), 10, True, TEAL)
    p = doc.add_paragraph()
    set_paragraph_format(p, before=22, after=10, line=1.0)
    set_run_font(p.add_run("四年級在地課程教案\n〈扇形車庫〉第2節樣稿"), 26, True, NAVY)
    p = doc.add_paragraph()
    set_paragraph_format(p, after=18, line=1.2)
    set_run_font(p.add_run("審查型活動標題 × 七類教材草圖 × 五風格比較"), 14, True, TEAL)
    add_metadata_table(doc)
    add_body(
        doc,
        "此檔是第2節詳案樣稿，已承接並稽核既有來源追溯、設計決定、2節骨架及第1節成果；不是教材完成版。",
        bold_label="前向測試界線｜",
        fill=PALE_BLUE,
    )
    add_body(
        doc,
        "本次新產出停在教師確認關卡。七類附圖是內容與操作草圖，五風格比較板只供選擇；"
        "教師確認活動、草圖內容及一種視覺風格後，才建立第2節正式學生版與教師解答版教材草稿。",
        bold_label="下一關卡｜",
        fill=PALE_AMBER,
    )

    add_heading(doc, "一、正式來源與前序成果稽核", 1)
    audit_rows = [
        ["正式課程", "《4年級 校定課程.pdf》第12頁；課名為在地課程，第1–2週各1節，共2節。", "通過"],
        ["原始目標與活動", "由來／構成要件；影片觀賞、分組討論與提問回答、書寫學習單；發表問答。", "通過"],
        ["追溯與設計決定", "01_來源追溯與設計決定.md；D-02、D-07、D-08、D-09、D-12均已確認。", "通過"],
        ["單元骨架", "02_扇形車庫_2節骨架.md；第2節為由來—構件—功能、學習單、短講與問答。", "通過"],
        ["第1節成果", "詳案、影片／Kahoot、圖片教材及QA既有成果只讀回，不重做。", "通過"],
        ["本次confirmed RDQ", "RDQ-spec-fan-roundhouse-lesson2-review-title-material-visuals-20260826.md；8張中文句條、4張英文句型卡與尺寸已確認。", "通過"],
    ]
    add_grid_table(doc, ["稽核項目", "證據與結論", "結果"], audit_rows, [1800, 6260, 1300], font_size=9.1)

    add_heading(doc, "二、第2節設計決定", 1)
    add_grid_table(doc, ["決定面向", "本樣稿具體化內容"], DECISIONS, [1900, 7460], font_size=9.1)

    add_heading(doc, "三、單元層級課綱", 1)
    unit = UNITS[0]
    curriculum_rows = [
        [
            "英語文",
            format_curriculum_entries(unit["curriculum"]["english"]["performance"]),
            format_curriculum_entries(unit["curriculum"]["english"]["content"]),
            "聽懂構件與位置指令；以單數／複數或Where’s問答完成圖像指認；運用圖例、手勢及同伴互動協助理解。",
        ],
        [
            "社會",
            format_curriculum_entries(unit["curriculum"]["social"]["performance"]),
            format_curriculum_entries(unit["curriculum"]["social"]["content"]),
            "從官方圖文摘取1922至1933年發展、轉車台及股道功能，說明鐵道設施與地方交通歷史的關係。",
        ],
        [
            "綜合活動",
            format_curriculum_entries(unit["curriculum"]["comprehensive"]["performance"]),
            format_curriculum_entries(unit["curriculum"]["comprehensive"]["content"]),
            "先完成個人解密單，再在四人導覽練習中依序發言與檢核，並由自願小組上台分享，體會合作及地方鐵道文化。",
        ],
    ]
    add_grid_table(
        doc,
        ["領域", "學習表現（代碼＋完整敘述）", "學習內容（代碼＋完整敘述）", "本單元活動與評量證據"],
        curriculum_rows,
        [1000, 3100, 2800, 2460],
        font_size=8.4,
    )

    add_heading(doc, "四、第2節學習目標與可見成果", 1)
    for goal in GOALS:
        add_real_bullet(doc, goal)
    add_body(
        doc,
        "個人《扇形車庫解密單》＋三格出口票＋四人30秒導覽練習；上台分享採自願。",
        bold_label="本節成果｜",
        fill=PALE_TEAL,
    )

    add_heading(doc, "五、第2節40分鐘四階段詳案", 1)
    add_lesson(doc)
    add_material_sketches(doc)
    add_style_board(doc)

    add_heading(doc, "八、評量判準、差異化與備援", 1)
    assessment_rows = [
        ["主要：由來", "能說出1922年啟用，並知道股道後來增加到12股。", "個人解密單與出口票"],
        ["主要：構件與功能", "至少2組構件—功能正確，且能依圖說明轉車台、軌道與車庫的關係。", "個人解密單、教師巡視口頭說明"],
        ["主要：小組表達", "30秒導覽練習含由來、構件、功能；四人都有一段口說任務。上台採自願。", "教師巡視組內練習；自願上台組另給口頭回饋"],
        ["輔助：英語功能", "完成1組可理解的構件指認、複數或位置問答。", "導覽口說；不評專有詞拼寫"],
    ]
    add_grid_table(doc, ["層級", "通過表現", "蒐集證據"], assessment_rows, [1800, 4700, 2860], font_size=9.0)
    for item in [
        "需要支援：提供三格圖示、中文資料句條及英文句型卡；可先指出位置再口說。",
        "一般任務：完成1項由來、2組構件—功能及1組位置問答，並參與組內30秒導覽練習。",
        "延伸任務：用中文比較『1922年初設6股道』與『1933年完成12股道』，說明資料中的先後變化。",
        "設備備援：無法投影時，改用預先列印的官方全景圖與資料句條；目標與評量不變。",
    ]:
        add_real_bullet(doc, item)

    add_heading(doc, "九、原始來源—設計決定—教學活動—評量證據", 1)
    add_grid_table(doc, ["原始來源", "設計決定", "教學活動", "評量證據"], TRACE_ROWS, [2400, 1200, 3000, 2760], font_size=8.7)

    add_heading(doc, "十、官方資料來源", 1)
    add_grid_table(doc, ["資源名稱", "教學用途與使用提醒", "路徑／網址"], SOURCES, [2300, 4000, 3060], font_size=8.6)
    add_body(
        doc,
        "正式教材須保留來源標示與查核日期；只擷取教學所需資料。轉車台、軌道與車庫關係須依官方影像人工核對，不以AI圖取代事實查證。",
        bold_label="來源使用界線｜",
        fill=PALE_AMBER,
    )

    add_heading(doc, "十一、教師樣稿審查點", 1)
    for item in [
        "四階段正式標題是否讓課程審查委員只看標題就能辨識活動內容、學生動作與教學目標。",
        "Presentation的『何時—構件—功能』三格圖示、8張中文句條及4張英文句型卡，是否能在10分鐘完成。",
        "七類教材草圖是否已足以看懂學生的排序、書寫、連線、定位、回答及導覽步驟。",
        "各小題的火車小偵探動作、功能圖示與句首提示是否適量，且不直接洩漏完整答案。",
        "五風格比較板要選哪一種作為整套正式教材風格。",
        "Production是否採8分鐘個人作答、7分鐘四人共編、5分鐘自願小組上台；上台不列為個人成績。",
    ]:
        add_real_bullet(doc, item)

    FINAL.parent.mkdir(parents=True, exist_ok=True)
    doc.save(FINAL)
    print(FINAL)


if __name__ == "__main__":
    build()
