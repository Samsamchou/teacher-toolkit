from pathlib import Path
import hashlib
import json
import re
import sys

from docx import Document
from docx.oxml.ns import qn


SCRIPT_DIR = Path(__file__).resolve().parent
UNIT_DIR = SCRIPT_DIR.parents[1]
ROOT = SCRIPT_DIR.parents[4]
FINAL = (
    Path(sys.argv[1])
    if len(sys.argv) > 1
    else UNIT_DIR / "12_扇形車庫_第2節四階段詳案樣稿_審查標題與七類教材草圖版_20260826.docx"
)
REFERENCE = ROOT / "在地課程4年級上學期教案" / "01_第1-2週_扇形車庫.docx"
EXPECTED_REFERENCE_SHA = "3779FB5E553287DF55027EBD64A26DB1F6809997520CE639373FF043FA0EB759"


def sha256(path):
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        parts.extend(p.text for p in section.header.paragraphs)
        parts.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(parts)


def table_geometry(table):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    width = int(tbl_w.get(qn("w:w"))) if tbl_w is not None else None
    grid = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
    return {"width": width, "grid_sum": sum(grid), "columns": len(grid)}


def main():
    assert FINAL.exists(), FINAL
    doc = Document(FINAL)
    text = all_text(doc)
    required = [
        "草稿待教師確認",
        "Warm-up",
        "Presentation",
        "Production",
        "Wrap-up",
        "1922年啟用初期6股道",
        "1933年形成12股道",
        "Where’s the turntable?",
        "It’s in the middle.",
        "What are these? They’re tracks.",
        "配對扇形車庫構件，複習名稱與空間位置",
        "兩人把4張構件名稱卡放回第1節全景圖",
        "分類扇形車庫年代、構件與功能，理解運作關係",
        "何時—構件—功能",
        "8張中文資料句條",
        "4張英文句型卡",
        "完成個人解密單，合作進行30秒扇形車庫導覽",
        "先依四區草圖完成個人解密單",
        "完成三格出口票，檢核由來、功能與位置",
        "個別寫下由來、功能與位置三項答案",
        "自願小組上台",
        "上台發表採自願，不以是否上台作為個人成績",
        "七類手作教材內容與操作草圖（詳案確認用）",
        "火車小偵探五風格比較板（教師選擇用）",
        "3D Q版溫暖手繪",
        "3D Q版剪紙",
        "3D家庭動畫電影感",
        "溫暖日式手繪動畫感",
        "日式Q版動漫",
        "中央轉車台可以……",
        "詳案理解用／非正式教材",
        "詳案獲教師確認後才另建教材草稿",
        "不評鐵道專有詞拼寫",
        "原始來源—設計決定—教學活動—評量證據",
    ]
    for token in required:
        assert token in text, token
    forbidden = [
        "1922年建有12股道",
        "第2節教材已完成",
        "本節教材已完成",
        "教師已確認第2節",
        "把相鄰兩組配對",
        "兩組各導覽30秒並互問1題",
        "構件找回家",
        "三格資料解密",
        "解密單到小導覽",
        "三格出口票收束",
        "何時啟用—怎麼運作—做什麼",
    ]
    for token in forbidden:
        assert token not in text, token

    flow_table = None
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == "階段／時間":
            flow_table = table
            break
    assert flow_table is not None
    assert len(flow_table.rows) == 5
    minutes = []
    flow_font_sizes = []
    for row in flow_table.rows[1:]:
        match = re.search(r"(\d+)分鐘", row.cells[0].text)
        assert match, row.cells[0].text
        minutes.append(int(match.group(1)))
    for row in flow_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        assert run.font.size is not None, run.text
                        flow_font_sizes.append(round(run.font.size.pt, 1))
    assert minutes == [5, 10, 20, 5]
    assert sum(minutes) == 40
    assert set(flow_font_sizes) == {11.5}, sorted(set(flow_font_sizes))
    expected_titles = [
        "配對扇形車庫構件，複習名稱與空間位置",
        "分類扇形車庫年代、構件與功能，理解運作關係",
        "完成個人解密單，合作進行30秒扇形車庫導覽",
        "完成三格出口票，檢核由來、功能與位置",
    ]
    for row, expected in zip(flow_table.rows[1:], expected_titles):
        title_cell = row.cells[1]
        nonempty = [p.text.strip() for p in title_cell.paragraphs if p.text.strip()]
        assert len(nonempty) >= 2, (expected, nonempty)
        assert nonempty[0] == expected, (expected, nonempty)

    image_descriptions = []
    for shape in doc.inline_shapes:
        image_descriptions.append(shape._inline.docPr.get("descr", ""))
    assert len(doc.inline_shapes) == 8, len(doc.inline_shapes)
    for expected in [
        "個人《扇形車庫解密單》",
        "四色圖例卡",
        "三格出口票",
        "三格圖示板",
        "中文資料句條8張",
        "英文句型卡4張",
        "30秒導覽順序卡",
        "五種一般化視覺風格比較板",
    ]:
        assert any(expected in description for description in image_descriptions), (expected, image_descriptions)

    section = doc.sections[0]
    geometry = {
        "page_width_in": round(section.page_width.inches, 2),
        "page_height_in": round(section.page_height.inches, 2),
        "left_margin_in": round(section.left_margin.inches, 2),
        "right_margin_in": round(section.right_margin.inches, 2),
        "top_margin_in": round(section.top_margin.inches, 2),
        "bottom_margin_in": round(section.bottom_margin.inches, 2),
    }
    assert geometry == {
        "page_width_in": 8.5,
        "page_height_in": 11.0,
        "left_margin_in": 1.0,
        "right_margin_in": 1.0,
        "top_margin_in": 1.0,
        "bottom_margin_in": 1.0,
    }
    table_checks = [table_geometry(table) for table in doc.tables]
    assert all(item["width"] == 9360 for item in table_checks)
    assert all(item["grid_sum"] == 9360 for item in table_checks)
    assert sha256(REFERENCE) == EXPECTED_REFERENCE_SHA
    result = {
        "final": str(FINAL),
        "bytes": FINAL.stat().st_size,
        "sha256": sha256(FINAL),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "flow_minutes": minutes,
        "flow_total": sum(minutes),
        "flow_font_sizes_pt": sorted(set(flow_font_sizes)),
        "page_geometry": geometry,
        "table_geometry": table_checks,
        "reference_sha256": sha256(REFERENCE),
        "required_tokens": "PASS",
        "forbidden_claims": "PASS",
        "teacher_gate": "PASS",
        "inline_images": len(doc.inline_shapes),
        "material_sketches": 7,
        "style_board": 1,
        "image_alt_text": image_descriptions,
        "activity_title_structure": "PASS",
        "voluntary_presentation": "PASS",
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
